from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "STANDARDET_E_PUNES_NE_PRIMEFLOW_V1_0_03.09.2026_PRIMEX.docx"
LOGO_PATH = ROOT / "backend" / "app" / "assets" / "primex_logo.png"

BLUE = "1F5FAE"
DARK_BLUE = "123B67"
LIGHT_BLUE = "E8F1FB"
PALE_BLUE = "F4F8FC"
TEXT = "1F2937"
MUTED = "5B6573"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D6DCE5"
WHITE = "FFFFFF"
RED = "DC2626"
DARK_RED = "991B1B"
GREEN = "C4FDC4"
PINK = "FFC4ED"
YELLOW = "FFFF00"
WAIT_CONFIRM = "FFEDD5"
WAIT_CLIENT = "E2C15B"

# Named design decision: compact_reference_guide with a PrimeFlow A4 override.
# A4, 0.5-inch margins and 0.08-inch header/footer distances reproduce the
# standard enforced by backend/app/services/word_standardizer.py.
CONTENT_DXA = 10464


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        tag = qn(f"w:{edge_name}")
        edge = tc_borders.find(tag)
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        for key, value in edge_data.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_border(paragraph, edge_name: str, color=MID_GRAY, size=8, space=2) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{edge_name}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)


def set_run_font(run, size=None, bold=None, italic=None, color=TEXT, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction: str, placeholder: str) -> None:
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)
    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    instr_run.append(instr)
    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    text_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "16")
    r_pr.append(sz)
    text_run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = placeholder
    text_run.append(text)
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph._p.extend([begin_run, instr_run, separate_run, text_run, end_run])


def set_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(18)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Prompt Text" not in [s.name for s in doc.styles]:
        prompt = doc.styles.add_style("Prompt Text", 1)
    else:
        prompt = doc.styles["Prompt Text"]
    prompt.font.name = "Consolas"
    prompt._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    prompt._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    prompt.font.size = Pt(9)
    prompt.font.color.rgb = RGBColor.from_string(TEXT)
    prompt.paragraph_format.space_before = Pt(0)
    prompt.paragraph_format.space_after = Pt(4)
    prompt.paragraph_format.line_spacing = 1.15


def configure_section(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.header_distance = Inches(0.08)
    section.footer_distance = Inches(0.08)
    section.different_first_page_header_footer = False


def build_header_footer(section) -> None:
    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(2)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(7.22), WD_TAB_ALIGNMENT.RIGHT)
    set_paragraph_border(hp, "bottom")
    if LOGO_PATH.exists():
        pic = hp.add_run().add_picture(str(LOGO_PATH), width=Inches(1.0))
        pic._inline.docPr.set("descr", "PrimEx company logo")
    else:
        r = hp.add_run("PRIMEX")
        set_run_font(r, size=11, bold=True, color=DARK_BLUE)
    hp.add_run("\t")
    add_field(hp, 'DATE \\@ "dd/MM/yyyy"', date.today().strftime("%d/%m/%Y"))

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(2)
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(7.22), WD_TAB_ALIGNMENT.RIGHT)
    set_paragraph_border(fp, "top")
    r = fp.add_run("PrimEx SH.P.K.  |  +383 49 937 863  |  info@primex.com  |  www.primexeu.com")
    set_run_font(r, size=7.5, color=MUTED)
    fp.add_run("\t")
    r = fp.add_run("Page ")
    set_run_font(r, size=8, color=MUTED)
    add_field(fp, "PAGE", "1")
    r = fp.add_run(" of ")
    set_run_font(r, size=8, color=MUTED)
    add_field(fp, "NUMPAGES", "1")


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r)


def add_callout(doc: Document, label: str, text: str, fill=LIGHT_BLUE, border=BLUE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), border)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    r = p.add_run(label + " ")
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, color=TEXT)


def add_heading(doc: Document, number: str, title: str, level=1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{number}. {title}" if number else title)
    set_run_font(r, size={1:16, 2:13, 3:12}[level], bold=True, color=BLUE if level < 3 else DARK_BLUE)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], fills: list[str] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, value in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=9.5, bold=True, color=DARK_BLUE)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if fills and row_idx < len(fills):
                set_cell_shading(cell, fills[row_idx])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, size=9.2)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def add_prompt_intro(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Prompt Text")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE_BLUE)
    p_pr.append(shd)
    r = p.add_run(text)
    set_run_font(r, size=9, bold=True, name="Consolas", color=DARK_BLUE)


def add_prompt_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Prompt Text")
    r = p.add_run(text)
    set_run_font(r, size=9, name="Consolas")


def page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_document() -> Document:
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        build_header_footer(section)
    set_update_fields(doc)
    doc.core_properties.title = "Standardet e punës në PrimeFlow"
    doc.core_properties.subject = "Manual operativ për detyrat, planifikimin dhe dokumentet"
    doc.core_properties.author = "PrimEx SH.P.K."
    doc.core_properties.keywords = "PrimeFlow, PrimEx, standarde, detyra, Word, Excel, planifikim"

    # Cover: editorial_cover adapted to PrimeFlow operating manual.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(52)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    r = kicker.add_run("MANUAL OPERATIV")
    set_run_font(r, size=10, bold=True, color=BLUE)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("STANDARDET E PUNËS\nNË PRIMEFLOW")
    set_run_font(r, size=28, bold=True, color=DARK_BLUE)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Detyrat, planifikimi, dokumentet Word/Excel dhe kontrolli para dorëzimit")
    set_run_font(r, size=14, color=MUTED)

    metadata = [
        ("Organizata", "PrimEx SH.P.K."),
        ("Platforma", "PrimeFlow"),
        ("Versioni", "1.0"),
        ("Data", "03.09.2026"),
    ]
    add_table(doc, ["Fusha", "Vlera"], [[label, value] for label, value in metadata], [2800, 7664])
    add_callout(
        doc,
        "Qëllimi:",
        "Të sigurojë që detyrat, planifikimi dhe dokumentet e kompanisë krijohen, kontrollohen dhe dorëzohen në mënyrë të njëjtë.",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("Dokument i brendshëm | Për përdorim nga ekipi PrimEx")
    set_run_font(r, size=9.5, italic=True, color=MUTED)

    page_break(doc)
    add_heading(doc, "", "Përmbajtja", 1)
    toc_items = [
        "1. Qëllimi, fusha dhe statusi i rregullave",
        "2. Standardi i titullit të detyrës",
        "3. Standardi i përshkrimit të detyrës",
        "4. Shkurtesat zyrtare të PrimeFlow",
        "5. Statuset, progresi dhe ngjyrat e detyrave",
        "6. Afatet, detyrat e reja dhe detyrat 08:00 EM",
        "7. Planifikimi ditor dhe javor",
        "8. System Tasks",
        "9. Standardi i dokumenteve Word",
        "10. Standardi i dokumenteve Excel",
        "11. Emërtimi, versionimi dhe ruajtja e skedarëve",
        "12. Kontrolli para dorëzimit",
        "13. Email-et, takimet dhe raportet",
        "14. Përdorimi i AI-së",
        "15. Përgjegjësia dhe mirëmbajtja e standardit",
        "Shtojcat A-E. Promptat operativë",
    ]
    for item in toc_items:
        add_bullet(doc, item)
    add_callout(
        doc,
        "Leximi i statusit:",
        "“Zbatuar në kod” nënkupton se rregulli ekziston në platformë. “Standard organizativ” nënkupton se rregulli duhet ndjekur nga ekipi dhe mund të kërkojë implementim ose konfigurim shtesë.",
        fill=LIGHT_GRAY,
        border=MUTED,
    )

    add_heading(doc, "1", "Qëllimi, fusha dhe statusi i rregullave", 1)
    doc.add_paragraph(
        "Ky manual përcakton mënyrën e krijimit dhe menaxhimit të detyrave në PrimeFlow, planifikimin ditor/javor, standardizimin e dokumenteve Word dhe Excel, si dhe kontrollin para dorëzimit. Qëllimi është qartësia, krahasueshmëria, gjurmueshmëria dhe zvogëlimi i gabimeve operative."
    )
    add_bullet(doc, "Rregullat e sistemit: të verifikuara në kodin aktual të PrimeFlow.")
    add_bullet(doc, "Vendimet organizative: rregulla të miratuara për përdorim të përbashkët.")
    add_bullet(doc, "Rekomandimet: praktika që duhen miratuar para se të konsiderohen detyruese.")
    add_callout(doc, "Parim:", "Standardi përmban rregulla pune, jo standarde dekorative të faqeve ose kartelave.")

    add_heading(doc, "2", "Standardi i titullit të detyrës", 1)
    add_callout(doc, "Struktura zyrtare:", "[INICIALET]: [PROJEKTI/KLIENTI]: [VEPRIMI KRYESOR]")
    doc.add_paragraph("Titulli duhet të jetë i shkurtër, i drejtpërdrejtë dhe të tregojë qartë kush e kryen punën, për cilin projekt/klient dhe cili është veprimi kryesor.")
    add_heading(doc, "2.1", "Rregullat bazë", 2)
    for text in [
        "Titulli shkruhet në një rresht dhe përmban vetëm një veprim kryesor.",
        "Përdoren inicialet e personit përgjegjës vetëm kur nevojiten në titull.",
        "Përdoret kodi zyrtar i projektit/klientit kur ekziston.",
        "Përdoren vetëm shkurtesat zyrtare të PrimeFlow.",
        "Gjatësia e rekomanduar është deri në 100 karaktere; kufiri maksimal është 120 karaktere.",
        "Hapat, listat, URL-të dhe shpjegimet teknike vendosen në Description/Notes.",
        "Lloji i detyrës nuk vendoset në titull, sepse caktohet në fushën përkatëse të sistemit.",
        "Statusi, prioriteti, AM/PM dhe data nuk vendosen në titull.",
        "Nuk shpiken shkurtesa të reja; p.sh. RREG nuk është shkurtesë zyrtare dhe duhet shkruar Rregullim.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "2.2", "Shembuj", 2)
    add_table(
        doc,
        ["Vlerësimi", "Titulli", "Arsyeja"],
        [
            ["Saktë", "RA: FRG: VERIF dokumentet e klientit", "Strukturë e shkurtër me projekt dhe veprim."],
            ["Saktë", "EK: WKF: ZGJ problemin e login-it", "Një veprim kryesor dhe kod zyrtar projekti."],
            ["Gabim", "R1: RA: FRG: VERIF dokumentet", "Lloji nuk duhet të jetë pjesë e titullit sipas këtij standardi."],
            ["Gabim", "RA: FRG: kontrollo dokumentet, dërgo email, krijo raport", "Përmban disa veprime; duhet ndarë ose kaluar në Description."],
        ],
        [1500, 4400, 4564],
    )
    add_callout(
        doc,
        "Harmonizim i nevojshëm:",
        "Auditimi aktual kontrollon disa prefikse të llojit (p.sh. R1:, P:, BLL:). Me miratimin e këtij standardi, ky validim duhet të përditësohet që lloji të ruhet vetëm në fushën e vet.",
        fill=WAIT_CONFIRM,
        border="C2410C",
    )

    add_heading(doc, "3", "Standardi i përshkrimit të detyrës", 1)
    doc.add_paragraph("Description shpjegon çfarë duhet realizuar dhe çfarë konsiderohet përfundim. Ai nuk duhet të përsërisë vetëm titullin.")
    for text in [
        "Përshkruaj rezultatin konkret që pritet.",
        "Shto hapat vetëm kur rendi i punës është i rëndësishëm.",
        "Vendos linket, dokumentet burimore dhe referencat në Description/Notes.",
        "Shëno personat ose ekipet nga të cilët varet detyra.",
        "Përcakto kriterin e përfundimit: çfarë duhet të ekzistojë, të dërgohet ose të konfirmohet.",
        "Mos vendos të dhëna të ndjeshme, kredenciale ose fjalëkalime.",
    ]:
        add_bullet(doc, text)
    add_callout(doc, "Model:", "Qëllimi: … | Veprimet: … | Rezultati: … | Varësitë: … | Referencat: …")

    add_heading(doc, "4", "Shkurtesat zyrtare të PrimeFlow", 1)
    doc.add_paragraph("Fjalori zyrtar i shkurtesave është versioni 2026.1, i përditësuar më 31.07.2026. Duhet të përdoren vetëm format e regjistruara në këtë fjalor.")
    abbreviations = [
        ("PX", "PRIMEX"), ("PV", "PUSHIM VJETOR"), ("MUNG", "MUNGESË"), ("VONS", "VONESË"),
        ("TAK EXT", "TAKIM EKSTERN"), ("TAK INT", "TAKIM INTERN"), ("BZ GA", "BARAZIM ME GANËN"),
        ("BLL", "BLLOK"), ("R1", "RAST I PARË"), ("P:", "PERSONALISHT"), ("PRJK", "PROJEKTE"),
        ("ANK", "ANKESA"), ("KRK", "KËRKESA"), ("PRZ", "PROPOZIME"), ("MBL", "MBLEDHJA"),
        ("BZ", "BARAZIM"), ("MOSBZ", "MOSBARAZIM"), ("TAK", "TAKIM"), ("JAV", "JAVOR"),
        ("CV", "COMMON VIEW"), ("PLNF", "PLANIFIKIMI"), ("RLZ", "REALIZIMI"), ("EM", "EMAIL"),
        ("DEP", "DEPARTAMENTI"), ("PCM", "PROJECT / PRODUCT CONTENT MANAGEMENT"),
        ("ZHVLL", "ZHVILLIMI"), ("GD", "GRAPHIC DESIGN"), ("ADM", "ADMINISTRATA"),
        ("RAP", "RAPORTIMI"), ("PRBL", "PROBLEME"), ("DET", "DETYRA"),
        ("DET SYS", "DETYRA TË SISTEMIT"), ("SYS", "SISTEM"), ("FT", "FAST TASK"), ("KO", "KONTROLLA"),
        ("T", "TODAY"), ("Y", "YESTERDAY"), ("O", "OVERVIEW"), ("PAG", "PAGESË"),
        ("INV", "INVOICE"), ("FAT", "FATURË"), ("KONF", "KONFIRMIM"),
        ("PF", "PRIMEFLOW / PLATFORMA"), ("DOC", "DOKUMENTE"), ("TRANS", "TRANSAKSIONE"),
        ("FINC", "FINANCA"), ("VERIF", "VERIFIKU"), ("GR", "GRUPI"),
        ("RJ", "RAPORTI JAVOR"), ("RN", "RAPORTI NESËR"), ("RS", "RAPORTI SOT"),
        ("DG", "DETYRA GRUPORE"), ("STAND", "STANDARDET"), ("J.T", "JAVËN TJETËR"),
        ("PROD", "PRODUKTE"), ("ZGJ", "ZGJIDHJE"), ("DRZ", "DORËZOHET"),
        ("APL", "APLIKANTA"), ("KOMPL", "KOMPLET"), ("REGJ", "REGJISTRATOR"),
        ("CHL", "CHECKLIST"), ("P/P", "PYETJE/PËRGJIGJE"), ("RIORG", "RIORGANIZIM"),
    ]
    # Two abbreviation pairs per row to keep the reference compact.
    abbr_rows = []
    for i in range(0, len(abbreviations), 2):
        a1 = abbreviations[i]
        a2 = abbreviations[i + 1] if i + 1 < len(abbreviations) else ("", "")
        abbr_rows.append([a1[0], a1[1], a2[0], a2[1]])
    add_table(doc, ["Shk.", "Kuptimi", "Shk.", "Kuptimi"], abbr_rows, [1100, 4132, 1100, 4132])
    add_heading(doc, "4.1", "Kodet e projekteve të njohura nga auditimi", 2)
    add_callout(doc, "Kode:", "WKF, VS, MST, ASC, TT, FRG, ARC, CRM, SMM")
    add_bullet(doc, "Kodet e panjohura nuk duhet të shpiken. Ato konfirmohen dhe pastaj shtohen në fjalorin zyrtar.")

    add_heading(doc, "5", "Statuset, progresi dhe ngjyrat e detyrave", 1)
    doc.add_paragraph("Në PrimeFlow progresi i punës komunikohet kryesisht përmes statusit. Nuk duhet të shpiket një përqindje progresi kur ajo nuk ekziston si e dhënë reale.")
    status_rows = [
        ["TODO", "Detyra nuk ka filluar", "#FFC4ED", "Rozë e hapur"],
        ["IN PROGRESS", "Detyra është në proces", "#FFFF00", "E verdhë"],
        ["WAITING CONFIRMATION", "Pret konfirmim", "#FFEDD5", "Portokalli e hapur"],
        ["WAITING CLIENT", "Pret përgjigje nga klienti", "#E2C15B", "E verdhë e errët"],
        ["DONE", "Detyra është përfunduar", "#C4FDC4", "E gjelbër e hapur"],
    ]
    add_table(doc, ["Statusi", "Kuptimi", "Kodi", "Ngjyra"], status_rows, [3000, 3900, 1600, 1964], [PINK, YELLOW, WAIT_CONFIRM, WAIT_CLIENT, GREEN])
    for text in [
        "Çdo detyrë ka vetëm një status aktiv.",
        "Ndryshimi i statusit duhet të ndryshojë automatikisht ngjyrën.",
        "Ngjyra përdoret si tregues ndihmës; emri i statusit duhet të mbetet i dukshëm.",
        "Ngjyrat e statusit përdoren njësoj në Weekly Planner, Common View dhe raportet përkatëse.",
        "Gri e hapur nuk është status zyrtar; një status i panjohur duhet korrigjuar.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "6", "Afatet, detyrat e reja dhe detyrat 08:00 EM", 1)
    add_heading(doc, "6.1", "Detyrat me afat të kaluar", 2)
    doc.add_paragraph("Detyra konsiderohet e vonuar kur due date ka kaluar dhe statusi nuk është DONE.")
    add_table(doc, ["Elementi", "Standardi"], [
        ["Sfondi", "#DC2626"], ["Border-i", "#991B1B"], ["Teksti", "I bardhë"], ["Treguesi", "AFATI KA KALUAR"],
    ], [2600, 7864])
    add_bullet(doc, "Ngjyra e vonesës ka përparësi ndaj ngjyrës së statusit.")
    add_bullet(doc, "Një detyrë DONE nuk paraqitet si e vonuar.")
    add_bullet(doc, "Start date dhe due date duhet të jenë në rend kronologjik.")

    add_heading(doc, "6.2", "Detyrat e reja", 2)
    add_bullet(doc, "Detyrë e re dhe e hapur: sfond #DBEAFE, border #1D4ED8.")
    add_bullet(doc, "Detyrë e re dhe e përfunduar: sfond #02E6C7, border #059669.")
    add_bullet(doc, "Treguesi “E RE” është i përkohshëm dhe nuk e zëvendëson statusin.")

    add_heading(doc, "6.3", "Detyrat 08:00 EM", 2)
    doc.add_paragraph("Detyrat e orës 08:00, veçanërisht kontrolli i email-eve, dallohen me border të kuq 2 px (#DC2626), tregues 08:00 me border #B91C1C dhe tekst të bardhë.")
    add_callout(doc, "Shembull:", "AK: 08:00 EM: VERIF email-et dhe kërkesat e reja")
    add_heading(doc, "6.4", "Përparësia vizuale", 2)
    for text in ["Detyra me afat të kaluar.", "Detyra e orës 08:00.", "Detyra e re.", "Ngjyra e statusit aktual."]:
        add_number(doc, text)

    add_heading(doc, "7", "Planifikimi ditor dhe javor", 1)
    add_heading(doc, "7.1", "Planifikimi ditor", 2)
    for text in [
        "Dita fillon me kontrollin 08:00 EM kur ky proces aplikohet për rolin.",
        "Detyrat me afat atë ditë dhe detyrat kritike planifikohen të parat.",
        "Takimet dhe angazhimet fikse vendosen para blloqeve fleksibile.",
        "Nuk planifikohen dy detyra në të njëjtin interval kohor.",
        "Detyrat e mëdha ndahen në hapa me rezultat të verifikueshëm.",
        "Lihet kohë rezervë për kërkesa të papritura dhe komunikim.",
        "Detyrat që presin klientin ose konfirmimin marrin statusin përkatës.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "7.2", "Planifikimi javor", 2)
    for text in [
        "Java planifikohet nga e hëna deri të premten, duke respektuar pushimet dhe mungesat.",
        "Bartja nga java e kaluar bëhet vetëm për detyra ende aktive.",
        "Ngarkesa shpërndahet në mënyrë realiste sipas afateve dhe kapacitetit.",
        "Çdo detyrë ka përgjegjës, start date, due date dhe status të vlefshëm.",
        "Detyra nuk planifikohet në ditë të Pushimit Vjetor.",
        "Kur kërkohet KO, caktohet personi përgjegjës i kontrollit.",
        "Konfliktet, varësitë dhe bllokimet evidentohen para fillimit të javës.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "7.3", "Kontrolli i realizimit", 2)
    add_bullet(doc, "Raportohen detyrat e realizuara sipas planit, në progres, pa progres, të shtyra dhe shtesë.")
    add_bullet(doc, "Kur kërkohet arsyetim për detyrë të papërfunduar, shënohen shkaku dhe hapi i ardhshëm.")
    add_bullet(doc, "Statusi DONE duhet të mbështetet nga rezultati i realizuar, jo vetëm nga ndryshimi manual i statusit.")

    add_heading(doc, "8", "System Tasks", 1)
    doc.add_paragraph("System Tasks janë detyra të krijuara ose përsëritura nga rregullat e sistemit. Ato duhet të menaxhohen përmes occurrence-it përkatës, pa ndryshuar pa nevojë modelin bazë.")
    for text in [
        "Titulli ndjek të njëjtin standard të shkurtër si detyrat e tjera.",
        "Lloji System Task ruhet në fushën përkatëse, jo në titull.",
        "Përsëritja, data dhe departamenti caktohen në konfigurimin e detyrës.",
        "Statusi i occurrence-it përditësohet sipas realizimit real.",
        "System Task me status DONE nuk redaktohet; rihapja bëhet vetëm përmes procesit të lejuar.",
        "Ndryshimi i template-it duhet të bëhet vetëm kur ndryshimi vlen për përsëritjet e ardhshme.",
    ]:
        add_bullet(doc, text)

    page_break(doc)
    add_heading(doc, "9", "Standardi i dokumenteve Word", 1)
    add_callout(doc, "Statusi:", "Rregulla të zbatuara nga Word Standardizer në PrimeFlow.")
    add_heading(doc, "9.1", "Hyrja dhe ruajtja e përmbajtjes", 2)
    for text in [
        "Pranohet vetëm formati .docx, deri në 20 MB.",
        "Dokumenti duhet të jetë i lexueshëm dhe jo bosh.",
        "Përmbajtja e trupit ruhet; standardizohen header-i, footer-i, margjinat dhe emri i skedarit.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "9.2", "Faqja, header-i dhe footer-i", 2)
    for text in [
        "Madhësia e faqes: A4.",
        "Margjinat: 0.5 inch në të katër anët.",
        "Distanca e header-it dhe footer-it: 0.08 inch.",
        "Logoja zyrtare PrimEx vendoset majtas në çdo section, me gjerësi 1 inch dhe pa shtrembërim.",
        "Data automatike vendoset djathtas në formatin dd/MM/yyyy, 8 pt, bold.",
        "Footer-i përmban të dhënat zyrtare të kompanisë.",
        "Numërimi automatik vendoset djathtas: Page {PAGE} of {NUMPAGES}.",
        "Header-i dhe footer-i janë të njëjtë në faqen e parë, faqet çift/tek dhe në të gjitha sections.",
        "Fushat automatike përditësohen kur dokumenti hapet.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "9.3", "Emri i dokumentit", 2)
    add_callout(doc, "Formati:", "[PËRSHKRIMI]_[DD.MM.YYYY]_[INICIALET].docx")
    add_bullet(doc, "Përshkrimi shkruhet me shkronja të mëdha dhe hapësirat zëvendësohen me underscore.")
    add_bullet(doc, "Hiqen karakteret e palejuara dhe përshkrimi kufizohet në 120 karaktere.")
    add_bullet(doc, "Shembull: RAPORTI_JAVOR_I_SHITJEVE_03.09.2026_AK.docx")

    add_heading(doc, "10", "Standardi i dokumenteve Excel", 1)
    add_callout(doc, "Statusi:", "Rregulla të zbatuara nga Excel Standardizer në PrimeFlow.")
    add_heading(doc, "10.1", "Hyrja dhe struktura", 2)
    for text in [
        "Pranohen .xlsx dhe .csv deri në 20 MB; CSV konvertohet në .xlsx.",
        "Workbook-u duhet të ketë të paktën një sheet me të dhëna; sheet-et bosh hiqen.",
        "Header-at bosh mbi kolona me të dhëna duhet të plotësohen nga përdoruesi; sistemi nuk i shpik.",
        "Emrat gjenerikë të sheets zëvendësohen me titull të sigurt dhe unik; referencat në formula përshtaten.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "10.2", "Titulli, header-i dhe të dhënat", 2)
    for text in [
        "Titulli vendoset në rreshtin 3, bashkohet mbi kolonat, shkruhet uppercase, Calibri 16 bold, centered dhe wrap.",
        "Header-i vendoset në rreshtin 6, Calibri 11 bold, left/bottom alignment dhe wrap.",
        "Kolona e parë është NR; shtohet ose rinumërohet në mënyrë sekuenciale dhe ka gjerësi 6.",
        "Të dhënat përdorin Calibri 11, left/bottom alignment, wrap dhe borders.",
        "Përdoren border-a të hollë brenda dhe border medium për skajin e jashtëm, header-in dhe fundin.",
        "Në fund ruhen saktësisht dy rreshta pune shtesë brenda tabelës; NR vazhdon rendin.",
        "Gjerësitë e kolonave të tjera përshtaten ndërmjet 10 dhe 45; lartësitë e rreshtave ndërmjet 18 dhe 300.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "10.3", "Numrat, formulat dhe printimi", 2)
    for text in [
        "Numrat e plotë: #,##0; numrat decimalë: #,##0.##; vlerat monetare: #,##0.00.",
        "Formulat ruhen; nuk zëvendësohen me vlera statike.",
        "Shkurtesat zyrtare aplikohen në header-a dhe etiketa kur ka përputhje të saktë.",
        "AutoFilter vendoset nga A6 deri në rreshtin e fundit të tabelës.",
        "Freeze panes vendoset në C7; rreshti 6 përsëritet në printim.",
        "Gridlines çaktivizohen. Faqja është A4; portrait deri në 7 kolona dhe landscape mbi 7 kolona.",
        "Fit to width = 1 dhe fit to height = pa kufizim.",
        "Margjinat: majtas/djathtas 0.25; sipër/poshtë 0.5; header/footer 0.2.",
        "Header-i djathtas përmban dd/MM/yyyy HH:mm; footer-i ka &P / &N në qendër dhe PUNOI: djathtas.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "10.4", "Emri i dokumentit", 2)
    add_callout(doc, "Formati:", "[PËRSHKRIMI]_[DD.MM.YYYY]_[INICIALET].xlsx")

    add_heading(doc, "11", "Emërtimi, versionimi dhe ruajtja e skedarëve", 1)
    add_callout(doc, "Statusi:", "Emërtimi Word/Excel zbatohet në kod; versionimi dhe struktura e folderëve janë standard organizativ i rekomanduar.", fill=LIGHT_GRAY, border=MUTED)
    add_bullet(doc, "Për dokumentet Word/Excel përdoret formati i standardizer-it pa shtuar fjalë si final, final2, new ose copy.")
    add_bullet(doc, "Kur nevojiten versione pune, përdoren V01, V02, V03 para inicialeve.")
    add_bullet(doc, "Versioni i dorëzuar ruhet në folderin përkatës të klientit/projektit; versionet e vjetra kalojnë në ARKIVË.")
    add_bullet(doc, "Qasja në dokumente përcaktohet sipas rolit dhe ndjeshmërisë së informacionit.")
    add_callout(doc, "Shembull me version:", "FRG_RAPORTI_JAVOR_03.09.2026_V02_AK.xlsx")

    add_heading(doc, "12", "Kontrolli para dorëzimit", 1)
    checklist = [
        "Kërkesa dhe rezultati i pritur janë realizuar.",
        "Titulli, përshkrimi, projekti, përgjegjësi dhe afati janë të sakta.",
        "Dokumenti përdor standardin e duhur Word ose Excel.",
        "Emri dhe versioni i skedarit janë të saktë.",
        "Nuk ka të dhëna testuese, dublime, formula të prishura ose faqe bosh të paqëllimshme.",
        "Dokumenti hapet pa gabime dhe është kontrolluar vizualisht.",
        "Janë hequr të dhënat konfidenciale që nuk duhet të shpërndahen.",
        "Skedari është ruajtur në vendin e duhur dhe personi përgjegjës është njoftuar.",
        "Detyra në PrimeFlow është përditësuar me statusin dhe rezultatin real.",
    ]
    for item in checklist:
        add_bullet(doc, "[ ] " + item)

    add_heading(doc, "13", "Email-et, takimet dhe raportet", 1)
    add_callout(doc, "Statusi:", "Standarde organizative të rekomanduara; nuk paraqiten si validime automatike të platformës.", fill=LIGHT_GRAY, border=MUTED)
    add_heading(doc, "13.1", "Email-et", 2)
    add_bullet(doc, "Subject-i tregon projektin/klientin, kërkesën dhe, kur duhet, datën.")
    add_bullet(doc, "Kërkesa përmban rezultatin e kërkuar, afatin dhe personin që duhet të përgjigjet.")
    add_bullet(doc, "Bashkëngjitjet emërtohen sipas standardit dhe kontrollohen para dërgimit.")
    add_bullet(doc, "Kërkesa që kërkon punë regjistrohet si detyrë në PrimeFlow.")
    add_heading(doc, "13.2", "Takimet", 2)
    add_bullet(doc, "Çdo takim ka qëllim, agjendë, pjesëmarrës dhe kohë të përcaktuar.")
    add_bullet(doc, "Në fund regjistrohen vendimi, detyra, përgjegjësi dhe afati.")
    add_bullet(doc, "Detyrat e dala nga takimi regjistrohen në PrimeFlow.")
    add_heading(doc, "13.3", "Raportet", 2)
    add_bullet(doc, "Raporti përmban periudhën, punët e realizuara, punët e hapura, problemet, rezultatet dhe hapat e ardhshëm.")
    add_bullet(doc, "Përdoren fakte dhe rezultate të matshme; shmangen formulimet e përgjithshme pa evidencë.")

    add_heading(doc, "14", "Përdorimi i AI-së", 1)
    add_callout(doc, "Statusi:", "Politikë organizative e rekomanduar që kërkon miratim të kompanisë.", fill=LIGHT_GRAY, border=MUTED)
    for text in [
        "Mos vendos fjalëkalime, kredenciale, të dhëna personale ose materiale konfidenciale në një shërbim AI pa autorizim.",
        "AI nuk duhet të shpikë klientë, data, shifra, burime ose vendime.",
        "Çdo rezultat i AI-së kontrollohet nga një person përpara përdorimit ose dërgimit.",
        "Dokumentet e prodhuara nga AI duhet të kalojnë standardin Word/Excel dhe kontrollin para dorëzimit.",
        "Vendimet ligjore, financiare, kontraktuale dhe të sigurisë kërkojnë verifikim nga personi kompetent.",
        "Ruhet versioni final i kontrolluar, jo vetëm përgjigjja e papërpunuar e AI-së.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "15", "Përgjegjësia dhe mirëmbajtja e standardit", 1)
    add_bullet(doc, "Çdo punonjës është përgjegjës për zbatimin e standardit në punën e vet.")
    add_bullet(doc, "Përgjegjësi i PrimeFlow mban fjalorin e shkurtesave dhe validimet të harmonizuara me dokumentin.")
    add_bullet(doc, "Ndryshimet miratohen, marrin version të ri dhe shoqërohen me datë.")
    add_bullet(doc, "Kur dokumenti dhe kodi bien ndesh, mospërputhja evidentohet dhe zgjidhet; nuk duhet fshehur.")
    add_table(doc, ["Versioni", "Data", "Ndryshimi", "Statusi"], [["1.0", "03.09.2026", "Krijimi i manualit të konsoliduar", "Për shqyrtim/miratim"]], [1400, 1900, 5064, 2100])

    page_break(doc)
    add_heading(doc, "", "Shtojcat - Promptat operativë", 1)
    doc.add_paragraph("Promptat më poshtë përdoren si mjete pune. Teksti në kllapa katrore zëvendësohet me të dhënat konkrete; AI duhet të kërkojë informacionin që mungon dhe nuk duhet ta shpikë.")

    add_heading(doc, "A", "Prompt për standardizimin e Word-it", 2)
    add_prompt_intro(doc, "STANDARDO DOKUMENTIN WORD SIPAS STANDARDIT PRIMEX / PRIMEFLOW")
    for line in [
        "Të dhënat: Përshkrimi [VENDOS]; Inicialet [VENDOS]; Logoja zyrtare [BASHKËNGJIT].",
        "Ruaj tekstin, tabelat, fotografitë, hyperlink-et dhe renditjen e përmbajtjes. Mos shpik dhe mos fshi të dhëna.",
        "Përdor A4; margjina 0.5 inch; header/footer 0.08 inch në çdo section.",
        "Në header vendos logon zyrtare majtas, 1 inch, me proporcione të ruajtura dhe alt text; djathtas datën automatike dd/MM/yyyy, 8 pt bold.",
        "Në footer majtas vendos PrimEx SH.P.K., +383 49 937 863, info@primex.com, www.primexeu.com; djathtas Page {PAGE} of {NUMPAGES}.",
        "Përdor të njëjtin header/footer në faqen e parë, faqet çift/tek dhe të gjitha sections. Aktivizo përditësimin e fushave kur dokumenti hapet.",
        "Emërto rezultatin [PËRSHKRIMI]_[DD.MM.YYYY]_[INICIALET].docx; përshkrimi uppercase, underscore, pa karaktere të palejuara, maksimumi 120 karaktere.",
        "Kontrollo që dokumenti hapet pa gabime, nuk ka mbivendosje ose faqe bosh dhe përmbajtja origjinale është ruajtur. Kthe .docx final dhe përmbledhjen e kontrolleve.",
    ]:
        add_prompt_line(doc, line)

    add_heading(doc, "B", "Prompt për standardizimin e Excel-it", 2)
    add_prompt_intro(doc, "STANDARDO WORKBOOK-UN EXCEL SIPAS STANDARDIT PRIMEX / PRIMEFLOW")
    for line in [
        "Të dhënat: Përshkrimi [VENDOS]; Inicialet [VENDOS]. Hyrja mund të jetë .xlsx ose .csv.",
        "Ruaj të dhënat, formulat, komentet dhe hyperlink-et. Mos shpik header-a; pyet kur mungojnë.",
        "Hiq vetëm sheet-et plotësisht bosh. Riemërto emrat gjenerikë me tituj të sigurt/unikë dhe përditëso referencat e formulave.",
        "Vendos titullin në rreshtin 3, merged, uppercase, Calibri 16 bold centered; header-in në rreshtin 6, Calibri 11 bold, wrap.",
        "Shto/rinumëro kolonën NR, width 6. Të dhënat Calibri 11, wrap, borders; dy rreshta pune në fund. Kolonat e tjera width 10-45; row height 18-300.",
        "Formatet numerike: #,##0; #,##0.##; monetare #,##0.00. Ruaj formulat.",
        "Vendos AutoFilter në rreshtin 6, freeze C7, repeat row 6, gridlines off, A4, portrait deri 7 kolona/landscape mbi 7, fit width 1.",
        "Margjinat 0.25 majtas/djathtas, 0.5 sipër/poshtë, 0.2 header/footer. Header djathtas dd/MM/yyyy HH:mm; footer &P / &N dhe PUNOI:.",
        "Emërto [PËRSHKRIMI]_[DD.MM.YYYY]_[INICIALET].xlsx. Kthe .xlsx final dhe përmbledhjen e kontrolleve.",
    ]:
        add_prompt_line(doc, line)

    add_heading(doc, "C", "Prompt për titujt e detyrave", 2)
    add_prompt_intro(doc, "STANDARDIZO TITUJT E DETYRAVE SIPAS PRIMEFLOW")
    for line in [
        "Për çdo detyrë përdor strukturën [INICIALET]: [PROJEKTI/KLIENTI]: [VEPRIMI KRYESOR].",
        "Titulli të jetë në një rresht, me një veprim, deri në 100 karaktere të rekomanduara dhe maksimumi 120.",
        "Përdor vetëm shkurtesat zyrtare. Mos vendos llojin, statusin, prioritetin, AM/PM ose datën në titull.",
        "Hapat, URL-të dhe shpjegimet vendosi në Description. Mos shpik iniciale, klient ose projekt.",
        "Kthe: Titulli i standardizuar; Description; Informacioni që mungon.",
    ]:
        add_prompt_line(doc, line)

    add_heading(doc, "D", "Prompt për planifikimin javor", 2)
    add_prompt_intro(doc, "KRIJO PLANIFIKIMIN JAVOR TË DETYRAVE PËR PRIMEFLOW")
    for line in [
        "Hyrjet: punonjësi/inicialet, departamenti, projektet aktive, detyrat e hapura, afatet, takimet, pushimet dhe bartjet nga java e kaluar.",
        "Planifiko nga e hëna në të premte. Vendos së pari detyrat me afat të afërt dhe angazhimet fikse.",
        "Mos krijo mbivendosje dhe mos planifiko më shumë se kapaciteti real. Mos planifiko në ditë pushimi.",
        "Ndaji detyrat e mëdha në rezultate të realizueshme; identifiko varësitë, KO-në dhe bllokimet.",
        "Standardizo titujt pa vendosur llojin/statusin/prioritetin në titull. Mos shpik të dhëna.",
        "Kthe tabelë me: Dita, Ora, Titulli, Description, Projekti/Klienti, Prioriteti, Afati, Varësia, Rezultati i pritshëm.",
        "Në fund jep objektivat e javës, detyrat kritike, ato që presin përgjigje, konfliktet dhe ngarkesën për çdo ditë.",
    ]:
        add_prompt_line(doc, line)

    add_heading(doc, "E", "Prompt për kontrollin para dorëzimit", 2)
    add_prompt_intro(doc, "KONTROLLO DHE STANDARDO KËTË MATERIAL SIPAS RREGULLAVE TË PRIMEX")
    for line in [
        "Lloji [DETYRË/WORD/EXCEL/EMAIL/RAPORT/TAKIM/PROJEKT]; Projekti/Klienti [VENDOS]; Përgjegjësi [VENDOS]; Afati [VENDOS]; Rezultati [VENDOS].",
        "Kontrollo që qëllimi, rezultati, përgjegjësi dhe afati janë të qarta; shkurtesat dhe standardi përkatës janë respektuar.",
        "Kontrollo emrin/versionin, plotësinë, dublimet, gabimet, informacionin e munguar dhe të dhënat konfidenciale.",
        "Mos shpik informacione. Kthe: Problemet; Korrigjimet; Informacioni që mungon; Versioni i standardizuar; Kontrolli final.",
    ]:
        add_prompt_line(doc, line)

    doc.add_paragraph()
    add_callout(doc, "Përfundim:", "Ky manual hyn në përdorim pasi të shqyrtohet dhe miratohet nga personat përgjegjës të PrimEx. Çdo mospërputhje me kodin e PrimeFlow duhet të kthehet në detyrë konkrete për harmonizim.")
    return doc


if __name__ == "__main__":
    document = build_document()
    document.save(OUT_PATH)
    print(OUT_PATH)
