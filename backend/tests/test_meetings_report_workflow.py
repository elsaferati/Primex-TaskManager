from __future__ import annotations

import unittest
import uuid
from io import BytesIO
from zipfile import ZipFile
from datetime import date, datetime, time
from types import SimpleNamespace
from unittest.mock import AsyncMock, Mock, patch
from zoneinfo import ZoneInfo

from fastapi import HTTPException

from app.api.deps import require_admin
from app.api.routers.meetings_report import DraftUpdate, RecipientsPayload, SectionPayload, router, update_draft
from app.models.enums import UserRole
from app.models.meetings_report_draft import MeetingsReportDraft
from app.services import meetings_report_scheduler
from app.services.meeting_point_manual_sync import is_known_report_title, is_manual_section_title
from app.services.meetings_report import (
    SECTION_TITLES,
    _common_meeting_lines,
    _common_task_metadata_by_title,
    _is_new_task_for_m3_day,
    _meeting_lines,
    _meeting_status_checkbox_table,
    _m3_department_label,
    _m3_added_week_label,
    _m3_status_table,
    _m3_task_type_label,
    _leave_lines,
    _render_ascii_table_html,
    _task_owners,
    _tomorrow_task_table,
    normalize_meetings_report_sections,
    render_section_report_docx,
    render_section_report_png,
    section_report_attachments,
)


class FakeDraftDb:
    def __init__(self, draft: MeetingsReportDraft) -> None:
        self.draft = draft

    async def get(self, _model, _draft_id):
        return self.draft

    async def commit(self) -> None:
        return None

    async def refresh(self, _row) -> None:
        return None


class ScalarResult:
    def __init__(self, value) -> None:
        self.value = value

    def scalars(self):
        return self

    def first(self):
        return self.value

    def scalar_one_or_none(self):
        return self.value


class SchedulerDb:
    def __init__(self, settings, draft) -> None:
        self.results = iter((ScalarResult(settings), ScalarResult(draft)))

    async def __aenter__(self):
        return self

    async def __aexit__(self, _exc_type, _exc, _traceback) -> None:
        return None

    async def execute(self, _query):
        return next(self.results)

    async def commit(self) -> None:
        return None


def make_draft() -> MeetingsReportDraft:
    return MeetingsReportDraft(
        id=uuid.uuid4(),
        report_date=date(2026, 8, 5),
        tomorrow_date=date(2026, 8, 6),
        subject="Saved subject",
        recipients={"to": ["old@example.com"], "cc": [], "bcc": []},
        sections=[{"title": "Section", "body": "Saved user edit"}],
        generated_snapshot={},
        status="DRAFT",
    )


class MeetingsReportWorkflowTests(unittest.IsolatedAsyncioTestCase):
    def test_settings_routes_require_admin(self) -> None:
        settings_routes = [route for route in router.routes if route.path == "/settings"]
        self.assertEqual(len(settings_routes), 2)
        for route in settings_routes:
            dependencies = [dependency.call for dependency in route.dependant.dependencies]
            self.assertIn(require_admin, dependencies)

    async def test_staff_can_edit_report_content_but_not_recipients(self) -> None:
        draft = make_draft()
        db = FakeDraftDb(draft)
        staff = SimpleNamespace(id=uuid.uuid4(), role=UserRole.STAFF, full_name="Staff User")

        result = await update_draft(
            draft.id,
            DraftUpdate(
                subject="Edited subject",
                sections=[
                    SectionPayload(
                        title="Section",
                        body="+---+---+---+\n| NR | WHO | TITLE |\n+---+---+---+\n| 1 | AB | Added task |\n+---+---+---+",
                    )
                ],
            ),
            db,
            staff,
        )

        self.assertEqual(result["subject"], "Edited subject")
        self.assertEqual(result["sections"][0]["title"], "Section")
        self.assertIn("Added task", result["sections"][0]["body"])

        with self.assertRaises(HTTPException) as raised:
            await update_draft(
                draft.id,
                DraftUpdate(recipients=RecipientsPayload(to=["changed@example.com"])),
                db,
                staff,
            )
        self.assertEqual(raised.exception.status_code, 403)

    async def test_scheduler_regenerates_sections_before_sending(self) -> None:
        timezone = ZoneInfo("Europe/Tirane")
        draft = make_draft()
        settings = SimpleNamespace(
            is_active=True,
            timezone="Europe/Tirane",
            weekdays=[2],
            send_time=time(16, 30),
            recipients={"to": ["report@example.com"], "cc": [], "bcc": []},
            updated_at=datetime(2026, 8, 5, 12, 0, tzinfo=timezone),
            last_run_date=None,
        )
        db = SchedulerDb(settings, draft)
        regenerated = [{"title": "Section", "body": "Fresh auto content"}]
        build = AsyncMock(return_value=(date(2026, 8, 6), regenerated, {"counts": {}}))
        send = AsyncMock(return_value={"id": "gmail-id", "threadId": "thread-id"})
        render_plain = Mock(return_value="plain")
        render_html = Mock(return_value="html")

        with (
            patch.object(meetings_report_scheduler, "SessionLocal", return_value=db),
            patch.object(meetings_report_scheduler, "build_meetings_report_sections", build),
            patch.object(meetings_report_scheduler, "send_meetings_report", send),
            patch.object(meetings_report_scheduler, "render_plain_text", render_plain),
            patch.object(meetings_report_scheduler, "render_html", render_html),
            patch(
                "app.services.meeting_point_manual_sync.merge_common_view_manual_sections",
                AsyncMock(side_effect=lambda _db, sections, *_args: sections),
            ),
        ):
            sent = await meetings_report_scheduler.run_meetings_report_scheduler_once(
                datetime(2026, 8, 5, 17, 0, tzinfo=timezone)
            )

        self.assertTrue(sent)
        build.assert_awaited_once()
        self.assertEqual(draft.sections, regenerated)
        self.assertEqual(render_plain.call_args.args[3], regenerated)
        self.assertEqual(draft.recipients["to"], ["report@example.com"])
        self.assertEqual(draft.gmail_message_id, "gmail-id")

    async def test_scheduler_sends_fresh_m3_report_at_both_delivery_times(self) -> None:
        timezone = ZoneInfo("Europe/Tirane")
        draft = make_draft()
        settings = SimpleNamespace(
            is_active=True,
            timezone="Europe/Tirane",
            weekdays=[2],
            send_time=time(16, 30),
            recipients={"to": ["report@example.com"], "cc": [], "bcc": []},
            updated_at=datetime(2026, 8, 5, 12, 0, tzinfo=timezone),
            last_run_date=None,
        )
        build = AsyncMock(return_value=(date(2026, 8, 6), [{"title": "Section", "body": "Fresh"}], {}))
        send = AsyncMock(return_value={"id": "gmail-id", "threadId": "thread-id"})
        session_local = Mock(
            side_effect=[SchedulerDb(settings, draft), SchedulerDb(settings, draft), SchedulerDb(settings, draft)]
        )

        with (
            patch.object(meetings_report_scheduler, "SessionLocal", session_local),
            patch.object(meetings_report_scheduler, "build_meetings_report_sections", build),
            patch.object(meetings_report_scheduler, "send_meetings_report", send),
            patch.object(meetings_report_scheduler, "render_plain_text", Mock(return_value="plain")),
            patch.object(meetings_report_scheduler, "render_html", Mock(return_value="html")),
            patch(
                "app.services.meeting_point_manual_sync.merge_common_view_manual_sections",
                AsyncMock(side_effect=lambda _db, sections, *_args: sections),
            ),
        ):
            first_sent = await meetings_report_scheduler.run_meetings_report_scheduler_once(
                datetime(2026, 8, 5, 15, 50, tzinfo=timezone)
            )
            second_sent = await meetings_report_scheduler.run_meetings_report_scheduler_once(
                datetime(2026, 8, 5, 16, 30, tzinfo=timezone)
            )
            no_third_send = await meetings_report_scheduler.run_meetings_report_scheduler_once(
                datetime(2026, 8, 5, 16, 31, tzinfo=timezone)
            )

        self.assertTrue(first_sent)
        self.assertTrue(second_sent)
        self.assertFalse(no_third_send)
        self.assertEqual(draft.auto_sent_slots, ["15:50", "16:30"])
        self.assertEqual(build.await_count, 2)
        self.assertEqual(send.await_count, 2)


class MeetingsReportAliasDedupTests(unittest.TestCase):
    def test_section_exports_keep_email_groups_tables_and_status_colors(self) -> None:
        sections = [
            {"title": "A JEMI BRENDA MESATARES ME PROJEKTE?", "body": "(Ploteso manualisht)"},
            {
                "title": "GA MBYLLJA E DET",
                "body": (
                    "TODO:\n+----+------+--------------------------------+\n"
                    "| NR | KUSH | TITULLI                        |\n"
                    "+----+------+--------------------------------+\n"
                    "| 1  | GA   | DetyrÃ« e re                    |\n"
                    "+----+------+--------------------------------+\n"
                    "DONE:\n+----+------+--------------------------------+\n"
                    "| NR | KUSH | TITULLI                        |\n"
                    "+----+------+--------------------------------+\n"
                    "| 1  | GA   | DetyrÃ« e kryer                 |\n"
                    "+----+------+--------------------------------+"
                ),
            },
        ]
        docx_bytes = render_section_report_docx("PrimeFlow M3", "M3", date(2026, 8, 10), sections)
        png_bytes = render_section_report_png("PrimeFlow M3", "M3", date(2026, 8, 10), sections)

        from docx import Document
        from PIL import Image

        document = Document(BytesIO(docx_bytes))
        document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        self.assertIn("PrimeFlow M3", document_text)
        self.assertIn("MANUAL QUESTIONS", table_text)
        self.assertIn("AUTO-FILLED FROM PRIMEFLOW", table_text)
        self.assertIn("TODO:", table_text)
        self.assertIn("DONE:", table_text)
        with ZipFile(BytesIO(docx_bytes)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn("fbcfe8", document_xml.lower())
        self.assertIn("d4ffe1", document_xml.lower())
        self.assertIn('w:w="432"', document_xml)  # Compact NR column, matching the email grid.
        image = Image.open(BytesIO(png_bytes))
        self.assertGreater(image.width, 1000)
        self.assertGreater(image.height, 400)

    def test_section_report_attachments_include_word_and_png(self) -> None:
        attachments = section_report_attachments(
            "PrimeFlow M3",
            "M3",
            date(2026, 8, 10),
            [{
                "title": "Section",
                "body": "TODO:\n+----+-----+\n| NR | TITLE |\n+----+-----+\n| 1  | Task  |\n+----+-----+",
            }],
            tomorrow=date(2026, 8, 11),
        )

        self.assertEqual([attachment[0] for attachment in attachments], [
            "PrimeFlow-M3-2026-08-10.docx",
            "PrimeFlow-M3-2026-08-10.png",
        ])
        self.assertTrue(attachments[0][1].startswith(b"PK"))
        self.assertTrue(attachments[1][1].startswith(b"\x89PNG"))

    def test_normalize_collapses_common_view_aliases_into_auto_sections(self) -> None:
        auto_tickets = "STD TICKETS: 2"
        normalized = normalize_meetings_report_sections(
            [
                {"title": "A JEMI BRENDA MESATARES ME PROJEKTE/", "body": "(Ploteso manualisht)"},
                {"title": "(GA) M3 DET GA MBYLLJA ME HV/OH?", "body": "(Ploteso manualisht)"},
                {"title": "(GA) TIKETAT E STD DHE TONAT? RAPORTOHEN NE M3", "body": "(Ploteso manualisht)"},
                {"title": SECTION_TITLES[1], "body": auto_tickets},
                {"title": "Brand new Common View pike", "body": "(Ploteso manualisht)"},
            ]
        )
        titles = [section["title"] for section in normalized]
        self.assertEqual(titles[0], SECTION_TITLES[0])  # manual first
        auto_titles = [title for title in titles if title == SECTION_TITLES[1]]
        self.assertEqual(auto_titles[0], SECTION_TITLES[1])  # tickets first among auto-filled
        self.assertEqual(titles.count(SECTION_TITLES[1]), 1)
        self.assertNotIn("GA MBYLLJA E DET", titles)
        self.assertNotIn("HV MBYLLJA E DET", titles)
        self.assertNotIn("(GA) M3 DET GA MBYLLJA ME HV/OH?", titles)
        self.assertNotIn("(GA) TIKETAT E STD DHE TONAT? RAPORTOHEN NE M3", titles)
        self.assertIn("Brand new Common View pike", titles)
        by_title = {section["title"]: section["body"] for section in normalized}
        self.assertEqual(by_title[SECTION_TITLES[1]], auto_tickets)

    def test_ticket_wording_variant_alone_maps_to_auto_title(self) -> None:
        normalized = normalize_meetings_report_sections(
            [
                {"title": "A JEMI BRENDA MESATARES ME PROJEKTE/", "body": "(Ploteso manualisht)"},
                {"title": "(GA) TIKETAT E STD DHE TONAT? RAPORTOHEN NE M3", "body": "(Ploteso manualisht)"},
            ]
        )
        titles = [section["title"] for section in normalized]
        self.assertNotIn("(GA) TIKETAT E STD DHE TONAT? RAPORTOHEN NE M3", titles)
        self.assertIn(SECTION_TITLES[1], titles)

    def test_ga_and_hv_closing_sections_are_removed_from_m3(self) -> None:
        normalized = normalize_meetings_report_sections(
            [
                {"title": SECTION_TITLES[0], "body": "(Ploteso manualisht)"},
                {"title": "GA MBYLLJA E DET", "body": "GA tasks"},
                {"title": "HV MBYLLJA E DET", "body": "HV tasks"},
            ]
        )

        self.assertEqual([section["title"] for section in normalized], [SECTION_TITLES[0]])

    def test_common_view_aliases_are_known_auto_not_manual(self) -> None:
        self.assertTrue(is_known_report_title("meetings", "(GA) M3 DET GA MBYLLJA ME HV/OH?"))
        self.assertTrue(is_known_report_title("meetings", "(GA) TIKETAT E STD DHE TONAT? RAPORTOHEN NE M3"))
        self.assertFalse(is_manual_section_title("meetings", "(GA) M3 DET GA MBYLLJA ME HV/OH?"))
        self.assertFalse(is_manual_section_title("meetings", "(GA) TIKETAT E STD DHE TONAT? RAPORTOHEN NE M3"))
        self.assertTrue(is_manual_section_title("meetings", "Brand new Common View pike"))


class MeetingsReportTaskTypeColumnTests(unittest.TestCase):
    def test_meeting_status_table_sorts_by_clock_time_not_source_date(self) -> None:
        meetings = [
            SimpleNamespace(
                id=uuid.uuid4(), title="13:15 meeting", starts_at=datetime(2026, 8, 1, 13, 15), recurrence_type="daily"
            ),
            SimpleNamespace(
                id=uuid.uuid4(), title="10:15 meeting", starts_at=datetime(2026, 8, 20, 10, 15), recurrence_type="once"
            ),
            SimpleNamespace(
                id=uuid.uuid4(), title="08:30 meeting", starts_at=datetime(2026, 8, 1, 8, 30), recurrence_type="daily"
            ),
        ]

        rows = _meeting_status_checkbox_table(meetings, {})
        meeting_rows = [row for row in rows if "meeting" in row]

        self.assertIn("08:30", meeting_rows[0])
        self.assertIn("10:15", meeting_rows[1])
        self.assertIn("13:15", meeting_rows[2])

    def test_pv_displays_date_range_or_one_day_partial_time_range(self) -> None:
        user_id = uuid.uuid4()
        entry = SimpleNamespace(assigned_to_user_id=user_id, created_by_user_id=None)

        rows = _leave_lines(
            [
                (entry, date(2026, 8, 10), date(2026, 8, 12), True, None, None, None, False),
                (entry, date(2026, 8, 13), date(2026, 8, 13), False, "08:00", "10:00", None, False),
            ],
            {user_id: "Finance Group"},
            {user_id: "FIN"},
        )

        header = next(row for row in rows if "KUSH" in row and "DEP" in row and "NGA" in row and "DERI" in row)
        self.assertLess(header.index("KUSH"), header.index("DEP"))
        self.assertLess(header.index("NGA"), header.index("DERI"))
        self.assertTrue(any("FG" in row and "FIN" in row for row in rows))
        self.assertTrue(any("10.08.2026" in row and "12.08.2026" in row for row in rows))
        self.assertTrue(any("08:00" in row and "10:00" in row for row in rows))

    def test_pv_uses_the_weekly_planner_user_order(self) -> None:
        first_user_id = uuid.uuid4()
        second_user_id = uuid.uuid4()
        first = SimpleNamespace(id=uuid.uuid4(), assigned_to_user_id=first_user_id, created_by_user_id=None)
        second = SimpleNamespace(id=uuid.uuid4(), assigned_to_user_id=second_user_id, created_by_user_id=None)

        rows = _leave_lines(
            [
                (second, date(2026, 8, 12), date(2026, 8, 12), True, None, None, None, False),
                (first, date(2026, 8, 12), date(2026, 8, 12), True, None, None, None, False),
            ],
            {first_user_id: "Zeta User", second_user_id: "Alpha User"},
            user_sort_keys={
                first_user_id: (0, "dev", 0, 1, "zeta user"),
                second_user_id: (0, "dev", 0, 2, "alpha user"),
            },
        )

        data_rows = [row for row in rows if "ZU" in row or "AU" in row]
        self.assertIn("ZU", data_rows[0])
        self.assertIn("AU", data_rows[1])

    def test_meeting_status_cells_use_green_and_red_backgrounds(self) -> None:
        html = _render_ascii_table_html(
            [
                "| NR | KOHA | MBAJTUR? | TITULLI |",
                "| 1 | 10:00 | ✓ | Held meeting |",
                "| 2 | 11:00 | ✕ | Canceled meeting |",
            ]
        )

        self.assertIn('class="n held"', html)
        self.assertIn('class="n canceled"', html)

    def test_non_daily_or_weekly_meetings_receive_a_blue_frame_and_title(self) -> None:
        html = _render_ascii_table_html(
            [
                "| NR | KOHA | TITULLI |",
                "| 1 | 10:00 | One-off meeting [[mt:non_daily_weekly]] |",
            ]
        )

        self.assertIn('class="highlight"', html)
        self.assertIn('class="title"', html)
        self.assertNotIn("[[mt:non_daily_weekly]]", html)

    def test_common_view_one_time_meetings_receive_the_highlight_marker(self) -> None:
        lines = _common_meeting_lines(
            [
                {"id": "one-time", "title": "One-time meeting", "date": "2026-08-11", "time": "10:00", "recurrence_type": "none"},
                {"id": "weekly", "title": "Weekly meeting", "date": "2026-08-11", "time": "11:00", "recurrence_type": "weekly"},
            ],
            date(2026, 8, 11),
        )

        self.assertIn("[[mt:non_daily_weekly]]", lines[0])
        self.assertNotIn("[[mt:non_daily_weekly]]", lines[1])

    def test_common_view_meetings_are_sorted_by_time(self) -> None:
        lines = _common_meeting_lines(
            [
                {"id": "late", "title": "Later meeting", "date": "2026-08-11", "time": "13:15"},
                {"id": "early", "title": "Earlier meeting", "date": "2026-08-11", "time": "10:15"},
                {"id": "first", "title": "First meeting", "date": "2026-08-11", "time": "08:30"},
            ],
            date(2026, 8, 11),
        )

        self.assertEqual(
            lines,
            [
                "- 08:30: First meeting [[mt:non_daily_weekly]]",
                "- 10:15: Earlier meeting [[mt:non_daily_weekly]]",
                "- 13:15: Later meeting [[mt:non_daily_weekly]]",
            ],
        )

    def test_database_meetings_are_sorted_by_clock_time_not_original_date(self) -> None:
        utc = ZoneInfo("UTC")
        lines = _meeting_lines(
            [
                SimpleNamespace(title="Recurring later", starts_at=datetime(2026, 8, 10, 11, 15, tzinfo=utc), recurrence_type="weekly"),
                SimpleNamespace(title="One-off earlier", starts_at=datetime(2026, 8, 11, 8, 15, tzinfo=utc), recurrence_type="none"),
            ]
        )

        self.assertTrue(lines[0].startswith("- 10:15: One-off earlier"))
        self.assertTrue(lines[1].startswith("- 13:15: Recurring later"))

    def test_many_assignees_display_as_all(self) -> None:
        assignee_ids = {uuid.uuid4() for _ in range(11)}
        task = SimpleNamespace(id=uuid.uuid4(), assigned_to=None)
        names = {user_id: f"User {index}" for index, user_id in enumerate(assignee_ids, start=1)}

        self.assertEqual(_task_owners(task, names, {task.id: assignee_ids}), "ALL")

    def test_task_type_labels(self) -> None:
        self.assertEqual(
            _m3_task_type_label(SimpleNamespace(system_template_origin_id=uuid.uuid4(), project_id=uuid.uuid4())),
            "SYS",
        )
        self.assertEqual(
            _m3_task_type_label(
                SimpleNamespace(
                    project_id=uuid.uuid4(),
                    ga_note_origin_id=uuid.uuid4(),
                    is_1h_report=True,
                )
            ),
            "PRJK",
        )
        self.assertEqual(
            _m3_task_type_label(SimpleNamespace(ga_note_origin_id=uuid.uuid4(), is_bllok=True)),
            "BLL",
        )
        self.assertEqual(_m3_task_type_label(SimpleNamespace(is_1h_report=True)), "1H")
        self.assertEqual(_m3_task_type_label(SimpleNamespace(is_r1=True)), "R1")
        self.assertEqual(_m3_task_type_label(SimpleNamespace(is_personal=True)), "P")
        self.assertEqual(
            _m3_task_type_label(SimpleNamespace(plan_note_origin_id=uuid.uuid4())),
            "FT",
        )
        self.assertEqual(_m3_task_type_label(SimpleNamespace()), "FT")

    def test_status_table_includes_type_column(self) -> None:
        task = SimpleNamespace(
            id=uuid.uuid4(),
            title="KA: MODECO KONT",
            status="TODO",
            is_bllok=True,
            is_r1=False,
            is_1h_report=False,
            is_personal=False,
            is_deadline_important=False,
            ga_note_origin_id=None,
            plan_note_origin_id=None,
            system_template_origin_id=None,
            project_id=None,
            assigned_to=None,
            fast_task_order=None,
            due_date=None,
            start_date=None,
            completed_at=None,
            created_at=None,
        )
        rows = _m3_status_table("TODO", [task], {}, include_type=True)
        header = next(row for row in rows if "LLOJI" in row and "TITULLI" in row)
        self.assertIn("LLOJI", header)
        self.assertTrue(any("BLL" in row and "MODECO" in row for row in rows))

    def test_new_task_table_keeps_each_task_status_for_row_coloring(self) -> None:
        task_template = {
            "id": uuid.uuid4(),
            "system_template_origin_id": None,
            "project_id": None,
            "assigned_to": None,
            "fast_task_order": None,
            "is_deadline_important": False,
            "due_date": None,
            "start_date": None,
            "completed_at": None,
            "created_at": None,
        }
        tasks = [
            SimpleNamespace(**task_template, title="Todo task", status="TODO"),
            SimpleNamespace(**{**task_template, "id": uuid.uuid4()}, title="Progress task", status="IN_PROGRESS"),
            SimpleNamespace(**{**task_template, "id": uuid.uuid4()}, title="Waiting task", status="WAITING_CONFIRMATION"),
            SimpleNamespace(**{**task_template, "id": uuid.uuid4()}, title="Done task", status="DONE"),
        ]

        rows = _m3_status_table("DETYRAT E REJA", tasks, {}, with_status=True)
        html = _render_ascii_table_html(rows)

        self.assertIn('class="todo"', html)
        self.assertIn('class="in-progress"', html)
        self.assertIn('class="waiting"', html)
        self.assertIn('class="done"', html)

    def test_status_table_orders_system_tasks_before_fast_tasks(self) -> None:
        system_task = SimpleNamespace(
            id=uuid.uuid4(), title="System task", status="TODO", system_template_origin_id=uuid.uuid4(),
            project_id=None, assigned_to=None, fast_task_order=None, is_deadline_important=False,
            due_date=None, start_date=None, completed_at=None, created_at=None,
        )
        fast_task = SimpleNamespace(
            id=uuid.uuid4(), title="Fast task", status="TODO", system_template_origin_id=None,
            project_id=None, assigned_to=None, fast_task_order=None, is_deadline_important=False,
            due_date=None, start_date=None, completed_at=None, created_at=None,
        )

        rows = _m3_status_table("TODO", [fast_task, system_task], {}, include_type=True)
        task_rows = [row for row in rows if "System task" in row or "Fast task" in row]

        self.assertIn("System task", task_rows[0])
        self.assertIn("Fast task", task_rows[1])

    def test_late_system_task_table_includes_department_after_who(self) -> None:
        department_id = uuid.uuid4()
        task = SimpleNamespace(
            id=uuid.uuid4(),
            title="Late system task",
            status="TODO",
            finish_period="PM",
            system_template_origin_id=uuid.uuid4(),
            department_id=department_id,
            assigned_to=None,
            fast_task_order=None,
            is_deadline_important=False,
            due_date=None,
            start_date=None,
            completed_at=None,
            created_at=None,
        )
        rows = _m3_status_table(
            "LATE",
            [task],
            {},
            include_late_days=True,
            include_department=True,
            include_am_pm=True,
            department_codes={department_id: "FIN"},
        )
        header = next(row for row in rows if "KUSH" in row and "DEP" in row and "AM/PM" in row and "TITULLI" in row)
        self.assertLess(header.index("KUSH"), header.index("DEP"))
        self.assertLess(header.index("DEP"), header.index("AM/PM"))
        self.assertLess(header.index("DEP"), header.index("TITULLI"))
        self.assertTrue(any("FIN" in row and "PM" in row and "Late system task" in row for row in rows))
        self.assertEqual(_m3_department_label(task, {department_id: "GDS"}), "GD")

    def test_bz_table_includes_department_and_am_pm_after_who(self) -> None:
        rows = _tomorrow_task_table(
            "BZ ME GA",
            ["- [TODO] FG: BZ task"],
            with_status=True,
            task_metadata={"BZ task": ("FIN", "AM")},
            include_am_pm_times=True,
        )

        header = next(row for row in rows if "KUSH" in row and "DEP" in row and "AM/PM" in row)
        self.assertLess(header.index("KUSH"), header.index("DEP"))
        self.assertLess(header.index("DEP"), header.index("AM/PM"))
        self.assertTrue(any("FIN" in row and "AM (08:15)" in row and "BZ task" in row for row in rows))

    def test_common_view_task_metadata_uses_string_department_id(self) -> None:
        department_id = uuid.uuid4()
        metadata = _common_task_metadata_by_title(
            [
                {
                    "title": "Blocked task",
                    "date": "2026-08-12",
                    "department_id": str(department_id),
                    "finish_period": "PM",
                }
            ],
            date(2026, 8, 12),
            {department_id: "PRODUCT CONTENT"},
        )

        self.assertEqual(metadata, {"Blocked task": ("PCM", "PM")})

    def test_bz_table_orders_tasks_by_department(self) -> None:
        rows = _tomorrow_task_table(
            "BZ ME GA",
            ["- [TODO] A: PCM task", "- [TODO] B: GD task", "- [TODO] C: DEV task"],
            with_status=True,
            task_metadata={
                "PCM task": ("PCM", "AM"),
                "GD task": ("GD", "AM"),
                "DEV task": ("DEV", "AM"),
            },
        )

        task_rows = [row for row in rows if "task" in row]
        self.assertIn("DEV task", task_rows[0])
        self.assertIn("GD task", task_rows[1])
        self.assertIn("PCM task", task_rows[2])

    def test_task_tables_honor_weekly_planner_department_and_user_order(self) -> None:
        def task(title: str, weekly_order: tuple) -> SimpleNamespace:
            return SimpleNamespace(
                id=uuid.uuid4(), title=title, status="TODO", assigned_to=None,
                fast_task_order=None, is_deadline_important=False, due_date=None,
                start_date=None, completed_at=None, created_at=None,
                _weekly_planner_report_sort=weekly_order,
            )

        rows = _m3_status_table(
            "TODO",
            [
                task("PCM task", (2, "pcm", 0, 0, "DM")),
                task("DEV EF task", (0, "dev", 0, 2, "EF")),
                task("GD task", (1, "gd", 0, 0, "DM")),
                task("DEV AT task", (0, "dev", 0, 1, "AT")),
            ],
            {},
        )
        task_rows = [row for row in rows if "task" in row]

        self.assertIn("DEV AT task", task_rows[0])
        self.assertIn("DEV EF task", task_rows[1])
        self.assertIn("GD task", task_rows[2])
        self.assertIn("PCM task", task_rows[3])

    def test_new_tomorrow_table_marks_this_and_last_week_after_department(self) -> None:
        department_id = uuid.uuid4()
        this_week = SimpleNamespace(
            id=uuid.uuid4(), title="This week", status="TODO", department_id=department_id,
            finish_period="AM", assigned_to=None, fast_task_order=None, is_deadline_important=False,
            due_date=None, start_date=None, completed_at=None, created_at=date(2026, 8, 10),
        )
        last_week = SimpleNamespace(
            id=uuid.uuid4(), title="Last week", status="TODO", department_id=department_id,
            finish_period="PM", assigned_to=None, fast_task_order=None, is_deadline_important=False,
            due_date=None, start_date=None, completed_at=None, created_at=date(2026, 8, 7),
        )
        rows = _m3_status_table(
            "DET TE REJA LAST WEEK DHE THIS WEEK", [this_week, last_week], {},
            include_department=True, include_added_week=True, include_am_pm=True,
            department_codes={department_id: "DEV"}, week_start=date(2026, 8, 10),
        )

        header = next(row for row in rows if "DEP" in row and "KRIJUAR" in row and "AM/PM" in row)
        self.assertLess(header.index("DEP"), header.index("KRIJUAR"))
        self.assertLess(header.index("KRIJUAR"), header.index("AM/PM"))
        self.assertTrue(any("This W" in row for row in rows))
        self.assertTrue(any("Last W" in row for row in rows))
        self.assertEqual(_m3_added_week_label(this_week, date(2026, 8, 10)), "This W")

    def test_new_m3_tasks_are_included_only_on_their_start_date(self) -> None:
        wednesday = date(2026, 8, 12)
        task = SimpleNamespace(
            start_date=wednesday,
            due_date=date(2026, 8, 14),
            created_at=date(2026, 8, 7),
            completed_at=None,
            status="TODO",
        )

        self.assertTrue(_is_new_task_for_m3_day(task, wednesday))
        self.assertFalse(_is_new_task_for_m3_day(task, date(2026, 8, 13)))

        no_start_date = SimpleNamespace(
            start_date=None,
            due_date=wednesday,
            created_at=date(2026, 8, 10),
            completed_at=None,
            status="TODO",
        )
        self.assertFalse(_is_new_task_for_m3_day(no_start_date, wednesday))

    def test_todo_table_includes_department_after_who(self) -> None:
        department_id = uuid.uuid4()
        task = SimpleNamespace(
            id=uuid.uuid4(),
            title="Pink task",
            status="TODO",
            system_template_origin_id=None,
            department_id=department_id,
            assigned_to=None,
            fast_task_order=None,
            is_deadline_important=False,
            due_date=None,
            start_date=None,
            completed_at=None,
            created_at=None,
        )
        rows = _m3_status_table(
            "TODO",
            [task],
            {},
            include_department=True,
            department_codes={department_id: "PCM"},
        )
        header = next(row for row in rows if "KUSH" in row and "DEP" in row and "TITULLI" in row)
        self.assertLess(header.index("KUSH"), header.index("DEP"))
        self.assertTrue(any("PCM" in row and "Pink task" in row for row in rows))


if __name__ == "__main__":
    unittest.main()
