from __future__ import annotations

import unittest
from datetime import date, datetime, timezone
from io import BytesIO
from types import SimpleNamespace
from unittest.mock import AsyncMock, patch

from docx import Document
from openpyxl import load_workbook
from PIL import Image

from app.services.tomorrow_closing_sections import (
    ClosingSection,
    ClosingTable,
    ClosingTableRow,
    _stack_start_due,
    _include_unfinished_system_task,
    _system_task_tyo_label,
)
from app.services.tomorrow_print_report import (
    _closing_sections_html,
    _docx_table_attachment,
    _excel_table_attachment,
    _png_table_attachment,
    _is_overdue_tyo_value,
    build_tomorrow_print_report,
)


def closing_fixture() -> list[ClosingSection]:
    return [
        ClosingSection(
            title="DET PA PROGRESS",
            tables=[ClosingTable(
                label="TODO",
                columns=["NR", "KUSH", "DEP", "AM/PM", "TITULLI", "ARSYEJA", "KOMENT"],
                rows=[ClosingTableRow(
                    values=["1", "EF", "DEV", "PM", "Pink task", "Urgjence", "Pres input"],
                    status="TODO",
                )],
                tone="todo",
            )],
        ),
        ClosingSection(
            title="DET SYS PA KRY",
            tables=[ClosingTable(
                label="DET SYS PA KRY",
                columns=[
                    "NR", "KUSH", "DEP", "AM/PM", "T/Y/O", "TITULLI",
                    "ARSYEJA", "KOMENT",
                ],
                rows=[ClosingTableRow(
                    values=[
                        "1", "RA", "DEV", "AM", "2", "System unfinished",
                        "Pa progres", "Pres sqarim",
                    ],
                    status="IN_PROGRESS",
                )],
            )],
        ),
        ClosingSection(
            title="DET E SHTYERA",
            tables=[ClosingTable(
                label="SHTYER DUE DATE",
                columns=["NR", "KUSH", "DEP", "AM/PM", "LLOJI", "NGA", "NE", "TITULLI"],
                rows=[ClosingTableRow(
                    values=[
                        "1", "FG", "GD", "PM", "P",
                        "START: 02.09.2026\nDUE: 02.09.2026",
                        "START: 03.09.2026\nDUE: 03.09.2026",
                        "Moved task",
                    ],
                    status="WAITING_CONFIRMATION",
                )],
            )],
        ),
        ClosingSection(
            title="NOTES PA DISK",
            tables=[ClosingTable(
                label="NOTES",
                columns=["NR", "DISK", "NOTE", "FROM", "TIME"],
                rows=[ClosingTableRow(values=["1", "NO", "Undiscussed note", "GA", "09:15"])],
                tone="notes",
            )],
        ),
    ]


class TomorrowClosingFormatParityTests(unittest.TestCase):
    def test_only_overdue_tyo_values_use_the_alert_style(self) -> None:
        self.assertFalse(_is_overdue_tyo_value("T"))
        self.assertFalse(_is_overdue_tyo_value("-"))
        self.assertTrue(_is_overdue_tyo_value("Y"))
        self.assertTrue(_is_overdue_tyo_value("2"))
        self.assertTrue(_is_overdue_tyo_value("12"))

    def test_m3_start_due_value_is_stacked(self) -> None:
        self.assertEqual(
            _stack_start_due("START: 02.09.2026 / DUE: 02.09.2026"),
            "START: 02.09.2026\nDUE: 02.09.2026",
        )

    def test_system_task_tyo_uses_daily_report_business_day_logic(self) -> None:
        task = SimpleNamespace(
            project_id=None,
            start_date=None,
            due_date=datetime(2026, 9, 4, 8, 0, tzinfo=timezone.utc),
        )

        # Friday to Monday is one business day late; the weekend is excluded.
        self.assertEqual(_system_task_tyo_label(task, date(2026, 9, 7)), "Y")
        self.assertEqual(_system_task_tyo_label(task, date(2026, 9, 8)), "2")

    def test_system_task_tyo_marks_due_today_and_missing_due_date(self) -> None:
        due_today = SimpleNamespace(
            project_id=None,
            start_date=None,
            due_date=datetime(2026, 9, 7, 8, 0, tzinfo=timezone.utc),
        )
        no_due_date = SimpleNamespace(project_id=None, start_date=None, due_date=None)

        self.assertEqual(_system_task_tyo_label(due_today, date(2026, 9, 7)), "T")
        self.assertEqual(_system_task_tyo_label(no_due_date, date(2026, 9, 7)), "-")

    def test_unfinished_overdue_system_task_is_included_without_daily_baseline(self) -> None:
        task = SimpleNamespace(
            id="overdue-system-task",
            system_template_origin_id="template-1",
            system_task_slot_id=None,
            status="TODO",
            completed_at=None,
            project_id=None,
            start_date=datetime(2026, 9, 4, 8, 0, tzinfo=timezone.utc),
            due_date=datetime(2026, 9, 4, 8, 0, tzinfo=timezone.utc),
            created_at=datetime(2026, 9, 4, 8, 0, tzinfo=timezone.utc),
        )

        self.assertTrue(_include_unfinished_system_task(task, date(2026, 9, 7), set()))
        self.assertEqual(_system_task_tyo_label(task, date(2026, 9, 7)), "Y")

    def test_old_undated_system_task_is_not_added_unless_it_is_in_baseline(self) -> None:
        task = SimpleNamespace(
            id="undated-system-task",
            system_template_origin_id="template-1",
            system_task_slot_id=None,
            status="TODO",
            completed_at=None,
            project_id=None,
            start_date=datetime(2026, 9, 4, 8, 0, tzinfo=timezone.utc),
            due_date=None,
            created_at=datetime(2026, 9, 4, 8, 0, tzinfo=timezone.utc),
        )

        self.assertFalse(_include_unfinished_system_task(task, date(2026, 9, 7), set()))
        self.assertTrue(
            _include_unfinished_system_task(task, date(2026, 9, 7), {"undated-system-task"})
        )

    def test_html_xlsx_png_and_docx_contain_same_closing_sections(self) -> None:
        sections = closing_fixture()
        expected = [section.title for section in sections] + [
            "Pink task", "System unfinished", "Pa progres", "Pres sqarim",
            "Moved task", "Undiscussed note",
        ]

        html = _closing_sections_html(sections)
        for value in expected:
            self.assertIn(value, html)
        self.assertLess(html.index("DET PA PROGRESS"), html.index("DET SYS PA KRY"))
        self.assertIn("background-color:#FFC4ED", html)
        self.assertIn("background-color:#FFFF00", html)
        self.assertIn("background-color:#DBEAFE", html)
        self.assertIn("table-layout:auto", html)
        self.assertIn('width="1%"', html)
        self.assertIn("white-space:nowrap;width:1%", html)
        self.assertIn('width="16%"', html)
        self.assertIn("border-bottom:3px solid #334155", html)
        self.assertIn("T/Y/O", html)
        self.assertIn("background-color:#DC2626;color:#FFFFFF", html)
        self.assertIn("font-weight:800;text-align:left", html)
        self.assertIn("START: 02.09.2026</div>", html)
        self.assertIn("DUE: 02.09.2026</div>", html)

        _, xlsx_bytes, _ = _excel_table_attachment([], [], date(2026, 9, 3), closing_sections=sections)
        workbook = load_workbook(BytesIO(xlsx_bytes))
        xlsx_text = "\n".join(str(cell.value or "") for row in workbook.active.iter_rows() for cell in row)
        for value in expected:
            self.assertIn(value, xlsx_text)
        overdue_cell = next(cell for row in workbook.active.iter_rows() for cell in row if cell.value == "2")
        self.assertEqual(overdue_cell.fill.fgColor.rgb, "00DC2626")
        self.assertEqual(overdue_cell.alignment.horizontal, "left")

        _, png_bytes, _ = _png_table_attachment([], date(2026, 9, 3), closing_sections=sections)
        image = Image.open(BytesIO(png_bytes))
        self.assertEqual(image.format, "PNG")
        self.assertGreater(image.height, 500)

        _, docx_bytes, _ = _docx_table_attachment([], date(2026, 9, 3), closing_sections=sections)
        document = Document(BytesIO(docx_bytes))
        self.assertGreater(document.sections[0].page_width, document.sections[0].page_height)
        self.assertGreaterEqual(len(document.tables), 5)
        docx_text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        )
        for value in expected:
            self.assertIn(value, docx_text)
        document_xml = document._element.xml
        for color in ("FFC4ED", "FFFF00", "FFEDD5", "DBEAFE", "FEE2E2"):
            self.assertIn(color, document_xml)
        self.assertIn('w:color="334155"', document_xml)
        self.assertIn('w:sz="12"', document_xml)


class TomorrowClosingBuildTests(unittest.IsolatedAsyncioTestCase):
    async def test_tomorrow_report_places_closing_sections_before_tasks_and_adds_all_formats(self) -> None:
        payload = {
            "items": {"oneH": [{
                "id": "task-next", "title": "Tomorrow task", "date": "2026-09-03",
                "status": "TODO", "oneHReportSlot": "10:00",
            }]},
            "users": [],
            "departments": [],
        }
        with patch(
            "app.services.tomorrow_print_report.build_tomorrow_closing_sections",
            new=AsyncMock(return_value=closing_fixture()),
        ):
            report = await build_tomorrow_print_report(
                date(2026, 9, 2), include_attachment=True, db=object(), payload=payload
            )
        self.assertLess(report["html"].index("DET PA PROGRESS"), report["html"].index("Tomorrow task"))
        self.assertEqual(
            [attachment[2] for attachment in report["attachments"]],
            [
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "image/png",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ],
        )


if __name__ == "__main__":
    unittest.main()
