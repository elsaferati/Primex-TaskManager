from __future__ import annotations

import io
import unittest
import zipfile
from datetime import date, datetime
from zoneinfo import ZoneInfo

from docx import Document
from openpyxl import load_workbook

from app.celery_app import celery_app
from app.config import settings
from app.services.px_jav_weekly_report import (
    EXCEL_MIME,
    PDF_MIME,
    WORD_MIME,
    PxJavNoteRow,
    PxJavWeeklyReport,
    build_attachments,
    classify_note_result,
    previous_thursday_cutoff,
    render_docx,
    render_pdf,
    render_plain_text,
    render_xlsx,
)


class PxJavWeeklyReportTests(unittest.TestCase):
    def setUp(self) -> None:
        generated_at = datetime(2026, 8, 20, 15, 50, tzinfo=ZoneInfo("Europe/Tirane"))
        common = {
            "note_status": "OPEN",
            "priority": "MEDIUM",
            "discussed": True,
            "created_at": generated_at,
            "created_by": "Test User",
            "department": "IT",
            "project": "PrimeFlow",
        }
        self.report = PxJavWeeklyReport(
            report_date=date(2026, 8, 20),
            generated_at=generated_at,
            period_start=datetime(2026, 8, 13, 15, 50, tzinfo=ZoneInfo("Europe/Tirane")),
            period_end=generated_at,
            timezone="Europe/Tirane",
            recipient="334primex.eu@gmail.com",
            source_note_count=7,
            period_note_count=6,
            year_end_comment_count=1,
            excluded_task_count=1,
            excluded_next_week_count=1,
            rows=[
                PxJavNoteRow(
                    note_id="note-only",
                    number=1,
                    content="Shënim i ri pa task",
                    next_week=False,
                    result="VETËM SHËNIM",
                    **common,
                ),
                PxJavNoteRow(
                    note_id="note-mismatch",
                    number=2,
                    content="Flag pa task real",
                    next_week=False,
                    result="MOSPËRPUTHJE",
                    **common,
                ),
                PxJavNoteRow(
                    note_id="note-next-week-task",
                    number=3,
                    content="Task i krijuar për javën tjetër",
                    next_week=True,
                    result="DETYRË",
                    task_count=1,
                    active_task_count=1,
                    assignees=["Test User"],
                    task_statuses=["OPEN"],
                    task_due_dates=[date(2026, 12, 31)],
                    year_end_task=True,
                    year_end_task_count=1,
                    **common,
                ),
                PxJavNoteRow(
                    note_id="note-year-end-comment",
                    number=4,
                    content="Shënim i vjetër i fundvitit",
                    comment="31.12.2026",
                    next_week=False,
                    result="VETËM SHËNIM",
                    year_end_comment=True,
                    **{
                        **common,
                        "created_at": datetime(2026, 6, 4, 10, 37, tzinfo=ZoneInfo("Europe/Tirane")),
                    },
                ),
            ],
        )

    def test_classification_uses_actual_linked_tasks(self) -> None:
        self.assertEqual(classify_note_result(1, False), "DETYRË")
        self.assertEqual(classify_note_result(1, True), "DETYRË")
        self.assertEqual(classify_note_result(0, False), "VETËM SHËNIM")
        self.assertEqual(classify_note_result(0, True), "MOSPËRPUTHJE")

    def test_previous_thursday_cutoff_is_report_to_report(self) -> None:
        zone = ZoneInfo("Europe/Tirane")
        self.assertEqual(
            previous_thursday_cutoff(datetime(2026, 8, 17, 9, 0, tzinfo=zone)),
            datetime(2026, 8, 13, 15, 50, tzinfo=zone),
        )
        self.assertEqual(
            previous_thursday_cutoff(datetime(2026, 8, 20, 15, 50, tzinfo=zone)),
            datetime(2026, 8, 13, 15, 50, tzinfo=zone),
        )

    def test_summary_includes_missing_tasks_and_created_next_week_tasks(self) -> None:
        self.assertEqual(self.report.summary(), {
            "period_notes": 6,
            "year_end_comments": 1,
            "report_notes": 4,
            "notes_without_task": 3,
            "next_week_tasks": 1,
            "note_only": 2,
            "inconsistencies": 1,
            "excluded_with_task": 1,
            "excluded_next_week": 1,
        })
        text = render_plain_text(self.report)
        self.assertIn("Pa task (në raport): 3", text)
        self.assertIn("Shënime me koment 31.12 / fundvit: 1", text)
        self.assertIn("Task i krijuar për J.T (në raport): 1", text)
        self.assertIn("J.T pa task real (përjashtuar): 1", text)
        self.assertIn("13.08.2026 15:50", text)

    def test_xlsx_contains_missing_and_created_next_week_tasks(self) -> None:
        workbook = load_workbook(io.BytesIO(render_xlsx(self.report)), data_only=False)
        self.assertEqual(workbook.sheetnames, ["PËRMBLEDHJE", "KONTROLLI PX JAV"])
        detail = workbook["KONTROLLI PX JAV"]
        self.assertEqual(detail.max_row, 5)
        self.assertEqual(detail.max_column, 17)
        self.assertEqual(detail["B2"].value, "PA TASK")
        self.assertEqual(detail["B3"].value, "FLAG CONVERTED, PA TASK")
        self.assertEqual(detail["B4"].value, "TASK PËR J.T")
        self.assertEqual(detail["H4"].value, "YES")
        self.assertEqual(detail["P4"].value, "31.12.2026 (FUNDVIT)")
        self.assertEqual(detail["B5"].value, "31.12 / PA TASK")
        self.assertEqual(detail["D5"].value, "31.12.2026")
        self.assertEqual(detail["I5"].value, "FUNDVIT")
        workbook.close()

    def test_docx_is_landscape_with_repeating_header_and_all_rows(self) -> None:
        payload = render_docx(self.report)
        document = Document(io.BytesIO(payload))
        section = document.sections[0]
        self.assertGreater(section.page_width, section.page_height)
        self.assertEqual(len(document.tables), 2)
        self.assertEqual(len(document.tables[1].rows), 5)
        self.assertEqual(len(document.tables[1].columns), 10)
        self.assertEqual(document.tables[1].cell(1, 1).text, "PA TASK")
        self.assertEqual(document.tables[1].cell(2, 1).text, "FLAG CONVERTED, PA TASK")
        self.assertEqual(document.tables[1].cell(3, 1).text, "TASK PËR J.T")
        self.assertIn("31.12.2026 (FUNDVIT)", document.tables[1].cell(3, 9).text)
        self.assertEqual(document.tables[1].cell(4, 1).text, "31.12 / PA TASK")
        self.assertIn("Koment: 31.12.2026", document.tables[1].cell(4, 2).text)
        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn("w:tblHeader", xml)
        self.assertIn("w:tblGrid", xml)

    def test_pdf_and_three_attachment_formats_are_generated(self) -> None:
        pdf = render_pdf(self.report)
        self.assertTrue(pdf.startswith(b"%PDF"))
        self.assertGreater(len(pdf), 2_000)
        attachments = build_attachments(self.report)
        self.assertEqual(
            [mime for _, _, mime in attachments],
            [EXCEL_MIME, WORD_MIME, PDF_MIME],
        )
        self.assertEqual(
            [name.rsplit(".", 1)[-1] for name, _, _ in attachments],
            ["xlsx", "docx", "pdf"],
        )
        self.assertTrue(all(payload for _, payload, _ in attachments))

    def test_schedule_recipient_and_timezone_match_request(self) -> None:
        entry = celery_app.conf.beat_schedule["px-jav-weekly-report-thursday-1550"]
        self.assertEqual(entry["task"], "app.celery_tasks.send_px_jav_weekly_report")
        self.assertIn("thu", str(entry["schedule"]))
        self.assertIn("15", str(entry["schedule"]))
        self.assertIn("50", str(entry["schedule"]))
        self.assertEqual(settings.PX_JAV_WEEKLY_REPORT_RECIPIENT, "334primex.eu@gmail.com")
        self.assertEqual(settings.PX_JAV_WEEKLY_REPORT_TIMEZONE, "Europe/Tirane")


if __name__ == "__main__":
    unittest.main()
