import io
import unittest
from datetime import datetime
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches

from app.services.word_standardizer import analyze_word_document, standardize_word_document


class TestWordStandardizer(unittest.TestCase):
    def _source(self) -> bytes:
        document = Document()
        document.add_heading("PrimEx Procedure", level=1)
        document.add_paragraph("This body content must remain unchanged.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "Status"
        table.cell(1, 0).text = "Header and footer"
        table.cell(1, 1).text = "Required"
        second_section = document.add_section(WD_SECTION.NEW_PAGE)
        second_section.different_first_page_header_footer = True
        document.add_paragraph("Continuation section content.")
        output = io.BytesIO()
        document.save(output)
        return output.getvalue()

    def test_analysis_detects_missing_automatic_header_and_footer(self) -> None:
        analysis = analyze_word_document(self._source(), "procedure.docx")

        self.assertEqual(analysis.sections, 2)
        self.assertEqual(analysis.tables, 1)
        self.assertFalse(all(check.compliant for check in analysis.checks))
        self.assertFalse(next(check for check in analysis.checks if check.id == "header_logo").compliant)
        self.assertFalse(next(check for check in analysis.checks if check.id == "automatic_pages").compliant)
        self.assertFalse(next(check for check in analysis.checks if check.id == "same_first_page").compliant)

    def test_standardization_creates_verified_word_fields_in_every_section(self) -> None:
        generated_at = datetime(2026, 8, 7, 12, 0, tzinfo=ZoneInfo("Europe/Belgrade"))
        result, filename, report = standardize_word_document(
            self._source(),
            "procedure.docx",
            initials="ak",
            description="PROCEDURA_E_PUNES",
            generated_at=generated_at,
        )

        self.assertEqual(filename, "PROCEDURA_E_PUNES_07.08.2026_AK.docx")
        self.assertIn("kaloi kontrollin final PrimEx", report["summary"])
        analysis = analyze_word_document(result, filename)
        self.assertTrue(all(check.compliant for check in analysis.checks))

        document = Document(io.BytesIO(result))
        self.assertEqual(document.paragraphs[1].text, "This body content must remain unchanged.")
        self.assertEqual(document.paragraphs[-1].text, "Continuation section content.")
        self.assertEqual(document.tables[0].cell(1, 1).text, "Required")
        update_fields = document.settings._element.find(qn("w:updateFields"))
        self.assertIsNotNone(update_fields)
        self.assertEqual(update_fields.get(qn("w:val")), "true")

        for section in document.sections:
            self.assertFalse(section.different_first_page_header_footer)
            self.assertGreaterEqual(section.top_margin, Inches(0.9))
            self.assertGreaterEqual(section.bottom_margin, Inches(0.9))
            header_codes = " ".join(
                node.text or "" for node in section.header._element.iter(qn("w:instrText"))
            )
            footer_codes = " ".join(
                node.text or "" for node in section.footer._element.iter(qn("w:instrText"))
            )
            self.assertIn('DATE \\@ "dd/MM/yyyy"', header_codes)
            self.assertRegex(footer_codes, r"\bPAGE\b")
            self.assertIn("NUMPAGES", footer_codes)
            self.assertTrue(any(rel.reltype == RT.IMAGE for rel in section.header.part.rels.values()))
            self.assertIn("PrimEx SH.P.K.", section.footer.tables[0].cell(0, 0).text)
            self.assertIn("Page 1 of 1", section.footer.tables[0].cell(0, 1).text)
            expected_width = round((section.page_width - section.left_margin - section.right_margin) / 635)
            for table in (section.header.tables[0], section.footer.tables[0]):
                table_width = table._tbl.tblPr.find(qn("w:tblW"))
                self.assertEqual(int(table_width.get(qn("w:w"))), expected_width)
                cell_widths = [
                    int(cell._tc.get_or_add_tcPr().find(qn("w:tcW")).get(qn("w:w")))
                    for cell in table.rows[0].cells
                ]
                self.assertEqual(sum(cell_widths), expected_width)


if __name__ == "__main__":
    unittest.main()
