from __future__ import annotations

import html
import os
import re
import textwrap
import uuid
from datetime import date, datetime, time, timedelta, timezone
from typing import Any

from sqlalchemy import or_, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.audit_log import AuditLog
from app.models.common_entry import CommonEntry
from app.models.daily_planner_snapshot import DailyPlannerSnapshot
from app.models.department import Department
from app.models.enums import CommonApprovalStatus, CommonCategory
from app.models.meeting import Meeting
from app.models.meeting_occurrence_status import MeetingOccurrenceStatus
from app.models.system_task_template import SystemTaskTemplate
from app.models.system_task_template_alignment_user import SystemTaskTemplateAlignmentUser
from app.models.task import Task
from app.models.task_assignee import TaskAssignee
from app.models.task_daily_rlz_state import TaskDailyRlzState
from app.models.user import User
from app.services.common_leave import parse_common_view_annual_leave
from app.services.daily_report_logic import business_days_between, planned_range_for_daily_report
from app.services.daily_rlz_compliance import REASON_LABELS
from app.services.primeflow_report import GmailService, report_timezone
from app.services.primeflow_report import PrimeFlowClient
from app.services.std_feedback_tickets import std_tickets_report_section
from app.services.system_task_schedule import matches_template_date
from app.services.task_title_rules import normalize_email_task_title, title_has_eight_am_indicator

REPORT_TYPE = "meetings_report"
SECTION_TITLES = [
    "A JEMI BRENDA MESATARES ME PROJEKTE?",
    "(GA) ZHV: TIKETAT E STD?",
    "SYSTEM TASK LATE",
    "DET PA PROGRES PINK (FT DHE PRJK)",
    "N- (GA) PV/FESTE?",
    "N- (GA) TAK EXT/TAK INT/BZ ME GA/BLLOK",
    "N- (GA) DET TE REJA LAST WEEK/THIS WEEK/08:00/ME DEADLINE?",
    "TAK STATUSI?",
    "N- DETYRA 1H PA SLOT?",
    "N- (GA) DET PERSONALISHT?",
    "DET E KRYERA SOT (AM/PM)",
    "DET TE SHTYERA",
    "PRODUKTE +/- SOT (PCM)",
]
DISPLAY_SECTION_TITLES = [
    SECTION_TITLES[0],  # Manual first
    SECTION_TITLES[1],  # STD tickets first among auto-filled
    SECTION_TITLES[2],
    SECTION_TITLES[3],
    SECTION_TITLES[11],
    SECTION_TITLES[10],
    SECTION_TITLES[12],
    SECTION_TITLES[7],
    SECTION_TITLES[4],
    SECTION_TITLES[6],
    SECTION_TITLES[5],
    SECTION_TITLES[8],
    SECTION_TITLES[9],
]
MANUAL_SECTION_TITLES = {
    SECTION_TITLES[0],
}
SECTION_TITLE_ALIASES = {
    "(GA) ZHV: TIKETAT E STD? RAPORTOHEN NE M3": SECTION_TITLES[1],
    "(GA) TIKETAT E STD DHE TONAT? RAPORTOHEN NE M3": SECTION_TITLES[1],
    "TAKIMET PA KRY (KONTROLLO PLATFORMEN)?": SECTION_TITLES[7],
    "TAKIMET E PA KRYERA ?": SECTION_TITLES[7],
    "DET NE PROCES SISTEMIT - SYSTEM TASKS REPORT - LATE?": SECTION_TITLES[2],
    "DET. PA PROGRES (PINK)?": SECTION_TITLES[3],
    "N- (GA) SHIKOHET COMMON VIEW NESER, VETEM DETYRAT E REJA ME TE KALTER, 08:00 DHE ME DEADLINE?": SECTION_TITLES[6],
    "N- (GA) DET TE REJA LAST WEEK DHE THIS WEEK, 08:00, ME DEADLINE?": SECTION_TITLES[6],
    "N- (GA) DET TE REJA LAST WEEK DHE THIS WEEK": SECTION_TITLES[6],
    "DET TE REJA LAST WEEK DHE THIS WEEK": SECTION_TITLES[6],
    "N- (GA) TAKIMET EXTERNE/ TAKIMET INTERNE/ BZ ME GA/BLLOK?": SECTION_TITLES[5],
    "N- A KA DETYRA 1H PA SLOT?": SECTION_TITLES[8],
    "(GA/KA) KUSH KA DET PERSONALISHT?": SECTION_TITLES[9],
    "N- (GA/KA) KUSH KA DET PERSONALISHT?": SECTION_TITLES[9],
}
RETIRED_CLOSING_SECTION_TITLES = {
    "GA MBYLLJA E DET",
    "HV MBYLLJA E DET",
    "(GA) M3 DET GA MBYLLJA ME HV?",
    "(GA) M3 DET GA MBYLLJA ME HV/OH?",
}
RETIRED_CLOSING_SECTION_KEYS = {
    re.sub(r"[^A-Z0-9]+", "", title.upper()) for title in RETIRED_CLOSING_SECTION_TITLES
}


def is_retired_meetings_section_title(title: str | None) -> bool:
    return _compact_section_title(title) in RETIRED_CLOSING_SECTION_KEYS
DEFAULT_MANUAL_BODY = "(Ploteso manualisht)"
# Same rule for M1/M2/M3: personal rows only when the title marks GA (not KA).
PERSONAL_GA = re.compile(r"[/:]\s*GA\b", re.I)
TECHNICAL_TAG = re.compile(r"\[\[\s*/?\s*(?:added|done)\s*\]\]", re.I)
DUE_SUFFIX = re.compile(r"\s+due\s+\d{1,2}:\d{2}\s*$", re.I)
COMPLETED_PRODUCTS = re.compile(r"completed_products\s*[:=]\s*(\d+)", re.I)
TITLE_PREFIX = re.compile(r"^[A-Z]{1,4}(?:/[A-Z]{1,4})?\s*:\s*", re.I)
TASK_LINE_STATUS = re.compile(r"^\[([A-Z_]+)\]\s*")
MEETING_HIGHLIGHT_MARKER = "[[mt:non_daily_weekly]]"
MEETING_HIGHLIGHT_PATTERN = re.compile(r"\s*\[\[\s*mt\s*:\s*non_daily_weekly\s*\]\]", re.I)
ALL_ASSIGNEES_DISPLAY_THRESHOLD = 10
WEEKLY_PLANNER_DEPARTMENT_ORDER = {"DEV": 0, "GD": 1, "PCM": 2}


def _compact_section_title(value: str | None) -> str:
    return re.sub(r"[^A-Z0-9]+", "", (value or "").upper())


def canonical_meetings_section_title(raw_title: str | None) -> str:
    """Map Common View near-duplicates onto the built-in M3 section titles."""
    raw = (raw_title or "").strip()
    if not raw:
        return raw
    if raw in SECTION_TITLE_ALIASES:
        return SECTION_TITLE_ALIASES[raw]

    compact = _compact_section_title(raw)
    for alias, canonical in SECTION_TITLE_ALIASES.items():
        if _compact_section_title(alias) == compact:
            return canonical
    for known in DISPLAY_SECTION_TITLES:
        if _compact_section_title(known) == compact:
            return known

    # Wording variants that still mean the same auto-filled questions.
    if "TIKETATESTD" in compact and "RAPORTOHENNEM3" in compact:
        return SECTION_TITLES[1]
    return raw


def next_working_day(day: date) -> date:
    candidate = day + timedelta(days=1)
    while candidate.weekday() >= 5:
        candidate += timedelta(days=1)
    return candidate


def subject_for(day: date) -> str:
    return f"M3 - PrimeFlow Mbyllja e dites - {day:%d.%m.%Y}"


def is_generated_subject(subject: str | None, day: date) -> bool:
    """Recognize the old and current automatically generated M3 subjects."""
    return (subject or "").strip() in {
        subject_for(day),
        f"PrimeFlow Mbyllja e dites M3 - {day:%d.%m.%Y}",
    }


def _prefer_section_body(current: str, incoming: str) -> str:
    """Prefer real content over the manual placeholder when collapsing aliases."""
    cur = (current or "").strip()
    inc = (incoming or "").strip()
    if not cur or cur == DEFAULT_MANUAL_BODY:
        return incoming if incoming else current
    if not inc or inc == DEFAULT_MANUAL_BODY:
        return current
    return current


def normalize_meetings_report_sections(sections: list[dict[str, Any]] | None) -> list[dict[str, str]]:
    """Keep saved drafts aligned with the current M3 question list without losing edits."""
    existing_sections = sections or []
    by_title: dict[str, dict[str, str]] = {}
    unknown_sections: list[dict[str, str]] = []
    seen_unknown: set[str] = set()
    for section in existing_sections:
        raw_title = str(section.get("title") or "").strip()
        section_key = str(section.get("section_key") or "").strip()
        identity = section_key or raw_title
        if is_retired_meetings_section_title(identity):
            continue
        title = canonical_meetings_section_title(identity)
        if not title:
            continue
        body = str(section.get("body") or "")
        if title in DISPLAY_SECTION_TITLES:
            if title not in by_title:
                by_title[title] = {
                    "section_key": title,
                    "title": raw_title if section_key else title,
                    "body": body,
                }
            else:
                by_title[title]["body"] = _prefer_section_body(by_title[title]["body"], body)
            continue

        compact = _compact_section_title(title)
        if not compact or compact in seen_unknown:
            continue
        seen_unknown.add(compact)
        unknown_sections.append({
            "section_key": section_key or f"manual:{compact}",
            "title": raw_title or title,
            "body": body,
        })

    ordered: list[dict[str, str]] = []
    for title in DISPLAY_SECTION_TITLES:
        if title in by_title:
            ordered.append(by_title[title])
        elif title in MANUAL_SECTION_TITLES:
            ordered.append({"section_key": title, "title": title, "body": DEFAULT_MANUAL_BODY})

    # Keep Common View–synced manuals with the other manuals (after built-in manuals).
    if not unknown_sections:
        return ordered
    insert_at = 0
    for index, section in enumerate(ordered):
        if section["title"] in MANUAL_SECTION_TITLES:
            insert_at = index + 1
    return ordered[:insert_at] + unknown_sections + ordered[insert_at:]


def _local_date(value: datetime | date | None) -> date | None:
    if value is None:
        return None
    if isinstance(value, datetime):
        if value.tzinfo:
            return value.astimezone(report_timezone()).date()
        return value.date()
    return value


def _audit_local_date(value: Any) -> date | None:
    """Normalize an audited date value using the M3 report timezone."""
    if value is None:
        return None
    if isinstance(value, (datetime, date)):
        return _local_date(value)
    raw = str(value).strip()
    if not raw:
        return None
    try:
        parsed = datetime.fromisoformat(raw.replace("Z", "+00:00"))
    except ValueError:
        try:
            return date.fromisoformat(raw[:10])
        except ValueError:
            return None
    return _local_date(parsed)


def _is_postponed_for_m3_day(task: Task, events: list[AuditLog], report_day: date) -> bool:
    """Return true when today's audit trail moved due/start date to a later day."""
    return any(
        event.action in {"task.due_date_changed", "task.start_date_changed"}
        and _audit_local_date((event.before or {}).get("value")) == report_day
        and (
            changed_date := _audit_local_date((event.after or {}).get("value"))
        ) is not None
        and changed_date > report_day
        and (
            _local_date(task.due_date if event.action == "task.due_date_changed" else task.start_date)
            == changed_date
        )
        for event in events
    )


def _postponement_events_for_m3_task(
    task: Task, events: list[AuditLog], report_day: date
) -> list[AuditLog]:
    return [
        event for event in events
        if event.action in {"task.due_date_changed", "task.start_date_changed"}
        and _audit_local_date((event.before or {}).get("value")) == report_day
        and (
            changed_date := _audit_local_date((event.after or {}).get("value"))
        ) is not None
        and changed_date > report_day
        and (
            _local_date(task.due_date if event.action == "task.due_date_changed" else task.start_date)
            == changed_date
        )
    ]


def _daily_baseline_task_ids(snapshots: list[DailyPlannerSnapshot]) -> set[uuid.UUID]:
    """Collect task IDs which belonged to at least one user's original daily plan."""
    task_ids: set[uuid.UUID] = set()
    for snapshot in snapshots:
        for person in (snapshot.payload or {}).get("people") or []:
            for item in person.get("tasks") or []:
                try:
                    task_ids.add(uuid.UUID(str(item.get("task_id"))))
                except (TypeError, ValueError):
                    continue
    return task_ids


async def _postponed_tasks_for_m3_day(
    db: AsyncSession,
    report_day: date,
    tasks: list[Task],
) -> tuple[list[Task], dict[Any, tuple[str, str]], list[Task], dict[Any, tuple[str, str]]]:
    """Load report-day postponed tasks without changing Realization state or code."""
    snapshots = (
        await db.execute(
            select(DailyPlannerSnapshot).where(DailyPlannerSnapshot.day_date == report_day)
        )
    ).scalars().all()
    baseline_task_ids = _daily_baseline_task_ids(list(snapshots))
    if not baseline_task_ids:
        return [], {}, [], {}

    local_start = datetime.combine(report_day, time.min, tzinfo=report_timezone())
    start_utc = local_start.astimezone(timezone.utc)
    end_utc = (local_start + timedelta(days=1)).astimezone(timezone.utc)
    events = (
        await db.execute(
            select(AuditLog)
            .where(
                AuditLog.entity_type == "task",
                AuditLog.entity_id.in_(baseline_task_ids),
                AuditLog.action.in_(("task.due_date_changed", "task.start_date_changed")),
                AuditLog.created_at >= start_utc,
                AuditLog.created_at < end_utc,
            )
            .order_by(AuditLog.created_at, AuditLog.id)
        )
    ).scalars().all()
    events_by_task: dict[Any, list[AuditLog]] = {}
    for event in events:
        events_by_task.setdefault(event.entity_id, []).append(event)

    postponed: list[Task] = []
    date_ranges: dict[Any, tuple[str, str]] = {}
    postponed_both: list[Task] = []
    both_date_ranges: dict[Any, tuple[str, str]] = {}
    for task in tasks:
        if task.id not in baseline_task_ids:
            continue
        matching_events = _postponement_events_for_m3_task(
            task, events_by_task.get(task.id, []), report_day
        )
        if not matching_events:
            continue
        due_events = [e for e in matching_events if e.action == "task.due_date_changed"]
        start_events = [e for e in matching_events if e.action == "task.start_date_changed"]
        if due_events and start_events:
            postponed_both.append(task)
            start_before = _audit_local_date((start_events[-1].before or {}).get("value"))
            start_after = _audit_local_date((start_events[-1].after or {}).get("value"))
            due_before = _audit_local_date((due_events[-1].before or {}).get("value"))
            due_after = _audit_local_date((due_events[-1].after or {}).get("value"))
            both_date_ranges[task.id] = (
                # ``|`` is the column delimiter used by the ASCII table
                # renderer.  Keeping it inside a cell makes the report
                # parser treat the DUE portion as another column and shifts
                # the title into the TITULLI column.  Use a visual separator
                # that is safe inside a cell instead.
                f"START: {start_before.strftime('%d.%m.%Y') if start_before else '-'} / DUE: {due_before.strftime('%d.%m.%Y') if due_before else '-'}",
                f"START: {start_after.strftime('%d.%m.%Y') if start_after else '-'} / DUE: {due_after.strftime('%d.%m.%Y') if due_after else '-'}",
            )
        elif due_events:
            postponed.append(task)
            matching_event = due_events[-1]
            before_date = _audit_local_date((matching_event.before or {}).get("value"))
            after_date = _audit_local_date((matching_event.after or {}).get("value"))
            date_ranges[task.id] = (
                before_date.strftime("%d.%m.%Y") if before_date else report_day.strftime("%d.%m.%Y"),
                after_date.strftime("%d.%m.%Y") if after_date else "-",
            )
    return postponed, date_ranges, postponed_both, both_date_ranges


def _local_time(value: datetime | None) -> str:
    if value is None:
        return "-"
    local = value.astimezone(report_timezone()) if value.tzinfo else value
    return local.strftime("%H:%M")


def _task_day(task: Task) -> date | None:
    return _local_date(task.due_date or task.start_date or task.created_at)


def _is_new_task_for_m3_day(task: Task, day: date) -> bool:
    """Whether a task is new for the M3 report generated for ``day``.

    A task is reported as new once only: on its planned start date.  Its
    creation date is intentionally not used here; it only determines the
    ``Last W`` / ``This W`` label shown in the table.
    """
    return _is_open(task) and _local_date(task.start_date) == day


def _is_system_task(task: Task) -> bool:
    """System-template occurrences are intentionally excluded from M3 new-task review."""
    return bool(
        getattr(task, "system_template_origin_id", None)
        or getattr(task, "system_task_slot_id", None)
    )


def _completed_tasks_for_report_day(tasks: list[Task], report_day: date) -> list[Task]:
    """Non-system tasks that are still completed and were closed on the M3 day."""
    return [
        task
        for task in tasks
        if not _is_system_task(task)
        and getattr(task, "completed_at", None) is not None
        and _local_date(task.completed_at) == report_day
        and _normalize_report_status(task.status) == "DONE"
    ]


def _meeting_occurs_on_date(meeting: Meeting, day: date) -> bool:
    recurrence = (meeting.recurrence_type or "").lower()
    if recurrence == "weekly":
        return bool(meeting.recurrence_days_of_week and day.weekday() in meeting.recurrence_days_of_week)
    if recurrence == "monthly":
        return bool(meeting.recurrence_days_of_month and day.day in meeting.recurrence_days_of_month)
    if recurrence == "yearly":
        month = meeting.starts_at.month if meeting.starts_at else None
        day_value = meeting.recurrence_days_of_month[0] if meeting.recurrence_days_of_month else None
        return bool(month and day_value and day.month == month and day.day == day_value)
    return _local_date(meeting.starts_at or meeting.created_at) == day


def _is_open(task: Task) -> bool:
    return not task.completed_at and str(task.status or "").upper() not in {"DONE", "COMPLETED"}


def _late_days(task: Task) -> int:
    due_day = _local_date(task.due_date)
    if due_day is None:
        return 0
    today = datetime.now(report_timezone()).date()
    if today <= due_day:
        return 0
    return business_days_between(due_day, today)


def _late_days_label(days: int) -> str:
    if days <= 0:
        return "-"
    return f"{days} day" if days == 1 else f"{days} days"


def _late_days_tyo_label(days: int) -> str:
    if days <= 0:
        return "-"
    return "Y" if days == 1 else str(days)


def _initials(name: str | None) -> str:
    parts = re.findall(r"[^\W\d_]+", name or "", flags=re.UNICODE)
    return "".join(part[0] for part in parts).upper() or "-"


def _is_report_all_participant(user: User) -> bool:
    return user.is_active and _initials(user.full_name) not in {"GA", "KA", "HV"}


async def _all_participant_user_ids(db: AsyncSession) -> set[Any]:
    users = (await db.execute(select(User).where(User.is_active.is_(True)))).scalars().all()
    return {user.id for user in users if _is_report_all_participant(user)}


async def _assignee_names(db: AsyncSession, tasks: list[Task]) -> dict[Any, str]:
    user_ids = {task.assigned_to for task in tasks if task.assigned_to}
    task_ids = [task.id for task in tasks]
    if task_ids:
        assignee_user_ids = (
            await db.execute(select(TaskAssignee.user_id).where(TaskAssignee.task_id.in_(task_ids)))
        ).scalars().all()
        user_ids.update(assignee_user_ids)
    if not user_ids:
        return {}
    users = (await db.execute(select(User).where(User.id.in_(user_ids)))).scalars().all()
    return {user.id: user.full_name or user.email for user in users}


async def _effective_task_assignee_ids(db: AsyncSession, tasks: list[Task]) -> dict[Any, set[Any]]:
    result: dict[Any, set[Any]] = {task.id: set() for task in tasks}
    for task in tasks:
        if task.assigned_to:
            result.setdefault(task.id, set()).add(task.assigned_to)
    task_ids = [task.id for task in tasks]
    if not task_ids:
        return result
    rows = (
        await db.execute(select(TaskAssignee.task_id, TaskAssignee.user_id).where(TaskAssignee.task_id.in_(task_ids)))
    ).all()
    for task_id, user_id in rows:
        result.setdefault(task_id, set()).add(user_id)
    return result


def _weekly_planner_department_code(department_id: Any, department_codes: dict[Any, str]) -> str:
    code = str(department_codes.get(department_id) or "").strip().upper()
    aliases = {
        "DEVELOPMENT": "DEV",
        "GRAPHIC DESIGN": "GD",
        "GDS": "GD",
        "PRODUCT CONTENT": "PCM",
        "PROJECT CONTENT MANAGER": "PCM",
    }
    return aliases.get(code, code or "-")


async def apply_weekly_planner_task_order(
    db: AsyncSession,
    tasks: list[Task],
    assignee_ids_by_task: dict[Any, set[Any]],
    department_codes: dict[Any, str] | None = None,
) -> None:
    """Attach the Weekly Planner department/person order to report task rows.

    The Weekly Planner is the source of truth: department order is DEV, GD,
    PCM, and users use their saved ``weekly_planner_sort_order`` within each
    department.  The attribute is request-local on the loaded ORM objects and
    is consumed by the existing task table renderers.
    """
    if not tasks:
        return
    codes = department_codes or {
        department_id: code
        for department_id, code in (await db.execute(select(Department.id, Department.code))).all()
    }
    user_ids = {
        user_id
        for task in tasks
        for user_id in ({task.assigned_to} if task.assigned_to else set()) | assignee_ids_by_task.get(task.id, set())
    }
    users = (
        await db.execute(select(User).where(User.id.in_(user_ids)))
    ).scalars().all() if user_ids else []
    users_by_id = {user.id: user for user in users}

    def user_key(user_id: Any) -> tuple[int, int, str]:
        user = users_by_id.get(user_id)
        if user is None:
            return (1, 0, "~")
        return (
            1 if user.weekly_planner_sort_order is None else 0,
            user.weekly_planner_sort_order or 0,
            (user.full_name or user.username or user.email or "").casefold(),
        )

    for task in tasks:
        assignee_ids = set(assignee_ids_by_task.get(task.id, set()))
        if task.assigned_to:
            assignee_ids.add(task.assigned_to)
        primary_user_id = task.assigned_to if task.assigned_to in users_by_id else None
        if primary_user_id is None and assignee_ids:
            primary_user_id = min(assignee_ids, key=user_key)
        primary_user = users_by_id.get(primary_user_id)
        department_id = primary_user.department_id if primary_user and primary_user.department_id else task.department_id
        department_code = _weekly_planner_department_code(department_id, codes)
        # Report tables must show the department of the person responsible for
        # the task, never the department of the project that owns the task.
        # Keep it on the request-local task instance alongside the existing
        # Weekly Planner ordering metadata.
        setattr(task, "_report_user_department_code", department_code if primary_user else "-")
        setattr(
            task,
            "_weekly_planner_report_sort",
            (
                WEEKLY_PLANNER_DEPARTMENT_ORDER.get(department_code, len(WEEKLY_PLANNER_DEPARTMENT_ORDER)),
                department_code.casefold(),
                *user_key(primary_user_id),
            ),
        )


async def weekly_planner_user_sort_keys(
    db: AsyncSession,
    user_ids: set[Any],
    department_codes: dict[Any, str] | None = None,
) -> dict[Any, tuple[int, str, int, int, str]]:
    """Return the same department/person order used by report task tables."""
    if not user_ids:
        return {}
    codes = department_codes or {
        department_id: code
        for department_id, code in (await db.execute(select(Department.id, Department.code))).all()
    }
    users = (await db.execute(select(User).where(User.id.in_(user_ids)))).scalars().all()
    result: dict[Any, tuple[int, str, int, int, str]] = {}
    for user in users:
        department_code = _weekly_planner_department_code(user.department_id, codes)
        result[user.id] = (
            WEEKLY_PLANNER_DEPARTMENT_ORDER.get(department_code, len(WEEKLY_PLANNER_DEPARTMENT_ORDER)),
            department_code.casefold(),
            1 if user.weekly_planner_sort_order is None else 0,
            user.weekly_planner_sort_order or 0,
            (user.full_name or user.username or user.email or "").casefold(),
        )
    return result


async def _users_by_initials(db: AsyncSession, initials: str) -> list[User]:
    users = (await db.execute(select(User).where(User.is_active.is_(True)))).scalars().all()
    target = initials.upper()
    return [
        user for user in users
        if _initials(user.full_name) == target
        or (user.username or "").strip().upper() == target
        or (user.email or "").split("@", 1)[0].strip().upper() == target
    ]


def _task_owners(
    task: Task,
    names: dict[Any, str],
    assignee_ids_by_task: dict[Any, set[Any]] | None = None,
    *,
    all_participant_ids: set[Any] | None = None,
) -> str:
    assignee_ids = set(assignee_ids_by_task.get(task.id, set()) if assignee_ids_by_task else set())
    if not assignee_ids and task.assigned_to:
        assignee_ids = {task.assigned_to}
    if (
        (all_participant_ids and assignee_ids and all_participant_ids.issubset(assignee_ids))
        or len(assignee_ids) > ALL_ASSIGNEES_DISPLAY_THRESHOLD
    ):
        return "ALL"
    owners = sorted({_initials(names.get(user_id)) for user_id in assignee_ids if _initials(names.get(user_id)) != "-"})
    return " ".join(owners) or _initials(names.get(task.assigned_to))


def common_view_task_sort_key(
    task: Task,
    names: dict[Any, str],
    assignee_ids_by_task: dict[Any, set[Any]] | None = None,
    *,
    all_participant_ids: set[Any] | None = None,
) -> tuple:
    """Mirror Common View compareTaskOrder for report task tables (M1/M2/M3)."""
    owner = _task_owners(task, names, assignee_ids_by_task, all_participant_ids=all_participant_ids)
    existing_order = (
        0 if bool(task.is_deadline_important) else 1,
        0 if title_has_eight_am_indicator(task.title) else 1,
        owner.casefold(),
        task.fast_task_order if task.fast_task_order is not None else 10**9,
        _clean_task_title(task.title).casefold(),
        str(task.created_at or ""),
    )
    weekly_planner_order = getattr(task, "_weekly_planner_report_sort", None)
    return (*weekly_planner_order, *existing_order) if weekly_planner_order else existing_order


# Backwards-compatible alias used by older call sites / tests.
_m3_task_sort_key = common_view_task_sort_key


def common_view_item_sort_key(item: dict[str, Any]) -> tuple:
    """Same ordering for Common View payload dicts used in report fallbacks."""
    owner = _common_owner(item)
    title = _common_title(item)
    important = bool(item.get("is_deadline_important") or item.get("isDeadlineImportant"))
    eight_am = title_has_eight_am_indicator(title)
    order = item.get("fast_task_order")
    if order is None:
        order = item.get("fastTaskOrder")
    if not isinstance(order, int):
        order = 10**9
    existing_order = (
        0 if important else 1,
        0 if eight_am else 1,
        owner.casefold(),
        order,
        title.casefold(),
        str(item.get("created_at") or item.get("createdAt") or ""),
    )
    weekly_planner_order = item.get("weekly_planner_sort") or item.get("weeklyPlannerSort")
    return (*tuple(weekly_planner_order), *existing_order) if weekly_planner_order else existing_order


def _task_line(
    task: Task,
    names: dict[Any, str],
    assignee_ids_by_task: dict[Any, set[Any]] | None = None,
    *,
    include_status: bool = False,
    all_participant_ids: set[Any] | None = None,
) -> str:
    owner = _task_owners(
        task, names, assignee_ids_by_task, all_participant_ids=all_participant_ids
    )
    title = _clean_task_title(task.title)
    line = f"- {owner}: {title}"
    if include_status:
        return f"- [{_normalize_report_status(task.status)}] {owner}: {title}"
    return line


def _task_line_with_late_days(task: Task, names: dict[Any, str]) -> str:
    owner = _initials(names.get(task.assigned_to))
    title_lines = _wrap_fixed_width(_clean_task_title(task.title), 48)
    days = _late_days(task)
    late_label = _late_days_label(days)
    lines = [f"- {owner:<4} | {title_lines[0]:<48} | {late_label}"]
    lines.extend(f"  {'':<4} | {line:<48} |" for line in title_lines[1:])
    return "\n".join(lines)


def _wrap_fixed_width(value: str, width: int) -> list[str]:
    cleaned = re.sub(r"\s+", " ", value).strip()
    return textwrap.wrap(cleaned, width=width, break_long_words=False, break_on_hyphens=False) or ["-"]


def _m3_task_type_label(task: Task) -> str:
    """Task kind for M3 GA/HV tables: system, project, or fast (with subtype).

    GA/PX note origins are sources, not types. Fast tasks show FT, or a subtype
    when flagged (1H/BLL/R1/P).
    """
    if getattr(task, "system_template_origin_id", None) is not None:
        return "SYS"
    if getattr(task, "project_id", None) is not None:
        return "PRJK"
    # Fast-task subtypes take priority over generic FT
    if getattr(task, "is_bllok", False):
        return "BLL"
    if getattr(task, "is_r1", False):
        return "R1"
    if getattr(task, "is_1h_report", False):
        return "1H"
    if getattr(task, "is_personal", False):
        return "P"
    return "FT"


def _m3_department_code_label(department_id: Any, department_codes: dict[Any, str] | None) -> str:
    """Return the short M3 label for a department id."""
    codes = department_codes or {}
    code = codes.get(department_id)
    if not code and department_id is not None:
        # Common View is an HTTP payload, so its UUIDs arrive as strings while
        # the report's database map is keyed by UUID objects.
        department_id_value = str(department_id)
        code = next(
            (value for key, value in codes.items() if str(key) == department_id_value),
            None,
        )
    code = str(code or "").strip().upper()
    aliases = {
        "PRODUCT CONTENT": "PCM",
        "PROJECT CONTENT MANAGER": "PCM",
        "DEVELOPMENT": "DEV",
        "GRAPHIC DESIGN": "GD",
        "GDS": "GD",
        "FINANCE": "FIN",
    }
    return aliases.get(code, code or "-")


def _m3_department_label(task: Task, department_codes: dict[Any, str] | None) -> str:
    """Return the responsible user's department for report task tables.

    ``apply_weekly_planner_task_order`` resolves the primary assignee and adds
    the code to each report task. The fallback supports isolated legacy callers
    that have not been prepared by that shared report pipeline.
    """
    user_department = str(getattr(task, "_report_user_department_code", "") or "").strip()
    if user_department:
        return _m3_department_code_label(user_department, {user_department: user_department})
    return _m3_department_code_label(getattr(task, "department_id", None), department_codes)


def _m3_am_pm_label(task: Task) -> str:
    period = str(getattr(task, "finish_period", None) or "").strip().upper()
    return period if period in {"AM", "PM", "AM/PM"} else "-"


def _m3_added_week_label(task: Task, week_start: date | None) -> str:
    created_day = _local_date(getattr(task, "created_at", None))
    if week_start is not None and created_day is not None and created_day >= week_start:
        return "This W"
    return "Last W"


def _m3_product_counts(task: Task) -> tuple[int, int] | None:
    """Planned and completed product counts for a PCM product task.

    ``daily_products`` is the day's planned count; the completed count is stored
    in ``internal_notes`` as ``completed_products=N``.
    """
    planned = getattr(task, "daily_products", None)
    if planned is None:
        return None
    planned = int(planned)
    if planned <= 0:
        return None
    match = COMPLETED_PRODUCTS.search(getattr(task, "internal_notes", None) or "")
    if not match:
        return planned, 0
    try:
        return planned, max(0, int(match.group(1)))
    except ValueError:
        return planned, 0


def _m3_product_delta_label(task: Task) -> str:
    counts = _m3_product_counts(task)
    if counts is None:
        return "-"
    planned, done = counts
    delta = done - planned
    return f"{done}/{planned} ({'+' if delta > 0 else ''}{delta})"


def _product_delta_tasks_for_m3_day(tasks: list[Task], report_day: date) -> list[Task]:
    """PCM product tasks that closed the day above or below their planned count.

    Days with no progress at all are left out: they already appear in the pink
    ``DET PA PROGRES`` section, so repeating them here would only add noise.
    """
    selected: list[Task] = []
    for task in tasks:
        # PCM product rows are PRODUCT-phase tasks. CONTROL tasks can also
        # carry a copied daily_products total after editing, but must not be
        # reported as product production rows.
        phase = str(getattr(task, "phase", None) or "PRODUCT").strip().upper()
        if phase != "PRODUCT":
            continue
        counts = _m3_product_counts(task)
        if counts is None:
            continue
        if _task_day(task) != report_day and _local_date(task.completed_at) != report_day:
            continue
        planned, done = counts
        if done <= 0 or done == planned:
            continue
        selected.append(task)
    return selected


M3_DEPARTMENT_ORDER = WEEKLY_PLANNER_DEPARTMENT_ORDER


def _m3_department_sort_key(task: Task, department_codes: dict[Any, str] | None) -> int:
    """Keep M3 department tables in the requested DEV, GD, PCM order."""
    return M3_DEPARTMENT_ORDER.get(_m3_department_label(task, department_codes), len(M3_DEPARTMENT_ORDER))


def _m3_status_table(
    status_label: str,
    tasks: list[Task],
    names: dict[Any, str],
    *,
    include_late_days: bool = False,
    with_status: bool = False,
    include_type: bool = False,
    include_department: bool = False,
    include_added_week: bool = False,
    include_am_pm: bool = False,
    department_codes: dict[Any, str] | None = None,
    week_start: date | None = None,
    assignee_ids_by_task: dict[Any, set[Any]] | None = None,
    all_participant_ids: set[Any] | None = None,
    daily_rlz_by_task: dict[Any, tuple[str, str]] | None = None,
    date_range_by_task: dict[Any, tuple[str, str]] | None = None,
    products_by_task: dict[Any, str] | None = None,
) -> list[str]:
    columns: list[tuple[str, int]] = [("NR", 2), ("KUSH", 5)]
    if include_department:
        columns.append(("DEP", 5))
    if include_added_week:
        columns.append(("KRIJUAR", 7))
    if include_am_pm:
        columns.append(("AM/PM", 5))
    if include_type:
        columns.append(("LLOJI", 7))
    if date_range_by_task is not None:
        columns.extend((("NGA", 10), ("NE", 10)))
    if products_by_task is not None:
        columns.append(("PRODUKTE", 12))
    if include_late_days:
        columns.append(("T/Y/O", 5))
    columns.append(("TITULLI", 64))
    if daily_rlz_by_task is not None:
        columns.extend((("ARSYEJA", 24), ("KOMENT", 36)))

    border = "+" + "+".join("-" * (width + 2) for _, width in columns) + "+"

    def table_row(values: list[str]) -> str:
        return "| " + " | ".join(
            f"{value:<{width}}" for value, (_, width) in zip(values, columns)
        ) + " |"

    rows = [
        f"{status_label}:",
        border,
        table_row([label for label, _ in columns]),
        border,
    ]
    empty_title = "(Asnje detyre)"
    if not tasks:
        values = ["-", "-"]
        if include_department:
            values.append("-")
        if include_added_week:
            values.append("-")
        if include_am_pm:
            values.append("-")
        if include_type:
            values.append("-")
        if date_range_by_task is not None:
            values.extend(("-", "-"))
        if products_by_task is not None:
            values.append("-")
        if include_late_days:
            values.append("-")
        values.append(empty_title)
        if daily_rlz_by_task is not None:
            values.extend(("-", "-"))
        rows.append(table_row(values))
        rows.append(border)
        return rows
    def sort_key(item: Task) -> tuple:
        ordered_key = common_view_task_sort_key(
            item, names, assignee_ids_by_task, all_participant_ids=all_participant_ids
        )
        # Completed work is always listed after active work in a task table.
        done_last = 1 if getattr(item, "completed_at", None) or _normalize_report_status(item.status) == "DONE" else 0
        if getattr(item, "_weekly_planner_report_sort", None):
            return (done_last, *ordered_key)
        # Preserve the legacy order for direct callers that do not have the
        # Weekly Planner metadata available.
        return (
            done_last,
            0 if include_type and getattr(item, "system_template_origin_id", None) is not None else 1,
            _m3_department_sort_key(item, department_codes) if include_department else 0,
            *ordered_key,
        )

    ordered = sorted(tasks, key=sort_key)
    for index, task in enumerate(ordered, start=1):
        owner = _task_owners(
            task, names, assignee_ids_by_task, all_participant_ids=all_participant_ids
        )
        task_type = _m3_task_type_label(task) if include_type else ""
        department = _m3_department_label(task, department_codes) if include_department else ""
        added_week = _m3_added_week_label(task, week_start) if include_added_week else ""
        am_pm = _m3_am_pm_label(task) if include_am_pm else ""
        postponed_from, postponed_to = (date_range_by_task or {}).get(task.id, ("-", "-"))
        display_title = _clean_task_title(task.title)
        title_lines = _wrap_fixed_width(display_title, 64)
        reason, comment = (daily_rlz_by_task or {}).get(task.id, ("-", "-"))
        reason_lines = _wrap_fixed_width(reason or "-", 24) if daily_rlz_by_task is not None else []
        comment_lines = _wrap_fixed_width(comment or "-", 36) if daily_rlz_by_task is not None else []
        status = _normalize_report_status(task.status) if with_status else None
        padded_titles = _append_status_marker_to_lines(title_lines, status, 64)
        values = [str(index), owner]
        if include_department:
            values.append(department)
        if include_added_week:
            values.append(added_week)
        if include_am_pm:
            values.append(am_pm)
        if include_type:
            values.append(task_type)
        if date_range_by_task is not None:
            values.extend((postponed_from, postponed_to))
        if products_by_task is not None:
            values.append(products_by_task.get(task.id, "-"))
        if include_late_days:
            values.append(_late_days_tyo_label(_late_days(task)))
        values.append(padded_titles[0])
        if daily_rlz_by_task is not None:
            values.extend((reason_lines[0], comment_lines[0]))
        rows.append(table_row(values))
        continuation_count = max(len(padded_titles), len(reason_lines), len(comment_lines))
        for line_index in range(1, continuation_count):
            continuation = ["", ""]
            if include_department:
                continuation.append("")
            if include_added_week:
                continuation.append("")
            if include_am_pm:
                continuation.append("")
            if include_type:
                continuation.append("")
            if date_range_by_task is not None:
                continuation.extend(("", ""))
            if products_by_task is not None:
                continuation.append("")
            if include_late_days:
                continuation.append("")
            continuation.append(padded_titles[line_index] if line_index < len(padded_titles) else "")
            if daily_rlz_by_task is not None:
                continuation.extend((
                    reason_lines[line_index] if line_index < len(reason_lines) else "",
                    comment_lines[line_index] if line_index < len(comment_lines) else "",
                ))
            rows.append(table_row(continuation))
        rows.append(border)
    return rows


def _task_late_lines(tasks: list[Task], names: dict[Any, str]) -> list[str]:
    if not tasks:
        return ["(Asnje detyre)"]
    ordered = sorted(
        tasks,
        key=lambda item: (-_late_days(item), *common_view_task_sort_key(item, names, None)),
    )
    return [_task_line_with_late_days(task, names) for task in ordered]


def _clean_task_title(value: str | None) -> str:
    """Normalize task titles for reports.

    Titles often store post-create edits as ``[[added]]...[[/added]]``. Keep that
    text (it is usually the real description) and only strip the markers.
    """
    cleaned = TECHNICAL_TAG.sub("", value or "")
    candidates = [line.strip() for line in cleaned.splitlines() if line.strip()]
    title_line = next((line for line in candidates if TITLE_PREFIX.search(line)), "")
    if not title_line:
        title_line = next((line for line in candidates if not re.match(r"^\d+\.", line)), "")
    title_line = DUE_SUFFIX.sub("", title_line)
    return normalize_email_task_title(re.sub(r"\s+", " ", title_line).strip()) or "-"


def _empty_aware(lines: list[str]) -> str:
    return "\n".join(lines) if lines else "(Asnje detyre)"


def _week_start(day: date) -> date:
    return day - timedelta(days=day.weekday())


def _is_without_progress_for_m3_day(task: Task, day: date) -> bool:
    """Match the unfinished pink rows shown in Daily Report My View.

    A task remains relevant throughout its planned range (and while overdue),
    rather than only on its due date. This is important for project tasks that
    start today but have a later deadline.
    """
    if _is_system_task(task) or not _is_open(task) or _normalize_report_status(task.status) != "TODO":
        return False
    planned_start, planned_end = planned_range_for_daily_report(task, None)
    if planned_start is None or planned_end is None:
        return False
    return planned_start <= day


async def _daily_rlz_values_by_task(
    db: AsyncSession,
    tasks: list[Task],
    day: date,
    names: dict[Any, str],
    assignee_ids_by_task: dict[Any, set[Any]],
) -> dict[Any, tuple[str, str]]:
    """Return the day-scoped reason/comment saved in Daily Report My View."""
    task_ids = [task.id for task in tasks]
    if not task_ids:
        return {}
    states = (
        await db.execute(
            select(TaskDailyRlzState)
            .where(TaskDailyRlzState.task_id.in_(task_ids))
            .where(TaskDailyRlzState.day_date == day)
        )
    ).scalars().all()
    grouped: dict[Any, list[TaskDailyRlzState]] = {}
    for state in states:
        grouped.setdefault(state.task_id, []).append(state)

    result: dict[Any, tuple[str, str]] = {}
    for task_id, task_states in grouped.items():
        task_states.sort(key=lambda state: (_initials(names.get(state.user_id)), str(state.user_id)))
        show_owner = len(assignee_ids_by_task.get(task_id, set())) > 1 or len(task_states) > 1
        reasons: list[str] = []
        comments: list[str] = []
        for state in task_states:
            owner = _initials(names.get(state.user_id))
            prefix = f"{owner}: " if show_owner else ""
            reason = REASON_LABELS.get(state.reason_code) if state.reason_code else None
            if reason:
                reasons.append(f"{prefix}{reason}")
            if state.comment and state.comment.strip():
                comments.append(f"{prefix}{state.comment.strip()}")
        result[task_id] = ("; ".join(reasons) or "-", "; ".join(comments) or "-")
    return result


async def build_meetings_report_sections(db: AsyncSession, report_day: date) -> tuple[date, list[dict[str, str]], dict[str, Any]]:
    tomorrow = next_working_day(report_day)
    week_start = _week_start(report_day)
    common_items = await _common_view_items(tomorrow)
    local_day_start = datetime.combine(report_day, time.min, tzinfo=report_timezone())
    completed_day_start = local_day_start.astimezone(timezone.utc)
    completed_day_end = (local_day_start + timedelta(days=1)).astimezone(timezone.utc)

    task_stmt = (
        select(Task)
        .where(Task.is_active.is_(True))
        .where(
            or_(
                Task.start_date.is_not(None),
                Task.due_date.is_not(None),
                Task.created_at.is_not(None),
            )
        )
    )
    tasks = (await db.execute(task_stmt)).scalars().all()
    # Completed tasks use a separate date-bounded query. This includes tasks
    # archived after completion without widening any other M3 section.
    completed_tasks = (
        await db.execute(
            select(Task).where(
                Task.completed_at >= completed_day_start,
                Task.completed_at < completed_day_end,
            )
        )
    ).scalars().all()
    report_tasks_by_id = {task.id: task for task in [*tasks, *completed_tasks]}
    report_tasks = list(report_tasks_by_id.values())
    names = await _assignee_names(db, report_tasks)
    assignee_ids_by_task = await _effective_task_assignee_ids(db, report_tasks)
    all_participant_ids = await _all_participant_user_ids(db)
    std_tickets_section = await std_tickets_report_section(db, report_day)

    system_tasks = [task for task in tasks if task.system_template_origin_id and _is_open(task)]
    system_late = _dedupe_system_task_rows([task for task in system_tasks if _late_days(task) > 0])
    department_codes = {
        department_id: code
        for department_id, code in (await db.execute(select(Department.id, Department.code))).all()
    }
    user_department_codes = {
        user_id: _m3_department_code_label(department_id, department_codes)
        for user_id, department_id in (await db.execute(select(User.id, User.department_id))).all()
    }
    await apply_weekly_planner_task_order(db, report_tasks, assignee_ids_by_task, department_codes)

    today_todo = [task for task in tasks if _is_without_progress_for_m3_day(task, report_day)]
    # Both pink no-progress tasks and late system tasks show the explanation
    # entered for this report day in Daily Report My View.
    daily_rlz_by_task = await _daily_rlz_values_by_task(
        db, [*today_todo, *system_late], report_day, names, assignee_ids_by_task
    )
    done_today = _completed_tasks_for_report_day(completed_tasks, report_day)
    postponed_today, postponed_date_ranges, postponed_both_today, postponed_both_date_ranges = await _postponed_tasks_for_m3_day(
        db, report_day, tasks
    )
    product_delta_today = _product_delta_tasks_for_m3_day(report_tasks, report_day)

    tomorrow_tasks = [task for task in tasks if _task_day(task) == tomorrow and _is_open(task)]
    new_task_review_tasks = [task for task in tomorrow_tasks if not _is_system_task(task)]
    # "Detyrat e reja" is a one-day list: a task is new on its planned start
    # date, rather than on every day up to its due date.
    new_tomorrow = [task for task in new_task_review_tasks if _is_new_task_for_m3_day(task, tomorrow)]
    at_0800 = [
        task for task in new_task_review_tasks
        if title_has_eight_am_indicator(task.title)
        or (task.due_date and _local_time(task.due_date) == "08:00")
    ]
    deadline = [task for task in new_task_review_tasks if task.is_deadline_important]
    one_h_no_slot = [task for task in tomorrow_tasks if task.is_1h_report and not task.one_h_report_slot]
    personal_ga = [
        task for task in tasks
        if task.is_personal
        and _is_open(task)
        and _task_day(task) == tomorrow
        and (
            PERSONAL_GA.search(_clean_task_title(task.title))
            or PERSONAL_GA.search(task.title or "")
        )
    ]
    blocked = [task for task in tomorrow_tasks if task.is_bllok]
    bz_tasks = [task for task in tomorrow_tasks if re.search(r"\bBZ\b", task.title or "", re.I)]
    bz_alignment_lines = await _bz_alignment_lines(
        db, tomorrow, tasks, names, assignee_ids_by_task, include_status=True
    )
    bz_template_metadata = await _bz_template_metadata(db)

    meeting_stmt = select(Meeting).where(Meeting.starts_at.is_not(None))
    meetings = (await db.execute(meeting_stmt)).scalars().all()
    today_meetings = [meeting for meeting in meetings if _meeting_occurs_on_date(meeting, report_day)]
    tomorrow_meetings = [meeting for meeting in meetings if _meeting_occurs_on_date(meeting, tomorrow)]
    external_meetings = [m for m in tomorrow_meetings if getattr(m, "meeting_type", None) == "external"]
    internal_meetings = [m for m in tomorrow_meetings if getattr(m, "meeting_type", None) != "external"]
    leave_entries = (
        await db.execute(select(CommonEntry).where(CommonEntry.category == CommonCategory.annual_leave))
    ).scalars().all()
    leave_tomorrow = []
    for entry in leave_entries:
        start_date, end_date, full_day, start_time, end_time, note, is_all_users = parse_common_view_annual_leave(entry)
        if start_date <= tomorrow <= end_date:
            leave_tomorrow.append(
                (entry, start_date, end_date, full_day, start_time, end_time, note, is_all_users)
            )
    leave_user_sort_keys = await weekly_planner_user_sort_keys(
        db,
        {
            entry.assigned_to_user_id or entry.created_by_user_id
            for entry, *_ in leave_tomorrow
            if entry.assigned_to_user_id or entry.created_by_user_id
        },
        department_codes,
    )

    table_kwargs = {
        "assignee_ids_by_task": assignee_ids_by_task,
        "all_participant_ids": all_participant_ids,
    }
    section_1 = _m3_status_table(
        "LATE",
        system_late,
        names,
        include_late_days=True,
        include_department=True,
        include_am_pm=True,
        department_codes=department_codes,
        daily_rlz_by_task=daily_rlz_by_task,
        **table_kwargs,
    )
    section_4 = _tomorrow_common_section(
        common_items=common_items,
        tomorrow=tomorrow,
        fallback_external=_meeting_lines(external_meetings),
        fallback_internal=_meeting_lines(internal_meetings),
        fallback_bz=bz_alignment_lines
        or _task_lines(
            bz_tasks, names, assignee_ids_by_task, include_status=True, all_participant_ids=all_participant_ids
        ),
        fallback_blocked=_task_lines(
            blocked, names, assignee_ids_by_task, include_status=True, all_participant_ids=all_participant_ids
        ),
        bz_task_metadata=_merged_task_metadata(
            _merged_task_metadata(
                _task_metadata_by_title(
                    tomorrow_tasks,
                    department_codes,
                    names,
                    assignee_ids_by_task,
                    all_participant_ids,
                ),
                bz_template_metadata,
            ),
            _common_task_metadata_by_title(
                common_items.get("bz") or [], tomorrow, department_codes, user_department_codes
            ),
        ),
        blocked_task_metadata=_merged_task_metadata(
            _task_metadata_by_title(
                blocked,
                department_codes,
                names,
                assignee_ids_by_task,
                all_participant_ids,
            ),
            _common_task_metadata_by_title(
                common_items.get("blocked") or [], tomorrow, department_codes, user_department_codes
            ),
        ),
        with_status=True,
    )
    section_5 = [
        *_m3_status_table(
            "DETYRAT E REJA",
            new_tomorrow,
            names,
            with_status=True,
            include_department=True,
            include_added_week=True,
            include_am_pm=True,
            department_codes=department_codes,
            week_start=week_start,
            **table_kwargs,
        ),
        "",
        *_m3_status_table(
            "DET 08:00",
            at_0800,
            names,
            include_department=True,
            include_am_pm=True,
            department_codes=department_codes,
            **table_kwargs,
        ),
        "",
        *_m3_status_table(
            "DET ME DEADLINE",
            deadline,
            names,
            include_department=True,
            include_am_pm=True,
            department_codes=department_codes,
            **table_kwargs,
        ),
    ]
    section_6 = await _today_meeting_status_section(db, today_meetings, report_day)

    by_title = {
        SECTION_TITLES[0]: "(Ploteso manualisht)",
        SECTION_TITLES[1]: std_tickets_section,
        SECTION_TITLES[2]: _normalize_section(section_1),
        SECTION_TITLES[3]: _normalize_section(
            _m3_status_table(
                "TODO",
                today_todo,
                names,
                include_department=True,
                include_am_pm=True,
                department_codes=department_codes,
                daily_rlz_by_task=daily_rlz_by_task,
                **table_kwargs,
            )
        ),
        SECTION_TITLES[11]: _normalize_section(
            _m3_status_table(
                "SHTYER START DHE DUE DATE",
                postponed_both_today,
                names,
                with_status=True,
                include_type=True,
                include_department=True,
                include_am_pm=True,
                department_codes=department_codes,
                date_range_by_task=postponed_both_date_ranges,
                **table_kwargs,
            )
            + ["", "SHTYER DUE DATE:"]
            + _m3_status_table(
                "",
                postponed_today,
                names,
                with_status=True,
                include_type=True,
                include_department=True,
                include_am_pm=True,
                department_codes=department_codes,
                date_range_by_task=postponed_date_ranges,
                **table_kwargs,
            )[1:]
        ),
        SECTION_TITLES[10]: _normalize_section(
            _m3_status_table(
                "DET E KRYERA SOT (AM/PM)",
                done_today,
                names,
                with_status=True,
                include_type=True,
                include_department=True,
                include_am_pm=True,
                department_codes=department_codes,
                **table_kwargs,
            )
        ),
        SECTION_TITLES[12]: _normalize_section(
            _m3_status_table(
                "PRODUKTE +/-",
                product_delta_today,
                names,
                with_status=True,
                include_department=True,
                include_am_pm=True,
                department_codes=department_codes,
                products_by_task={
                    task.id: _m3_product_delta_label(task) for task in product_delta_today
                },
                **table_kwargs,
            )
        ),
        SECTION_TITLES[7]: section_6,
        SECTION_TITLES[4]: _empty_aware(
            _leave_lines(leave_tomorrow, names, user_department_codes, leave_user_sort_keys)
        ),
        SECTION_TITLES[6]: _normalize_section(section_5),
        SECTION_TITLES[5]: _normalize_section(section_4),
        SECTION_TITLES[8]: _normalize_section(
            _m3_status_table(
                "1H PA SLOT",
                one_h_no_slot,
                names,
                with_status=True,
                include_department=True,
                include_am_pm=True,
                department_codes=department_codes,
                **table_kwargs,
            )
        ),
        SECTION_TITLES[9]: _normalize_section(
            _m3_status_table(
                "PERSONAL GA",
                personal_ga,
                names,
                with_status=True,
                include_department=True,
                include_am_pm=True,
                department_codes=department_codes,
                **table_kwargs,
            )
        ),
    }
    sections = [{"title": title, "body": by_title[title]} for title in DISPLAY_SECTION_TITLES]
    snapshot = {
        "report_day": report_day.isoformat(),
        "tomorrow": tomorrow.isoformat(),
        "counts": {section["title"]: section["body"].count("\n- ") + (1 if section["body"].startswith("- ") else 0) for section in sections},
    }
    return tomorrow, sections, snapshot


def _normalize_report_status(value: str | None) -> str:
    status = str(value or "TODO").strip().upper().replace(" ", "_")
    if status in {"COMPLETED", "COMPLETE"}:
        return "DONE"
    if status in {"TO_DO", "TO-DO"}:
        return "TODO"
    if status in {"INPROGRESS", "IN-PROGRESS"}:
        return "IN_PROGRESS"
    if status in {"WAITINGCLIENT", "WAITING_CLIENT", "WAITING_FOR_CLIENT"}:
        return "WAITING_CLIENT"
    if status in {"WAITING", "WAITING_CONFIRMATION", "PENDING_CONFIRMATION"}:
        return "WAITING_CONFIRMATION"
    if status in {"TODO", "IN_PROGRESS", "WAITING_CLIENT", "WAITING_CONFIRMATION", "DONE"}:
        return status
    return "TODO"


def _task_lines(
    tasks: list[Task],
    names: dict[Any, str],
    assignee_ids_by_task: dict[Any, set[Any]] | None = None,
    *,
    include_status: bool = False,
    all_participant_ids: set[Any] | None = None,
) -> list[str]:
    if not tasks:
        return ["(Asnje detyre)"]
    ordered = sorted(
        tasks,
        key=lambda task: common_view_task_sort_key(
            task, names, assignee_ids_by_task, all_participant_ids=all_participant_ids
        ),
    )
    return [
        _task_line(
            task,
            names,
            assignee_ids_by_task,
            include_status=include_status,
            all_participant_ids=all_participant_ids,
        )
        for task in ordered
    ]


def _task_metadata_by_title(
    tasks: list[Task],
    department_codes: dict[Any, str] | None,
    names: dict[Any, str] | None = None,
    assignee_ids_by_task: dict[Any, set[Any]] | None = None,
    all_participant_ids: set[Any] | None = None,
) -> dict[str, tuple[str, str]]:
    """Metadata for line-based BZ/BLLOK tables.

    A title can legitimately occur for different users.  Store an owner+title
    key for the table renderer, while retaining a title-only fallback for old
    callers and non-owned sources.
    """
    metadata: dict[str, tuple[str, str]] = {}
    for task in tasks:
        title = _clean_task_title(task.title)
        value = (_m3_department_label(task, department_codes), _m3_am_pm_label(task))
        metadata.setdefault(title, value)
        if names is not None:
            owner = _task_owners(
                task, names, assignee_ids_by_task, all_participant_ids=all_participant_ids
            )
            metadata[f"{owner}\0{title}"] = value
    return metadata


def _common_task_metadata_by_title(
    items: list[dict[str, Any]],
    day: date,
    department_codes: dict[Any, str] | None,
    user_department_codes: dict[Any, str] | None = None,
) -> dict[str, tuple[str, str]]:
    """Read the responsible user's department and AM/PM from Common View rows.

    The BZ and Block tables may use Common View rows whose effective date does
    not match the report's fallback task list.  Those rows already include the
    resolved assignee department and finish period, so use them instead of
    rendering empty metadata.
    """
    metadata: dict[str, tuple[str, str]] = {}
    for item in items:
        if _item_date(item) != day:
            continue
        title = _common_title(item)
        if not title:
            continue
        user_id = str(item.get("user_id") or item.get("userId") or "")
        department = next(
            (code for candidate_id, code in (user_department_codes or {}).items() if str(candidate_id) == user_id),
            "",
        )
        if not department:
            department = _m3_department_code_label(
                item.get("department_id") or item.get("departmentId"), department_codes
            )
        period = str(item.get("finish_period") or item.get("finishPeriod") or "").strip().upper()
        am_pm = period if period in {"AM", "PM", "AM/PM"} else "-"
        metadata[title] = (department, am_pm)
    return metadata


def _merged_task_metadata(
    fallback: dict[str, tuple[str, str]], common: dict[str, tuple[str, str]]
) -> dict[str, tuple[str, str]]:
    """Prefer Common View's resolved fields while retaining fallback values."""
    merged = dict(fallback)
    for title, (common_department, common_am_pm) in common.items():
        fallback_department, fallback_am_pm = merged.get(title, ("-", "-"))
        merged[title] = (
            common_department if common_department != "-" else fallback_department,
            common_am_pm if common_am_pm != "-" else fallback_am_pm,
        )
    return merged


def _status_group_section(
    title: str | None,
    tasks: list[Task],
    names: dict[Any, str],
    report_day: date,
    *,
    assignee_ids_by_task: dict[Any, set[Any]] | None = None,
    all_participant_ids: set[Any] | None = None,
) -> list[str]:
    today_tasks = [task for task in tasks if _task_day(task) == report_day]
    todo = [task for task in today_tasks if str(task.status or "").upper() == "TODO" and _is_open(task)]
    in_progress = [task for task in today_tasks if str(task.status or "").upper() == "IN_PROGRESS" and _is_open(task)]
    done = [
        task for task in tasks
        if str(task.status or "").upper() in {"DONE", "COMPLETED"}
        and (_task_day(task) == report_day or _local_date(task.completed_at) == report_day)
    ]
    late = [task for task in tasks if _is_open(task) and _late_days(task) > 0]
    table_kwargs = {
        "assignee_ids_by_task": assignee_ids_by_task,
        "all_participant_ids": all_participant_ids,
        "include_type": True,
    }
    return [
        *([f"{title}:"] if title else []),
        *_m3_status_table("TODO", todo, names, **table_kwargs),
        "",
        *_m3_status_table("IN PROGRESS", in_progress, names, **table_kwargs),
        "",
        *_m3_status_table("DONE", done, names, **table_kwargs),
        "",
        *_m3_status_table("LATE", late, names, include_late_days=True, **table_kwargs),
    ]


async def _m3_finance_ga_sections(
    db: AsyncSession, tasks: list[Task], names: dict[Any, str], report_day: date
) -> tuple[str, str]:
    assignee_ids_by_task = await _effective_task_assignee_ids(db, tasks)
    all_participant_ids = await _all_participant_user_ids(db)
    ga_users = await _users_by_initials(db, "GA")
    ga_user_ids = {user.id for user in ga_users}
    ga_tasks = [
        task for task in tasks
        if ga_user_ids.intersection(assignee_ids_by_task.get(task.id, set()))
        and (
            _task_day(task) == report_day
            or _local_date(task.completed_at) == report_day
            or _late_days(task) > 0
        )
    ]

    finance_department = (
        await db.execute(select(Department).where(Department.name.ilike("finance")))
    ).scalar_one_or_none()
    if finance_department is None:
        hv_tasks: list[Task] = []
    else:
        finance_tasks = [task for task in tasks if task.department_id == finance_department.id]
        hv_tasks = [
            task for task in finance_tasks
            if _task_day(task) == report_day
            or _local_date(task.completed_at) == report_day
            or _late_days(task) > 0
        ]

    return (
        _normalize_section(_status_group_section(
            None,
            ga_tasks,
            names,
            report_day,
            assignee_ids_by_task=assignee_ids_by_task,
            all_participant_ids=all_participant_ids,
        )),
        _normalize_section(_status_group_section(
            None,
            hv_tasks,
            names,
            report_day,
            assignee_ids_by_task=assignee_ids_by_task,
            all_participant_ids=all_participant_ids,
        )),
    )


def _system_row_key(task: Task) -> tuple[str, str, str, str, str]:
    return (
        str(task.system_template_origin_id or ""),
        task.title or "",
        (_local_date(task.start_date) or "").isoformat() if _local_date(task.start_date) else "",
        (_local_date(task.due_date) or "").isoformat() if _local_date(task.due_date) else "",
        (_local_date(getattr(task, "planned_date", None) or task.original_due_date) or "").isoformat()
        if _local_date(getattr(task, "planned_date", None) or task.original_due_date)
        else "",
    )


def _dedupe_system_task_rows(tasks: list[Task]) -> list[Task]:
    by_key: dict[tuple[str, str, str, str, str], Task] = {}
    for task in tasks:
        key = _system_row_key(task)
        existing = by_key.get(key)
        if existing is None or _late_days(task) > _late_days(existing):
            by_key[key] = task
    return sorted(
        by_key.values(),
        key=lambda task: (
            -_late_days(task),
            *common_view_task_sort_key(task, {}, None),
        ),
    )


def _meeting_clock_sort_key(meeting: Meeting) -> tuple[int, int, str]:
    """Order report meetings by their displayed clock time, never their source date."""
    starts_at = meeting.starts_at
    if starts_at is None:
        return (24, 0, _clean_task_title(meeting.title).casefold())
    local = starts_at.astimezone(report_timezone()) if starts_at.tzinfo else starts_at
    return (local.hour, local.minute, _clean_task_title(meeting.title).casefold())


def _meeting_lines(meetings: list[Meeting]) -> list[str]:
    if not meetings:
        return ["(Asnje takim)"]

    return [
        f"- {_local_time(meeting.starts_at)}: {_meeting_title_with_highlight(meeting)}"
        for meeting in sorted(meetings, key=_meeting_clock_sort_key)
    ]


def _meeting_title_with_highlight(meeting: Meeting) -> str:
    title = _clean_task_title(meeting.title)
    recurrence = str(getattr(meeting, "recurrence_type", None) or "").strip().lower()
    if recurrence not in {"daily", "weekly"}:
        return f"{title} {MEETING_HIGHLIGHT_MARKER}"
    return title


def _split_meeting_highlight_marker(value: str) -> tuple[str, bool]:
    highlighted = bool(MEETING_HIGHLIGHT_PATTERN.search(value or ""))
    return MEETING_HIGHLIGHT_PATTERN.sub("", value or "").strip(), highlighted


def _append_meeting_highlight_marker(lines: list[str], highlighted: bool, width: int) -> list[str]:
    padded = [f"{line:<{width}}" for line in lines] or [f"{'-':<{width}}"]
    if highlighted:
        padded[-1] = f"{padded[-1].rstrip()} {MEETING_HIGHLIGHT_MARKER}"
    return padded


def _meeting_group_title(title: str) -> list[str]:
    return [title]


def _meeting_status_checkbox_table(meetings: list[Meeting], status_by_meeting: dict[Any, str]) -> list[str]:
    border = "+----+-------+----------+------------------------------------------------------------------+"
    rows = [
        border,
        f"| {'NR':<2} | {'KOHA':<5} | {'MBAJTUR?':<8} | {'TITULLI':<64} |",
        border,
    ]
    if not meetings:
        rows.append(f"| {'-':<2} | {'-':<5} | {'':<8} | {'(Asnje takim)':<64} |")
        rows.append(border)
        return rows
    for index, meeting in enumerate(sorted(meetings, key=_meeting_clock_sort_key), start=1):
        status = status_by_meeting.get(meeting.id, "")
        title_value, highlighted = _split_meeting_highlight_marker(_meeting_title_with_highlight(meeting))
        title_lines = _append_meeting_highlight_marker(_wrap_fixed_width(title_value, 64), highlighted, 64)
        status_icon = "\u2713" if status == "held" else "\u2715" if status == "canceled" else ""
        rows.append(
            f"| {index:<2} | {_local_time(meeting.starts_at):<5} | {status_icon:<8} | {title_lines[0]} |"
        )
        for line in title_lines[1:]:
            rows.append(f"| {'':<2} | {'':<5} | {'':<8} | {line} |")
        rows.append(border)
    return rows


async def _today_meeting_status_section(db: AsyncSession, meetings: list[Meeting], report_day: date) -> str:
    statuses = (
        await db.execute(
            select(MeetingOccurrenceStatus).where(MeetingOccurrenceStatus.occurrence_date == report_day)
        )
    ).scalars().all()
    status_by_meeting = {row.meeting_id: row.status for row in statuses}
    external = [meeting for meeting in meetings if getattr(meeting, "meeting_type", None) == "external"]
    internal = [meeting for meeting in meetings if getattr(meeting, "meeting_type", None) != "external"]

    return _normalize_section([
        *_meeting_group_title("TAK EXTERNE"),
        *_meeting_status_checkbox_table(external, status_by_meeting),
        "",
        *_meeting_group_title("TAK INTERNE"),
        *_meeting_status_checkbox_table(internal, status_by_meeting),
    ])


async def _common_view_items(day: date) -> dict[str, list[dict[str, Any]]]:
    base_url = os.getenv("PRIMEFLOW_API_BASE_URL")
    if not base_url:
        return {}
    client = PrimeFlowClient(
        base_url.rstrip("/"),
        os.getenv("PRIMEFLOW_EMAIL"),
        os.getenv("PRIMEFLOW_PASSWORD"),
        os.getenv("PRIMEFLOW_ACCESS_TOKEN"),
    )
    try:
        payload = await client.common_view(day)
    except Exception:
        return {}
    return payload.get("items") or {}


def _item_date(item: dict[str, Any]) -> date | None:
    for key in ("date", "day", "report_date", "startDate", "planned_for", "due_date"):
        raw = item.get(key)
        if raw:
            try:
                return date.fromisoformat(str(raw)[:10])
            except ValueError:
                continue
    return None


def _common_title(item: dict[str, Any]) -> str:
    raw = item.get("task_title") or item.get("title") or item.get("task") or item.get("note") or ""
    return _clean_task_title(str(raw))


def _common_owner(item: dict[str, Any]) -> str:
    person = str(item.get("person") or item.get("owner") or item.get("employee") or item.get("assignee_name") or "").strip()
    if person.upper() == "ALL":
        return "ALL"
    assignees = item.get("assignees") or item.get("assigned_users") or item.get("owners")
    if isinstance(assignees, list):
        initials = []
        for assignee in assignees:
            if isinstance(assignee, dict):
                label = assignee.get("full_name") or assignee.get("username") or assignee.get("email") or assignee.get("name")
            else:
                label = assignee
            value = _initials(str(label or ""))
            if value != "-" and value not in initials:
                initials.append(value)
        if initials:
            return " ".join(initials)
    bz_with_label = str(item.get("bzWithLabel") or item.get("bz_with_label") or "").strip()
    if bz_with_label:
        initials = [
            value.strip().upper()
            for value in re.split(r"[,;/\s]+", bz_with_label)
            if value.strip()
        ]
        if initials:
            return " ".join(dict.fromkeys(initials))
    return _initials(person)


def _common_task_status(item: dict[str, Any]) -> str:
    raw = item.get("status") or item.get("task_status") or item.get("state") or ""
    return _normalize_report_status(str(raw) if raw else "TODO")


def _common_task_lines(items: list[dict[str, Any]], day: date, *, include_status: bool = False) -> list[str]:
    rows: list[dict[str, Any]] = []
    seen = set()
    for item in items:
        if _item_date(item) != day:
            continue
        title = _common_title(item)
        owner = _common_owner(item)
        key = (item.get("id"), title, owner)
        if key in seen:
            continue
        seen.add(key)
        rows.append(item)
    rows.sort(key=common_view_item_sort_key)
    lines = []
    for item in rows:
        title = _common_title(item)
        owner = _common_owner(item)
        if include_status:
            lines.append(f"- [{_common_task_status(item)}] {owner}: {title}")
        else:
            lines.append(f"- {owner}: {title}")
    return lines


async def _bz_alignment_lines(
    db: AsyncSession,
    day: date,
    tasks: list[Task],
    names: dict[Any, str],
    assignee_ids_by_task: dict[Any, set[Any]],
    *,
    include_status: bool = False,
) -> list[str]:
    templates = (
        await db.execute(
            select(SystemTaskTemplate)
            .where(SystemTaskTemplate.is_active.is_(True))
            .where(SystemTaskTemplate.approval_status == CommonApprovalStatus.approved)
        )
    ).scalars().all()
    template_ids = [template.id for template in templates]
    if not template_ids:
        return []

    rows = (
        await db.execute(
            select(SystemTaskTemplateAlignmentUser.template_id, SystemTaskTemplateAlignmentUser.user_id)
            .where(SystemTaskTemplateAlignmentUser.template_id.in_(template_ids))
        )
    ).all()
    alignment_users_map: dict[Any, list[Any]] = {}
    user_ids: set[Any] = set()
    for template_id, user_id in rows:
        alignment_users_map.setdefault(template_id, []).append(user_id)
        user_ids.add(user_id)
    if not user_ids:
        return []

    for template in templates:
        for user_id in (template.assignee_ids or []):
            user_ids.add(user_id)
        if template.default_assignee_id:
            user_ids.add(template.default_assignee_id)

    users = (await db.execute(select(User).where(User.id.in_(user_ids)))).scalars().all()
    users_map = {user.id: user for user in users}
    department_codes = {
        department_id: code
        for department_id, code in (await db.execute(select(Department.id, Department.code))).all()
    }
    ga_user = next((user for user in users if (user.username or "").lower() == "gane.arifaj"), None)
    ga_user_id = ga_user.id if ga_user else None
    if ga_user_id is None:
        ga_candidates = [user for user in users if _initials(user.full_name or user.username or user.email) == "GA"]
        ga_user_id = ga_candidates[0].id if ga_candidates else None

    task_owners_by_template: dict[Any, list[str]] = {}
    task_status_by_template: dict[Any, str] = {}
    status_rank = {"IN_PROGRESS": 4, "WAITING_CLIENT": 3, "WAITING_CONFIRMATION": 2, "TODO": 1, "DONE": 0}
    for task in tasks:
        template_id = task.system_template_origin_id
        if not template_id:
            continue
        if _local_date(task.start_date or task.due_date or task.created_at) != day:
            continue
        owners = task_owners_by_template.setdefault(template_id, [])
        for owner in _task_owners(task, names, assignee_ids_by_task).split():
            if owner != "-" and owner not in owners:
                owners.append(owner)
        status = _normalize_report_status(task.status)
        current = task_status_by_template.get(template_id)
        if current is None or status_rank.get(status, 0) > status_rank.get(current, 0):
            task_status_by_template[template_id] = status

    def template_planner_sort_key(template: SystemTaskTemplate) -> tuple:
        assignee_ids = list(template.assignee_ids or [])
        # BZ rows use the first configured assignee as their one report
        # reference user; the template default is used only when no list exists.
        primary_user_id = assignee_ids[0] if assignee_ids else template.default_assignee_id
        primary_user = users_map.get(primary_user_id)
        department = _m3_department_code_label(
            getattr(primary_user, "department_id", None), department_codes
        ) if primary_user else "-"
        return (
            WEEKLY_PLANNER_DEPARTMENT_ORDER.get(department, len(WEEKLY_PLANNER_DEPARTMENT_ORDER)),
            department.casefold(),
            1 if not primary_user or primary_user.weekly_planner_sort_order is None else 0,
            primary_user.weekly_planner_sort_order if primary_user and primary_user.weekly_planner_sort_order is not None else 10**9,
            (primary_user.full_name or primary_user.username or primary_user.email or "").casefold() if primary_user else "~",
            template.alignment_time or datetime.max.time(),
            (template.title or "").casefold(),
        )

    lines: list[str] = []
    seen: set[tuple[Any, str]] = set()
    for template in sorted(templates, key=template_planner_sort_key):
        alignment_ids = alignment_users_map.get(template.id, [])
        if not alignment_ids or (ga_user_id is not None and ga_user_id not in alignment_ids):
            continue
        if not matches_template_date(template, day):
            continue
        owners = list(task_owners_by_template.get(template.id) or [])
        if not owners:
            assignee_ids = list(template.assignee_ids or [])
            if not assignee_ids and template.default_assignee_id:
                assignee_ids = [template.default_assignee_id]
            owners = [
                _initials(users_map[user_id].full_name or users_map[user_id].username or users_map[user_id].email)
                for user_id in assignee_ids
                if user_id in users_map
            ]
        owner_label = " ".join(dict.fromkeys([owner for owner in owners if owner != "-"])) or "-"
        title = _clean_task_title(template.title)
        key = (template.id, day.isoformat())
        if key in seen:
            continue
        seen.add(key)
        status = task_status_by_template.get(template.id, "TODO")
        if include_status:
            lines.append(f"- [{status}] {owner_label}: {title}")
        else:
            lines.append(f"- {owner_label}: {title}")
    return lines


async def _bz_template_metadata(db: AsyncSession) -> dict[str, tuple[str, str]]:
    """Resolve BZ metadata from each template's first assigned user.

    This keeps BZ's displayed department and sorting reference consistent for
    templates that have multiple assignees.
    """
    templates = (
        await db.execute(
            select(SystemTaskTemplate)
            .where(SystemTaskTemplate.is_active.is_(True))
            .where(SystemTaskTemplate.approval_status == CommonApprovalStatus.approved)
        )
    ).scalars().all()
    user_ids = {
        (list(template.assignee_ids or [])[0] if list(template.assignee_ids or []) else template.default_assignee_id)
        for template in templates
        if list(template.assignee_ids or []) or template.default_assignee_id
    }
    users = (
        await db.execute(select(User).where(User.id.in_(user_ids)))
    ).scalars().all() if user_ids else []
    users_by_id = {user.id: user for user in users}
    department_codes = {
        department_id: code
        for department_id, code in (await db.execute(select(Department.id, Department.code))).all()
    }
    metadata: dict[str, tuple[str, str]] = {}
    for template in templates:
        assignee_ids = list(template.assignee_ids or [])
        primary_user_id = assignee_ids[0] if assignee_ids else template.default_assignee_id
        primary_user = users_by_id.get(primary_user_id)
        metadata[_clean_task_title(template.title)] = (
            _m3_department_code_label(getattr(primary_user, "department_id", None), department_codes)
            if primary_user else "-",
            "-",
        )
    return metadata


def _common_meeting_lines(items: list[dict[str, Any]], day: date) -> list[str]:
    rows: list[tuple[tuple[int, int, str], str]] = []
    seen = set()
    for item in items:
        if _item_date(item) != day:
            continue
        title = _clean_task_title(str(item.get("title") or item.get("task_title") or "Meeting"))
        recurrence = str(item.get("recurrence_type") or item.get("recurrenceType") or "").strip().lower()
        if recurrence not in {"daily", "weekly"}:
            title = f"{title} {MEETING_HIGHLIGHT_MARKER}"
        time_value = str(item.get("time") or item.get("when") or "").strip()
        prefix = f"{time_value}: " if time_value else ""
        key = (item.get("id"), title, time_value)
        if key in seen:
            continue
        seen.add(key)
        time_match = re.search(r"\b([01]?\d|2[0-3]):([0-5]\d)\b", time_value)
        time_sort_key = (
            int(time_match.group(1)) if time_match else 24,
            int(time_match.group(2)) if time_match else 0,
            title.casefold(),
        )
        rows.append((time_sort_key, f"- {prefix}{title}"))
    return [line for _, line in sorted(rows, key=lambda row: row[0])]


def _prefer_common(common_lines: list[str], fallback_lines: list[str]) -> list[str]:
    return common_lines if common_lines else [line for line in fallback_lines if not line.startswith("(Asnje")]


def _prefer_owned_common(common_lines: list[str], fallback_lines: list[str]) -> list[str]:
    usable_common = [
        line for line in common_lines
        if not re.match(r"^\s*-\s*-+\s*:", line)
    ]
    return usable_common if usable_common else [line for line in fallback_lines if not line.startswith("(Asnje")]


def _strip_list_marker(value: str) -> str:
    return re.sub(r"^\s*-\s*", "", value).strip()


def _tomorrow_meeting_table(title: str, lines: list[str]) -> list[str]:
    border = "+----+-------+------------------------------------------------------------------+"
    rows = [
        f"{title}:",
        border,
        f"| {'NR':<2} | {'KOHA':<5} | {'TITULLI':<64} |",
        border,
    ]
    values = [_strip_list_marker(line) for line in lines if line and not line.startswith("(")]
    if not values:
        rows.append(f"| {'-':<2} | {'-':<5} | {'(Asnje takim)':<64} |")
        rows.append(border)
        return rows
    for index, value in enumerate(values, start=1):
        match = re.match(r"^(\d{1,2}:\d{2})\s*:\s*(.+)$", value)
        time_value = match.group(1) if match else "-"
        title_value = match.group(2) if match else value
        title_value, highlighted = _split_meeting_highlight_marker(title_value)
        title_lines = _append_meeting_highlight_marker(_wrap_fixed_width(title_value, 64), highlighted, 64)
        rows.append(f"| {index:<2} | {time_value:<5} | {title_lines[0]} |")
        for line in title_lines[1:]:
            rows.append(f"| {'':<2} | {'':<5} | {line} |")
        rows.append(border)
    return rows


def _parse_owned_task_line(value: str) -> tuple[str, str, str]:
    """Return (status, owner, title) from ``[STATUS] WHO: title`` or ``WHO: title``."""
    status = ""
    remainder = value
    status_match = TASK_LINE_STATUS.match(remainder)
    if status_match:
        status = _normalize_report_status(status_match.group(1))
        remainder = remainder[status_match.end() :]
    owner = "-"
    title_value = remainder
    if ":" in remainder:
        owner, title_value = remainder.split(":", 1)
        owner = owner.strip() or "-"
        title_value = title_value.strip()
    return status, owner, title_value


def _title_with_status_marker(title: str, status: str | None) -> str:
    cleaned = _strip_status_markers(title or "").rstrip()
    if not status:
        return cleaned
    return f"{cleaned} [[st:{_normalize_report_status(status)}]]"


def _split_status_marker(value: str) -> tuple[str, str]:
    """Return (title, status). Marker may sit mid-cell after wrap/merge, so search anywhere."""
    text = value or ""
    matches = list(re.finditer(r"\s*\[\[\s*st\s*:?\s*([A-Z_]+)\s*\]\]", text, flags=re.I))
    if not matches:
        return text.rstrip(), ""
    status = _normalize_report_status(matches[-1].group(1))
    cleaned = text
    for match in reversed(matches):
        cleaned = cleaned[: match.start()] + cleaned[match.end() :]
    return re.sub(r"\s+", " ", cleaned).strip(), status


def _strip_status_markers(value: str) -> str:
    return re.sub(r"\s*\[\[\s*st\s*:?\s*[A-Z_]+\s*\]\]", "", value or "", flags=re.I)


def _append_status_marker_to_lines(title_lines: list[str], status: str | None, width: int) -> list[str]:
    """Keep [[st:STATUS]] on the last wrapped line so merge/display still finds it."""
    lines = list(title_lines) or ["-"]
    if not status:
        return [f"{lines[0]:<{width}}", *[f"{line:<{width}}" for line in lines[1:]]]
    marker = f" [[st:{_normalize_report_status(status)}]]"
    padded = [f"{line:<{width}}" for line in lines]
    padded[-1] = f"{padded[-1].rstrip()}{marker}"
    return padded


def _tomorrow_task_table(
    title: str,
    lines: list[str],
    *,
    with_status: bool = False,
    task_metadata: dict[str, tuple[str, str]] | None = None,
    include_am_pm_times: bool = False,
) -> list[str]:
    who_width = 20
    department_width = 5
    am_pm_width = 11 if include_am_pm_times else 5
    title_width = 64
    border = (
        f"+----+{'-' * (who_width + 2)}+{'-' * (department_width + 2)}+"
        f"{'-' * (am_pm_width + 2)}+{'-' * (title_width + 2)}+"
    )
    rows = [
        f"{title}:",
        border,
        f"| {'NR':<2} | {'KUSH':<{who_width}} | {'DEP':<{department_width}} | {'AM/PM':<{am_pm_width}} | {'TITULLI':<{title_width}} |",
        border,
    ]
    values = [_strip_list_marker(line) for line in lines if line and not line.startswith("(")]

    def metadata_for(value: str) -> tuple[str, str]:
        """Resolve metadata for the exact rendered owner/title pair.

        Different tasks can share a title (for example one task for DV and one
        for LH).  Looking up only the title lets one person's department leak
        into the other's row.
        """
        _, owner, raw_title = _parse_owned_task_line(value)
        owner_display = re.sub(r"\s+", " ", owner).strip() or "-"
        title_value = _strip_status_markers(raw_title).strip()
        metadata = task_metadata or {}
        return metadata.get(f"{owner_display}\0{title_value}", metadata.get(title_value, ("-", "-")))

    department_order = M3_DEPARTMENT_ORDER
    values.sort(
        key=lambda value: (
            1 if _normalize_report_status(_parse_owned_task_line(value)[0]) == "DONE" else 0,
            department_order.get(metadata_for(value)[0], len(department_order)),
        )
    )
    if not values:
        rows.append(
            f"| {'-':<2} | {'-':<{who_width}} | {'-':<{department_width}} | {'-':<{am_pm_width}} | {'(Asnje detyre)':<{title_width}} |"
        )
        rows.append(border)
        return rows
    for index, value in enumerate(values, start=1):
        status, owner, title_value = _parse_owned_task_line(value)
        if with_status:
            title_value = _strip_status_markers(title_value)
            status = status or "TODO"
        # Keep all assignees in one WHO cell so one task is never split by a mid-row border.
        owner_display = re.sub(r"\s+", " ", owner).strip() or "-"
        department, am_pm = metadata_for(value)
        if include_am_pm_times:
            am_pm = {"AM": "AM (08:15)", "PM": "PM (13:30)"}.get(am_pm, am_pm)
        title_lines = _wrap_fixed_width(title_value, title_width)
        padded_titles = _append_status_marker_to_lines(
            title_lines, status if with_status else None, title_width
        )
        rows.append(
            f"| {index:<2} | {owner_display:<{who_width}} | {department:<{department_width}} | {am_pm:<{am_pm_width}} | {padded_titles[0]} |"
        )
        for line in padded_titles[1:]:
            rows.append(
                f"| {'':<2} | {'':<{who_width}} | {'':<{department_width}} | {'':<{am_pm_width}} | {line} |"
            )
        rows.append(border)
    return rows


def _tomorrow_common_section(
    *,
    common_items: dict[str, list[dict[str, Any]]],
    tomorrow: date,
    fallback_external: list[str],
    fallback_internal: list[str],
    fallback_bz: list[str],
    fallback_blocked: list[str],
    bz_task_metadata: dict[str, tuple[str, str]] | None = None,
    blocked_task_metadata: dict[str, tuple[str, str]] | None = None,
    with_status: bool = False,
) -> list[str]:
    external = _prefer_common(_common_meeting_lines(common_items.get("external") or [], tomorrow), fallback_external)
    internal = _prefer_common(_common_meeting_lines(common_items.get("internal") or [], tomorrow), fallback_internal)
    bz = _prefer_owned_common(
        _common_task_lines(common_items.get("bz") or [], tomorrow, include_status=with_status),
        fallback_bz,
    )
    blocked = _prefer_common(
        _common_task_lines(common_items.get("blocked") or [], tomorrow, include_status=with_status),
        fallback_blocked,
    )
    return [
        *_tomorrow_meeting_table("TAK EXTERNE", external),
        "",
        *_tomorrow_meeting_table("TAK INTERNE", internal),
        "",
        *_tomorrow_task_table(
            "BZ ME GA",
            bz,
            with_status=with_status,
            task_metadata=bz_task_metadata,
            include_am_pm_times=True,
        ),
        "",
        *_tomorrow_task_table("Bllok 14:30 - 16:00", blocked, with_status=with_status, task_metadata=blocked_task_metadata),
    ]


def _leave_lines(
    entries: list[
        tuple[CommonEntry, date, date, bool, str | None, str | None, str | None, bool]
    ],
    names: dict[Any, str],
    user_department_codes: dict[Any, str] | None = None,
    user_sort_keys: dict[Any, tuple[int, str, int, int, str]] | None = None,
) -> list[str]:
    who_width = 5
    department_width = 5
    date_width = 14
    border = f"+----+{'-' * (who_width + 2)}+{'-' * (department_width + 2)}+{'-' * (date_width + 2)}+{'-' * (date_width + 2)}+"
    lines = [
        border,
        f"| {'NR':<2} | {'KUSH':<{who_width}} | {'DEP':<{department_width}} | {'NGA':<{date_width}} | {'DERI':<{date_width}} |",
        border,
    ]
    if not entries:
        lines.append(f"| {'-':<2} | {'-':<{who_width}} | {'-':<{department_width}} | {'(Asnje detyre)':<{date_width}} | {'-':<{date_width}} |")
        lines.append(border)
        return lines
    ordered_entries = sorted(
        entries,
        key=lambda row: (
            (10**6, "~", 1, 10**6, "~")
            if row[7]
            else (user_sort_keys or {}).get(
                row[0].assigned_to_user_id or row[0].created_by_user_id,
                (10**6, "~", 1, 10**6, _initials(names.get(row[0].assigned_to_user_id or row[0].created_by_user_id)).casefold()),
            ),
            row[1],
            row[2],
            str(getattr(row[0], "id", "")),
        ),
    )
    for index, (entry, start_date, end_date, full_day, start_time, end_time, note, is_all_users) in enumerate(ordered_entries, start=1):
        person = "ALL" if is_all_users else _initials(names.get(entry.assigned_to_user_id or entry.created_by_user_id) or entry.title)
        user_id = entry.assigned_to_user_id or entry.created_by_user_id
        department = "-" if is_all_users else (user_department_codes or {}).get(user_id, "-")
        if start_date == end_date and not full_day:
            from_value, to_value = start_time or "-", end_time or "-"
        else:
            from_value, to_value = f"{start_date:%d.%m.%Y}", f"{end_date:%d.%m.%Y}"
        cleaned_note = TECHNICAL_TAG.sub("", note or "").strip() if note else ""
        cleaned_note = re.sub(r"\s+", " ", cleaned_note).strip()
        if cleaned_note:
            to_value = f"{to_value} - {cleaned_note}"
        from_lines = _wrap_fixed_width(from_value, date_width)
        to_lines = _wrap_fixed_width(to_value, date_width)
        row_count = max(len(from_lines), len(to_lines))
        lines.append(f"| {index:<2} | {person:<{who_width}} | {department:<{department_width}} | {from_lines[0]:<{date_width}} | {to_lines[0]:<{date_width}} |")
        for line_index in range(1, row_count):
            from_line = from_lines[line_index] if line_index < len(from_lines) else ""
            to_line = to_lines[line_index] if line_index < len(to_lines) else ""
            lines.append(f"| {'':<2} | {'':<{who_width}} | {'':<{department_width}} | {from_line:<{date_width}} | {to_line:<{date_width}} |")
        lines.append(border)
    return lines


def _normalize_section(lines: list[str]) -> str:
    return "\n".join(lines).replace("\n\n\n", "\n\n").strip()


def render_plain_text(subject: str, report_day: date, tomorrow: date, sections: list[dict[str, str]]) -> str:
    blocks = [subject, f"Sot: {report_day:%d.%m.%Y}", f"Neser: {tomorrow:%d.%m.%Y}", ""]
    current_group = ""
    for index, section in enumerate(sections, 1):
        group = _section_group_label(section["title"], section.get("section_key"))
        if group != current_group:
            blocks.append(group)
            current_group = group
        blocks.append(
            f"{index}. {section['title']}\n{_strip_status_markers(section.get('body') or '')}".strip()
        )
    return "\n\n".join(blocks)


def _section_group_label(title: str, section_key: str | None = None) -> str:
    from app.services.meeting_point_manual_sync import section_group_label

    return section_group_label("meetings", title, section_key)


def _render_group_label_html(label: str) -> str:
    # Table cell (not a bare div) so Outlook desktop keeps the band visible.
    return (
        "<table role=\"presentation\" width=\"100%\" cellspacing=\"0\" cellpadding=\"0\" border=\"0\" "
        "style=\"width:100%;border-collapse:collapse;margin:22px 0 10px;\">"
        "<tr><td style=\"padding:9px 11px;background:#f1f5f9;border:1px solid #d7dee8;"
        "color:#334155;font-family:Arial,sans-serif;font-size:12px;font-weight:700;"
        "text-transform:uppercase;letter-spacing:.02em;\">"
        f"{html.escape(label)}"
        "</td></tr></table>"
    )


def _render_section_block_html(index: int, title: str, body: str) -> str:
    return (
        "<table role=\"presentation\" width=\"100%\" cellspacing=\"0\" cellpadding=\"0\" border=\"0\" "
        "style=\"width:100%;border-collapse:collapse;margin:22px 0 0;\">"
        "<tr><td style=\"padding:0;\">"
        f"<h2 style=\"font-size:14px;margin:0 0 8px;color:#0f172a;font-family:Arial,sans-serif;\">"
        f"{index}. {html.escape(title)}</h2>"
        f"{_render_section_body_html(body)}"
        "</td></tr></table>"
    )


def _wrap_report_email_html(subject: str, subtitle_html: str, section_html: str) -> str:
    """Outlook-safe HTML shell shared by M1 / M2 / M3.

    Outlook desktop (Word engine) ignores many CSS rules and max-width. Use a fixed
    width=600 table, inline styles, and avoid media queries that force all cells to
    max-width:100% (that collapses report tables in Outlook).
    """
    return f"""<!doctype html>
<html><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1.0">
<!--[if mso]>
<style type="text/css">
table, td {{ font-family: Arial, sans-serif !important; }}
</style>
<![endif]-->
<style>
body{{font-family:Arial,sans-serif;color:#111827;background:#f8fafc;margin:0;padding:24px}}
h1{{font-size:22px;margin:0 0 8px}}p{{margin:0 0 18px;color:#475569}}
h2{{font-size:14px;margin:22px 0 8px;color:#0f172a}}
.report-table{{width:100%;border-collapse:collapse;table-layout:auto;font:12px/1.3 Arial,sans-serif}}
.report-table th{{background:#e5e7eb;color:#111827;text-align:left;font-weight:700;border:1px solid #cbd5e1;padding:4px 5px;vertical-align:top}}
.report-table td{{border:1px solid #cbd5e1;padding:4px 5px;vertical-align:top}}
.report-table .n{{white-space:nowrap}}.report-table tr.todo td{{background:#fbcfe8;color:#111827}}.report-table tr.in-progress td{{background:#fef3c7;color:#111827}}.report-table tr.waiting td{{background:#ffedd5;color:#9a3412}}.report-table tr.waiting-client td{{background:#e2c15b;color:#4f3a00}}.report-table tr.done td{{background:#d4ffe1;color:#111827}}.report-table tr.late td{{background:#fee2e2;color:#111827}}.report-table tr.deadline td{{background:#dc2626;color:#fff}}.report-table tr.eight-am td{{background:#fff;color:#111827;border-top:3px solid #dc2626;border-bottom:3px solid #dc2626}}.report-table tr.eight-am td:first-child{{border-left:3px solid #dc2626}}.report-table tr.eight-am td:last-child{{border-right:3px solid #dc2626}}.report-table tr.notes td{{background:#dbeafe;color:#111827}}.report-table .disk-yes,.report-table .held{{background:#dcfce7!important;color:#166534!important;font-weight:700;text-align:center}}.report-table .disk-no,.report-table .canceled{{background:#fee2e2!important;color:#991b1b!important;font-weight:700;text-align:center}}.report-table tr.highlight td{{border-top:3px solid #2563eb;border-bottom:3px solid #2563eb}}.report-table tr.highlight td:first-child{{border-left:3px solid #2563eb}}.report-table tr.highlight td:last-child{{border-right:3px solid #2563eb}}.report-table tr.highlight .title{{color:#2563eb;font-weight:700}}
@media only screen and (max-width:600px){{
body{{padding:8px!important}}
h1{{font-size:18px!important;line-height:1.2!important}}
h2{{font-size:13px!important;line-height:1.25!important}}
pre{{font-size:12px!important;padding:10px!important}}
.report-table th,.report-table td{{font-size:11px!important;padding:3px 4px!important;line-height:1.25!important}}
}}
</style></head>
<body style="font-family:Arial,sans-serif;color:#111827;background:#f8fafc;margin:0;padding:0;">
<table role="presentation" width="100%" cellspacing="0" cellpadding="0" border="0" style="width:100%;background:#f8fafc;border-collapse:collapse;">
<tr><td align="center" style="padding:16px 8px;">
<!--[if mso]>
<table role="presentation" width="600" cellspacing="0" cellpadding="0" border="0"><tr><td>
<![endif]-->
<table role="presentation" width="600" cellspacing="0" cellpadding="0" border="0" style="width:600px;max-width:100%;background:#ffffff;border:1px solid #e5e7eb;border-collapse:collapse;">
<tr><td style="padding:16px 18px;font-family:Arial,sans-serif;">
<h1 style="font-size:22px;margin:0 0 8px;font-family:Arial,sans-serif;color:#111827;">{html.escape(subject)}</h1>
<p style="margin:0 0 18px;color:#475569;font-family:Arial,sans-serif;">{subtitle_html}</p>
{section_html}
</td></tr></table>
<!--[if mso]>
</td></tr></table>
<![endif]-->
</td></tr></table>
</body></html>"""


def _parse_ascii_cells(line: str) -> list[str]:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return [
        re.sub(r"^(.+?)\s+/\s+(DUE:)", r"\1\n\2", cell, count=1)
        if cell.startswith("START:")
        else cell
        for cell in cells
    ]


def _ascii_table_block(lines: list[str], start: int) -> tuple[list[str], int] | None:
    """Read one ASCII table, allowing blank editor rows between its lines.

    Report sections are editable one line at a time.  Older saved drafts can
    therefore contain blank lines between a table border, header, and data
    rows.  Those blanks are formatting only; they must not turn the table into
    a plain-text block in the email or its native attachment.
    """
    if start >= len(lines) or not lines[start].lstrip().startswith("+-"):
        return None

    table_lines: list[str] = []
    index = start
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("+-") or stripped.startswith("|"):
            table_lines.append(stripped)
            index += 1
            continue
        if stripped:
            break

        # Keep blanks only when they separate two lines of the same table.
        next_index = index + 1
        while next_index < len(lines) and not lines[next_index].strip():
            next_index += 1
        if next_index < len(lines) and (
            lines[next_index].lstrip().startswith("+-")
            or lines[next_index].lstrip().startswith("|")
        ):
            index = next_index
            continue
        break

    has_header = any(line.startswith("|") for line in table_lines)
    return (table_lines, index) if has_header else None


def _normalized_table_header(value: str) -> str:
    """Map Albanian report headers (and older English drafts) to semantic names."""
    normalized = value.strip().upper()
    return {
        "KOHA": "TIME",
        "TITULLI": "TITLE",
        "KUSH": "WHO",
        "NGA": "FROM",
        "NE": "TO",
        "DERI": "TO",
        "TOTALI": "COUNT",
        "LLOJI": "TYPE",
        "KRIJUAR": "ADDED",
        "ARSYEJA": "REASON",
        "KOMENT": "COMMENT",
        "PRODUKTE": "PRODUCTS",
    }.get(normalized, normalized)


def _table_tone_from_label(label: str) -> str:
    normalized = label.strip().upper().rstrip(":")
    if (
        normalized == "TODO"
        or "DETYRAT E REJA" in normalized
        or "DET TE REJA" in normalized
    ):
        return "todo"
    if "IN PROGRESS" in normalized:
        return "in-progress"
    if (
        "WAITING FOR CLIENT" in normalized
        or "WAITING CLIENT" in normalized
        or "WAITING_CLIENT" in normalized
        or "DT WFE" in normalized
    ):
        return "waiting-client"
    if "WAITING" in normalized:
        return "waiting"
    if normalized == "DONE" or "DET E KRYERA" in normalized:
        return "done"
    if "LATE" in normalized:
        return "late"
    if "DEADLINE" in normalized:
        return "deadline"
    if "NOTES" in normalized:
        return "notes"
    return ""


def _table_tone_from_status(status: str) -> str:
    return _table_tone_from_label(status.replace("_", " "))


def _table_tone_from_type(task_type: str) -> str:
    normalized = task_type.strip().upper()
    if "08:00" in normalized:
        return "eight-am"
    if "DEADLINE" in normalized:
        return "deadline"
    return ""


def _has_negative_product_delta(value: str) -> bool:
    return bool(re.search(r"\(\s*-\d+\s*\)", value or ""))


def _priority_task_type_rank(task_type: str) -> int:
    normalized = task_type.strip().upper()
    if normalized == "08:00":
        return 0
    if "08:00" in normalized:
        return 1
    if "DEADLINE" in normalized:
        return 2
    return 3


def _sort_priority_task_rows(header: list[str], rows: list[list[str]]) -> list[list[str]]:
    type_index = next(
        (index for index, cell in enumerate(header) if _normalized_table_header(cell) == "TYPE"),
        None,
    )
    if type_index is None or not any(
        len(row) > type_index and _priority_task_type_rank(row[type_index]) < 3 for row in rows
    ):
        return rows
    return sorted(
        rows,
        key=lambda row: _priority_task_type_rank(row[type_index]) if len(row) > type_index else 3,
    )


def _table_tone_styles(tone: str) -> tuple[str, str]:
    if tone == "todo":
        return "#fbcfe8", "#111827"
    if tone == "in-progress":
        return "#fef3c7", "#111827"
    if tone == "waiting":
        return "#ffedd5", "#9a3412"
    if tone == "waiting-client":
        return "#e2c15b", "#4f3a00"
    if tone == "done":
        return "#d4ffe1", "#111827"
    if tone == "late":
        return "#fee2e2", "#111827"
    if tone == "deadline":
        return "#dc2626", "#ffffff"
    if tone == "notes":
        return "#dbeafe", "#111827"
    if tone == "product-negative":
        # Keep the normal in-progress background while making the entire
        # under-plan row's text red in HTML, Word, and PNG output.
        return "#fef3c7", "#dc2626"
    return "#f8fafc", "#111827"


def _table_cell_style_override(header: str, value: str) -> tuple[str, str, bool] | None:
    """Return per-cell styling shared by email-equivalent native exports."""
    header_name = _normalized_table_header(header)
    normalized = value.strip().upper()
    if header_name == "DISK":
        if normalized == "YES":
            return "#DCFCE7", "#166534", True
        if normalized == "NO":
            return "#FEE2E2", "#991B1B", True
    if header_name in {"MBAJTUR?", "MBAJTUR"}:
        if value.strip() == "\u2713":
            return "#DCFCE7", "#166534", True
        if value.strip() == "\u2715":
            return "#FEE2E2", "#991B1B", True
    if header_name == "ADDED":
        if normalized == "THIS W":
            return "#BAE6FD", "#0C4A6E", True
        if normalized == "LAST W":
            return "#FDE68A", "#78350F", True
    return None


def _is_stacked_start_due_cell(header: str, value: str) -> bool:
    return (
        _normalized_table_header(header) in {"FROM", "TO"}
        and value.startswith("START:")
        and "\nDUE:" in value
    )


def _is_m3_start_due_table(caption: str) -> bool:
    return caption.strip().upper().rstrip(":") == "SHTYER START DHE DUE DATE"


def _render_table_cell_html(
    header: str,
    value: str,
    *,
    strong_start_due_divider: bool = False,
) -> str:
    if not _is_stacked_start_due_cell(header, value):
        return html.escape(value).replace(chr(10), "<br>")
    start_line, due_line = value.split("\n", 1)
    divider = "3px solid #334155" if strong_start_due_divider else "1px solid #94a3b8"
    return (
        f'<div style="padding:4px 5px 2px;border-bottom:{divider};">'
        f"{html.escape(start_line)}</div>"
        '<div style="padding:2px 5px 4px;">'
        f"{html.escape(due_line)}</div>"
    )


def _is_overdue_tyo_value(value: str) -> bool:
    normalized = str(value or "").strip().upper()
    if normalized == "Y":
        return True
    try:
        return int(normalized) >= 2
    except ValueError:
        return False


def _append_cell_style_attribute(attribute: str, css: str) -> str:
    if not attribute:
        return f' style="{css}"'
    if attribute.endswith('"') and ' style="' in attribute:
        return f'{attribute[:-1]}{css}"'
    return attribute


def _is_m1_ga_hv_dv_table(caption: str) -> bool:
    return caption.strip().upper().rstrip(":") in {"GA TASKS", "HV TASKS", "DV TASKS"}


def _render_ascii_table_html(lines: list[str], tone: str = "", caption: str = "") -> str:
    table_rows = [_parse_ascii_cells(line) for line in lines if line.startswith("|")]
    if not table_rows:
        return ""
    header, body_rows = table_rows[0], table_rows[1:]
    header, body_rows = _normalize_meeting_status_table(header, body_rows)
    body_rows = _merge_ascii_continuation_rows(header, body_rows)
    body_rows = _sort_priority_task_rows(header, body_rows)
    status_index = next((index for index, cell in enumerate(header) if _normalized_table_header(cell) == "STATUS"), None)
    type_index = next((index for index, cell in enumerate(header) if _normalized_table_header(cell) == "TYPE"), None)
    title_index = next((index for index, cell in enumerate(header) if _normalized_table_header(cell) == "TITLE"), None)
    products_index = next((index for index, cell in enumerate(header) if _normalized_table_header(cell) == "PRODUCTS"), None)
    row_tones: list[str] = []
    highlighted_meeting_rows: list[bool] = []
    cleaned_body_rows: list[list[str]] = []
    for row in body_rows:
        row = list(row)
        row_tone = tone
        if status_index is not None and len(row) > status_index:
            row_tone = _table_tone_from_status(row[status_index]) or row_tone
        if type_index is not None and len(row) > type_index:
            row_tone = _table_tone_from_type(row[type_index]) or row_tone
        if title_index is not None and len(row) > title_index:
            cleaned_title, marker_status = _split_status_marker(row[title_index])
            cleaned_title, is_highlighted_meeting = _split_meeting_highlight_marker(cleaned_title)
            row[title_index] = cleaned_title
            if marker_status:
                row_tone = _table_tone_from_status(marker_status) or row_tone
        else:
            is_highlighted_meeting = False
        if products_index is not None and len(row) > products_index and _has_negative_product_delta(row[products_index]):
            row_tone = "product-negative"
        row_tones.append(row_tone)
        highlighted_meeting_rows.append(is_highlighted_meeting)
        cleaned_body_rows.append(row)
    body_rows = cleaned_body_rows
    if status_index is not None:
        header = [cell for index, cell in enumerate(header) if index != status_index]
        body_rows = [
            [cell for index, cell in enumerate(row) if index != status_index]
            for row in body_rows
        ]
    column_widths = _email_column_widths(header)
    header_html = "".join(
        f"<th{_email_column_width_attr(column_widths[index])}{_email_column_class(cell)}>{html.escape(cell)}</th>"
        for index, cell in enumerate(header)
    )
    colgroup_html = "".join(
        f"<col{_email_column_width_attr(width)}>"
        for width in column_widths
    )
    canceled_index = next((index for index, cell in enumerate(header) if _normalized_table_header(cell) == "ANULUAR"), None)
    meeting_status_index = next((index for index, cell in enumerate(header) if _normalized_table_header(cell) == "MBAJTUR?"), None)
    disk_index = next((index for index, cell in enumerate(header) if _normalized_table_header(cell) == "DISK"), None)
    added_index = next((index for index, cell in enumerate(header) if _normalized_table_header(cell) == "ADDED"), None)
    am_pm_index = next((index for index, cell in enumerate(header) if _normalized_table_header(cell) == "AM/PM"), None)
    table_class = f"report-table report-table-{tone}" if tone else "report-table"
    body_html_parts = []
    for row_index, row in enumerate(body_rows):
        row_tone = row_tones[row_index] if row_index < len(row_tones) else tone
        is_highlighted_meeting = highlighted_meeting_rows[row_index] if row_index < len(highlighted_meeting_rows) else False
        row_background, row_foreground = _table_tone_styles(row_tone)
        is_canceled = (
            canceled_index is not None
            and len(row) > canceled_index
            and bool(row[canceled_index].strip())
        )
        previous_period = next(
            (
                body_rows[previous_index][am_pm_index].strip().upper()
                for previous_index in range(row_index - 1, -1, -1)
                if am_pm_index is not None
                and len(body_rows[previous_index]) > am_pm_index
                and body_rows[previous_index][am_pm_index].strip().upper() in {"AM", "PM"}
            ),
            "",
        )
        current_period = (
            row[am_pm_index].strip().upper()
            if am_pm_index is not None and len(row) > am_pm_index
            else ""
        )
        has_am_pm_divider = (
            _is_m1_ga_hv_dv_table(caption)
            and current_period == "PM"
            and previous_period == "AM"
        )
        row_cells = []
        for index, cell in enumerate(row):
            cell_classes = [_email_column_class_name(header[index])]
            cell_style = ""
            current_cell = cell
            if is_canceled:
                cell_classes.append("canceled")
            if tone == "notes" and disk_index is not None and index == disk_index:
                disk_value = cell.strip().upper()
                if disk_value == "YES":
                    cell_classes.append("disk-yes")
                elif disk_value == "NO":
                    cell_classes.append("disk-no")
            if meeting_status_index is not None and index == meeting_status_index:
                symbol = cell.strip()
                if symbol == "\u2713":
                    cell_classes.append("held")
                elif symbol == "\u2715":
                    cell_classes.append("canceled")
            if added_index is not None and index == added_index:
                created_week = cell.strip().upper()
                if created_week == "THIS W":
                    cell_classes.append("created-this-week")
                    cell_style = (
                        ' bgcolor="#bae6fd" style="background-color:#bae6fd!important;'
                        'color:#0c4a6e!important;font-weight:700;"'
                    )
                elif created_week == "LAST W":
                    cell_classes.append("created-last-week")
                    cell_style = (
                        ' bgcolor="#fde68a" style="background-color:#fde68a!important;'
                        'color:#78350f!important;font-weight:700;"'
                    )
            if _normalized_table_header(header[index]) == "T/Y/O" and _is_overdue_tyo_value(current_cell):
                cell_classes.append("tyo-overdue")
                cell_style = (
                    ' bgcolor="#dc2626" style="background-color:#dc2626!important;'
                    'color:#ffffff!important;font-weight:400;text-align:left;"'
                )
            if _is_stacked_start_due_cell(header[index], current_cell):
                cell_style = ' style="padding:0!important;"'
            if has_am_pm_divider:
                cell_classes.append("am-pm-divider")
                cell_style = _append_cell_style_attribute(
                    cell_style,
                    "border-top:3px solid #334155!important;",
                )
            if is_highlighted_meeting:
                if _normalized_table_header(header[index]) == "TITLE":
                    cell_classes.append("title")
            row_cells.append(
                f"<td{_email_column_width_attr(column_widths[index])}{cell_style} class=\"{' '.join(filter(None, cell_classes))}\">"
                f"{_render_table_cell_html(header[index], current_cell, strong_start_due_divider=_is_m3_start_due_table(caption))}</td>"
            )
        row_classes = " ".join(filter(None, (row_tone, "highlight" if is_highlighted_meeting else "")))
        row_style = (
            f' bgcolor="{row_background}" style="background-color:{row_background}!important;'
            f'color:{row_foreground}!important;"'
            if row_tone
            else ""
        )
        body_html_parts.append(
            f"<tr{row_style} class=\"{row_classes}\">"
            + "".join(row_cells)
            + "</tr>"
        )
    body_html = "".join(body_html_parts)
    caption_html = ""
    if caption.strip():
        caption_html = (
            "<tr><td style=\"background:#f8fafc;color:#111827;font-weight:700;"
            "border:1px solid #cbd5e1;border-bottom:0;padding:8px 10px;"
            "font-family:Arial,sans-serif;font-size:13px;\">"
            f"{html.escape(caption.strip())}</td></tr>"
        )
    return (
        "<table role=\"presentation\" width=\"100%\" cellspacing=\"0\" cellpadding=\"0\" border=\"0\" "
        "style=\"width:100%;border-collapse:collapse;margin:8px 0 12px;\">"
        f"{caption_html}<tr><td style=\"padding:0;\">"
        f"<table class=\"{table_class}\" width=\"100%\" cellspacing=\"0\" cellpadding=\"0\" border=\"0\" "
        "style=\"width:100%;border-collapse:collapse;\">"
        f"<colgroup>{colgroup_html}</colgroup><thead><tr>{header_html}</tr></thead><tbody>{body_html}</tbody></table>"
        "</td></tr></table>"
    )


def _normalize_meeting_status_table(header: list[str], body_rows: list[list[str]]) -> tuple[list[str], list[list[str]]]:
    """Render older saved meeting drafts with the current compact status column."""
    normalized = [_normalized_table_header(cell) for cell in header]
    if not {"MBAJTUR", "ANULUAR", "PA STATUS"}.issubset(set(normalized)):
        return header, body_rows
    if "TITLE" not in normalized:
        return header, body_rows

    def cell_at(row: list[str], name: str) -> str:
        try:
            index = normalized.index(name)
        except ValueError:
            return ""
        return row[index].strip() if len(row) > index else ""

    new_header: list[str] = []
    for cell in header:
        name = _normalized_table_header(cell)
        if name in {"MBAJTUR", "ANULUAR", "PA STATUS"}:
            continue
        if name == "TITLE":
            new_header.append("MBAJTUR?")
        new_header.append(cell)

    new_rows: list[list[str]] = []
    for row in body_rows:
        status_icon = "\u2713" if cell_at(row, "MBAJTUR") else "\u2715" if cell_at(row, "ANULUAR") else ""
        new_row: list[str] = []
        for index, cell in enumerate(header):
            name = _normalized_table_header(cell)
            if name in {"MBAJTUR", "ANULUAR", "PA STATUS"}:
                continue
            if name == "TITLE":
                new_row.append(status_icon)
            new_row.append(row[index] if len(row) > index else "")
        new_rows.append(new_row)
    return new_header, new_rows


def _email_column_width_attr(width: str) -> str:
    return "" if width == "auto" else f" width=\"{width}\""


def _email_column_class_name(header_cell: str) -> str:
    name = _normalized_table_header(header_cell)
    if name in {"NR", "TIME", "ORA", "KOHA", "DISK", "LATE", "T/Y/O", "FROM", "TO", "DATA", "DATE", "DEP", "TYPE", "AM/PM", "MBAJTUR", "MBAJTUR?", "ANULUAR", "PA STATUS"}:
        return "n"
    return ""


def _email_column_class(header_cell: str) -> str:
    class_name = _email_column_class_name(header_cell)
    return f' class="{class_name}"' if class_name else ""


def _email_column_widths(header: list[str]) -> list[str]:
    """Give utility columns only the space they need; reserve the rest for title/note text."""
    if not header:
        return []
    # Keep utility columns content-sized; TITLE/NOTE take the remaining width.
    fixed_by_name = {
        "NR": "24",
        "WHO": "34",
        "DEP": "34",
        "ADDED": "48",
        "TYPE": "42",
        "AM/PM": "42",
        "FROM": "38",
        "TO": "38",
        "TIME": "46",
        "ORA": "46",
        "KOHA": "46",
        "DATA": "62",
        "DATE": "62",
        "DISK": "42",
        "LATE": "54",
        "T/Y/O": "42",
        "KATEGORIA": "1%",
        "LISTA": "1%",
        "MBAJTUR": "58",
        "MBAJTUR?": "62",
        "ANULUAR": "58",
        "PA STATUS": "64",
    }
    content_names = {
        "TITLE", "NOTE", "SHENIMI", "PERSHKRIMI", "DESCRIPTION", "PYETJA",
        "REASON", "COMMENT",
    }
    normalized = [_normalized_table_header(cell) for cell in header]
    widths = ["auto" if name in content_names else fixed_by_name.get(name, "56") for name in normalized]
    if normalized and not any(name in content_names for name in normalized):
        widths[-1] = "auto"
    return widths


def _primary_text_column_index(header: list[str]) -> int:
    normalized = [_normalized_table_header(cell) for cell in header]
    for name in ("NOTE", "TITLE", "PYETJA", "SHENIMI", "PERSHKRIMI", "DESCRIPTION"):
        if name in normalized:
            return normalized.index(name)
    return min(2, max(len(header) - 1, 0))


def _merge_ascii_continuation_rows(header: list[str], rows: list[list[str]]) -> list[list[str]]:
    if not header:
        return rows
    width = len(header)
    text_index = _primary_text_column_index(header)
    nr_index = next((index for index, cell in enumerate(header) if _normalized_table_header(cell) == "NR"), 0)
    merged: list[list[str]] = []
    for row in rows:
        normalized = (row + [""] * width)[:width]
        has_text_continuation = bool(normalized[text_index].strip())
        other_cells_empty = all(
            not normalized[index].strip()
            for index in range(width)
            if index != text_index
        )
        is_wrapped_row = merged and not normalized[nr_index].strip() and any(cell.strip() for cell in normalized)
        if is_wrapped_row:
            previous = merged[-1]
            for index, value in enumerate(normalized):
                stripped = value.strip()
                if not stripped:
                    continue
                previous[index] = (
                    f"{previous[index]}\n{stripped}"
                    if previous[index].strip()
                    else stripped
                )
        elif merged and has_text_continuation and other_cells_empty:
            previous = merged[-1]
            previous[text_index] = (
                f"{previous[text_index]}\n{normalized[text_index].strip()}"
                if previous[text_index].strip()
                else normalized[text_index].strip()
            )
        else:
            merged.append(normalized)
    return merged


def _ascii_table_is_empty(lines: list[str]) -> bool:
    table_rows = [_parse_ascii_cells(line) for line in lines if line.startswith("|")]
    if len(table_rows) != 2:
        return False
    row = table_rows[1]
    return any(cell in {"(Asnje detyre)", "(Asnje takim)"} for cell in row)


def _is_keyed_prompt_line(value: str) -> bool:
    stripped = value.strip()
    keyed = re.match(r"^([A-Z][A-Z0-9 /&()?.:+-]*:)\s*(.*)$", stripped)
    return bool(keyed and keyed.group(1)[:-1] == keyed.group(1)[:-1].upper())


def _render_keyed_prompt_html(line: str) -> str:
    stripped = line.strip()
    keyed = re.match(r"^([A-Z][A-Z0-9 /&()?.:+-]*:)\s*(.*)$", stripped)
    if keyed and keyed.group(1)[:-1] == keyed.group(1)[:-1].upper():
        label, rest = keyed.group(1), keyed.group(2)
        rendered = f"<strong>{html.escape(label)}</strong>"
        if rest:
            rendered += f" {html.escape(rest)}"
    else:
        rendered = html.escape(stripped)
    return (
        "<table role=\"presentation\" width=\"100%\" cellspacing=\"0\" cellpadding=\"0\" border=\"0\" "
        "style=\"width:100%;border-collapse:collapse;margin:0 0 8px;\">"
        "<tr><td style=\"background:#f8fafc;border:1px solid #e5e7eb;padding:8px 10px;"
        "font-family:Arial,sans-serif;font-size:13px;line-height:1.45;color:#0f172a;\">"
        f"{rendered}</td></tr></table>"
    )


def _is_guidance_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped or re.match(r"^\d+\.\s", stripped):
        return False
    # Indented descriptions under numbered questions (even ALL-CAPS).
    return bool(re.match(r"^\s{2,}\S", line))


def _render_text_block_html(lines: list[str]) -> str:
    non_empty = [line for line in lines if line.strip()]
    if non_empty and all(_is_keyed_prompt_line(line) for line in non_empty):
        return "".join(_render_keyed_prompt_html(line) for line in non_empty)

    rendered_lines = []
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        keyed = re.match(r"^([A-Z][A-Z0-9 /&()?.:+-]*:)\s*(.*)$", stripped)
        if keyed and keyed.group(1)[:-1] == keyed.group(1)[:-1].upper():
            label, rest = keyed.group(1), keyed.group(2)
            rendered = f"<strong>{html.escape(label)}</strong>"
            if rest:
                rendered += f" {html.escape(rest)}"
            rendered_lines.append(rendered)
        elif stripped and len(stripped) <= 45 and stripped.endswith((": 0", ":")):
            rendered_lines.append(f"<strong>{html.escape(line)}</strong>")
        elif re.match(r"^\d+\.\s", stripped):
            # Keep question + following indented guidance in one visual unit.
            block = html.escape(stripped)
            while index + 1 < len(lines) and _is_guidance_line(lines[index + 1]):
                index += 1
                guidance = lines[index].strip()
                block += (
                    "<br>"
                    f"<span style=\"display:inline-block;padding-left:1.1em;margin-top:2px;"
                    f"color:#64748b;font-style:italic;font-size:12px;font-weight:normal;\">"
                    f"{html.escape(guidance)}</span>"
                )
            rendered_lines.append(block)
        elif _is_guidance_line(line):
            rendered_lines.append(
                f"<span style=\"display:inline-block;padding-left:1.1em;color:#64748b;"
                f"font-style:italic;font-size:12px;font-weight:normal;\">{html.escape(stripped)}</span>"
            )
        else:
            rendered_lines.append(html.escape(line))
        index += 1
    return (
        "<pre style=\"white-space:pre-wrap;font-family:Arial,sans-serif;font-size:13px;line-height:1.45;"
        "background:#f8fafc;border:1px solid #e5e7eb;padding:12px;margin:0;\">"
        f"{chr(10).join(rendered_lines).strip()}</pre>"
    )


def _render_section_body_html(body: str) -> str:
    lines = body.splitlines()
    chunks: list[str] = []
    text_buffer: list[str] = []
    index = 0

    def flush_text() -> None:
        if any(line.strip() for line in text_buffer):
            chunks.append(_render_text_block_html(text_buffer))
        text_buffer.clear()

    def current_table_tone() -> str:
        for previous in reversed(text_buffer):
            if previous.strip():
                return _table_tone_from_label(previous)
        return ""

    def mark_current_label_empty() -> None:
        for previous_index in range(len(text_buffer) - 1, -1, -1):
            stripped = text_buffer[previous_index].strip()
            if stripped:
                text_buffer[previous_index] = (
                    f"{text_buffer[previous_index]} 0"
                    if stripped.endswith(":")
                    else f"{text_buffer[previous_index]}: 0"
                )
                return

    def pop_current_table_label() -> str:
        for previous_index in range(len(text_buffer) - 1, -1, -1):
            stripped = text_buffer[previous_index].strip()
            if not stripped:
                continue
            if len(stripped) <= 45 and (stripped.endswith(":") or stripped.isupper() or re.match(r"^\d{1,2}:\d{2}:?$", stripped)):
                label = stripped
                del text_buffer[previous_index:]
                return label
            return ""
        return ""

    while index < len(lines):
        table_block = _ascii_table_block(lines, index)
        if table_block:
            table_lines, index = table_block
            tone = current_table_tone()
            if _ascii_table_is_empty(table_lines):
                mark_current_label_empty()
                continue
            caption = pop_current_table_label()
            flush_text()
            chunks.append(_render_ascii_table_html(table_lines, tone, caption))
            continue
        text_buffer.append(lines[index])
        index += 1
    flush_text()
    return "".join(chunk for chunk in chunks if chunk)


def render_html(subject: str, report_day: date, tomorrow: date, sections: list[dict[str, str]]) -> str:
    section_chunks: list[str] = []
    current_group = ""
    for index, section in enumerate(sections, 1):
        group = _section_group_label(section["title"], section.get("section_key"))
        if group != current_group:
            section_chunks.append(_render_group_label_html(group))
            current_group = group
        section_chunks.append(
            _render_section_block_html(index, section["title"], section.get("body") or "")
        )
    return _wrap_report_email_html(
        subject,
        f"Sot: {report_day:%d.%m.%Y} &nbsp; Neser: {tomorrow:%d.%m.%Y}",
        "".join(section_chunks),
    )


def _section_report_table_rows(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = [_parse_ascii_cells(line) for line in lines if line.startswith("|")]
    if not rows:
        return [], []
    header, body_rows = _normalize_meeting_status_table(rows[0], rows[1:])
    body_rows = _merge_ascii_continuation_rows(header, body_rows)
    return header, _sort_priority_task_rows(header, body_rows)


def _section_report_table_model(lines: list[str], tone: str = "") -> tuple[list[str], list[list[str]], list[str], list[bool]]:
    """Return the same cleaned table data and row treatment used in the email."""
    header, body_rows = _section_report_table_rows(lines)
    if not header:
        return [], [], [], []
    status_index = next((index for index, cell in enumerate(header) if _normalized_table_header(cell) == "STATUS"), None)
    type_index = next((index for index, cell in enumerate(header) if _normalized_table_header(cell) == "TYPE"), None)
    title_index = next((index for index, cell in enumerate(header) if _normalized_table_header(cell) == "TITLE"), None)
    products_index = next((index for index, cell in enumerate(header) if _normalized_table_header(cell) == "PRODUCTS"), None)
    cleaned_rows: list[list[str]] = []
    row_tones: list[str] = []
    highlights: list[bool] = []
    for source_row in body_rows:
        row = (list(source_row) + [""] * len(header))[:len(header)]
        row_tone = tone
        if status_index is not None and row[status_index]:
            row_tone = _table_tone_from_status(row[status_index]) or row_tone
        if type_index is not None and row[type_index]:
            row_tone = _table_tone_from_type(row[type_index]) or row_tone
        highlighted = False
        if title_index is not None:
            title, marker_status = _split_status_marker(row[title_index])
            title, highlighted = _split_meeting_highlight_marker(title)
            row[title_index] = title
            row_tone = _table_tone_from_status(marker_status) or row_tone
        if products_index is not None and _has_negative_product_delta(row[products_index]):
            row_tone = "product-negative"
        if status_index is not None:
            row = [cell for index, cell in enumerate(row) if index != status_index]
        cleaned_rows.append(row)
        row_tones.append(row_tone)
        highlights.append(highlighted)
    if status_index is not None:
        header = [cell for index, cell in enumerate(header) if index != status_index]
    return header, cleaned_rows, row_tones, highlights


def _section_report_blocks(
    body: str,
    *,
    preserve_empty_tables: bool = False,
) -> list[dict[str, Any]]:
    """Parse a section once so native attachments keep the email's structure."""
    blocks: list[dict[str, Any]] = []
    text_buffer: list[str] = []
    lines = body.splitlines()
    position = 0

    def flush_text() -> None:
        if any(line.strip() for line in text_buffer):
            blocks.append({"kind": "text", "lines": list(text_buffer)})
        text_buffer.clear()

    def current_table_tone() -> str:
        for previous in reversed(text_buffer):
            if previous.strip():
                return _table_tone_from_label(previous)
        return ""

    def mark_current_label_empty() -> None:
        for previous_index in range(len(text_buffer) - 1, -1, -1):
            if text_buffer[previous_index].strip():
                label = text_buffer[previous_index].rstrip()
                text_buffer[previous_index] = label if label.endswith("0") else f"{label} 0"
                return

    def pop_current_table_label() -> str:
        for previous_index in range(len(text_buffer) - 1, -1, -1):
            stripped = text_buffer[previous_index].strip()
            if not stripped:
                continue
            if len(stripped) <= 45 and (stripped.endswith(":") or stripped.isupper() or re.match(r"^\d{1,2}:\d{2}:?$", stripped)):
                del text_buffer[previous_index:]
                return stripped
            return ""
        return ""

    while position < len(lines):
        table_block = _ascii_table_block(lines, position)
        if table_block:
            table_lines, position = table_block
            if _ascii_table_is_empty(table_lines) and not preserve_empty_tables:
                mark_current_label_empty()
                continue
            tone = current_table_tone()
            caption = pop_current_table_label()
            header, rows, row_tones, highlights = _section_report_table_model(table_lines, tone)
            flush_text()
            if header:
                blocks.append({"kind": "table", "caption": caption, "header": header, "rows": rows, "tone": tone, "row_tones": row_tones, "highlights": highlights})
            continue
        text_buffer.append(lines[position])
        position += 1
    flush_text()
    return blocks


def _section_report_group_for_code(
    report_code: str, title: str | None, section_key: str | None = None,
) -> str:
    from app.services.meeting_point_manual_sync import section_group_label

    report_kind = {"M1": "morning", "M2": "after_break", "M3": "meetings"}.get(report_code.upper(), "meetings")
    return section_group_label(report_kind, title, section_key)


def _legacy_render_section_report_docx(
    subject: str,
    report_code: str,
    report_day: date,
    sections: list[dict[str, str]],
    *,
    tomorrow: date | None = None,
) -> bytes:
    """Render the M1/M2/M3 section and table format as a Word attachment."""
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    def shade(cell: Any, color: str) -> None:
        fill = OxmlElement("w:shd")
        fill.set(qn("w:fill"), color.lstrip("#"))
        cell._tc.get_or_add_tcPr().append(fill)

    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.55)
    section.left_margin = section.right_margin = Inches(0.5)
    header = document.add_table(rows=1, cols=1)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cell = header.cell(0, 0)
    shade(header_cell, "#2563EB")
    title = header_cell.paragraphs[0]
    title_run = title.add_run(subject)
    title_run.bold = True
    title_run.font.size = Pt(17)
    title_run.font.color.rgb = RGBColor(255, 255, 255)
    date_label = f"{report_code} · {report_day:%d.%m.%Y}"
    if tomorrow is not None:
        date_label += f" · Neser: {tomorrow:%d.%m.%Y}"
    metadata = header_cell.add_paragraph(date_label)
    metadata.runs[0].font.size = Pt(9)
    metadata.runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for index, report_section in enumerate(sections, 1):
        document.add_paragraph()
        section_title = document.add_table(rows=1, cols=1).cell(0, 0)
        shade(section_title, "#E2E8F0")
        title_paragraph = section_title.paragraphs[0]
        title_run = title_paragraph.add_run(f"{index}. {report_section.get('title') or 'Untitled'}")
        title_run.bold = True
        title_run.font.size = Pt(11)

        body_lines = str(report_section.get("body") or "").splitlines()
        position = 0
        while position < len(body_lines):
            line = body_lines[position]
            if line.startswith("+-"):
                table_lines: list[str] = []
                while position < len(body_lines) and (
                    body_lines[position].startswith("+-") or body_lines[position].startswith("|")
                ):
                    table_lines.append(body_lines[position])
                    position += 1
                table_header, table_body = _section_report_table_rows(table_lines)
                if table_header:
                    table = document.add_table(rows=1, cols=len(table_header))
                    table.style = "Table Grid"
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for column, value in enumerate(table_header):
                        cell = table.rows[0].cells[column]
                        cell.text = value
                        shade(cell, "#CBD5E1")
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
                            run.font.size = Pt(8)
                    for row_values in table_body:
                        row = table.add_row()
                        for column, value in enumerate(row_values[: len(table_header)]):
                            cell = row.cells[column]
                            cell.text = value
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(8)
                continue
            if line.strip():
                paragraph = document.add_paragraph(line.strip())
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            position += 1

    output = __import__("io").BytesIO()
    document.save(output)
    return output.getvalue()


def _legacy_render_section_report_png(
    subject: str,
    report_code: str,
    report_day: date,
    sections: list[dict[str, str]],
    *,
    tomorrow: date | None = None,
) -> bytes:
    """Render a readable, single-image PNG version of an M1/M2/M3 report."""
    from PIL import Image, ImageDraw, ImageFont

    font_path = os.getenv("PRIMEFLOW_REPORT_FONT_PATH", r"C:\Windows\Fonts\segoeui.ttf")
    bold_path = r"C:\Windows\Fonts\arialbd.ttf"
    try:
        font = ImageFont.truetype(font_path, 20)
        bold = ImageFont.truetype(bold_path, 21)
        heading = ImageFont.truetype(bold_path, 30)
    except OSError:
        font = bold = heading = ImageFont.load_default()

    rows: list[tuple[str, str]] = []
    for index, report_section in enumerate(sections, 1):
        rows.append(("section", f"{index}. {report_section.get('title') or 'Untitled'}"))
        body_lines = str(report_section.get("body") or "").splitlines()
        position = 0
        while position < len(body_lines):
            if body_lines[position].startswith("+-"):
                table_lines: list[str] = []
                while position < len(body_lines) and (
                    body_lines[position].startswith("+-") or body_lines[position].startswith("|")
                ):
                    table_lines.append(body_lines[position])
                    position += 1
                table_header, table_body = _section_report_table_rows(table_lines)
                if table_header:
                    rows.append(("table-header", " | ".join(table_header)))
                    rows.extend(("table", " | ".join(value for value in row if value)) for row in table_body)
                continue
            text = body_lines[position].strip()
            if text:
                rows.append(("text", text))
            position += 1

    wrapped_rows: list[tuple[str, str]] = []
    for kind, value in rows:
        width = 92 if kind == "section" else 112
        wrapped_rows.extend((kind, line) for line in (textwrap.wrap(value, width=width) or [""]))
    width, margin, line_height = 1600, 48, 31
    height = max(420, 175 + len(wrapped_rows) * line_height + margin)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((margin, 30, width - margin, 130), radius=12, fill="#2563EB")
    draw.text((margin + 20, 50), subject, fill="white", font=heading)
    date_label = f"{report_code} · {report_day:%d.%m.%Y}"
    if tomorrow is not None:
        date_label += f" · Neser: {tomorrow:%d.%m.%Y}"
    draw.text((margin + 22, 96), date_label, fill="white", font=font)
    y = 152
    for kind, value in wrapped_rows:
        if kind == "section":
            draw.rectangle((margin, y, width - margin, y + 27), fill="#E2E8F0")
            draw.text((margin + 12, y + 3), value, fill="#0F172A", font=bold)
        elif kind == "table-header":
            draw.rectangle((margin + 8, y, width - margin - 8, y + 25), fill="#CBD5E1")
            draw.text((margin + 16, y + 2), value, fill="#0F172A", font=bold)
        elif kind == "table":
            draw.text((margin + 16, y + 2), value, fill="#1F2937", font=font)
        else:
            draw.text((margin + 12, y + 2), value, fill="#1F2937", font=font)
        y += line_height
    output = __import__("io").BytesIO()
    image.save(output, format="PNG", optimize=True)
    return output.getvalue()


def render_section_report_docx(
    subject: str,
    report_code: str,
    report_day: date,
    sections: list[dict[str, str]],
    *,
    tomorrow: date | None = None,
) -> bytes:
    """Native Word equivalent of the M1/M2/M3 email report."""
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor, Twips

    def shade(cell: Any, color: str) -> None:
        item = OxmlElement("w:shd")
        item.set(qn("w:fill"), color.lstrip("#"))
        cell._tc.get_or_add_tcPr().append(item)

    def border(cell: Any, color: str = "CBD5E1", size: str = "6") -> None:
        properties = cell._tc.get_or_add_tcPr()
        borders = properties.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            properties.append(borders)
        for edge_name in ("top", "left", "bottom", "right"):
            edge = borders.find(qn(f"w:{edge_name}"))
            if edge is None:
                edge = OxmlElement(f"w:{edge_name}")
                borders.append(edge)
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), size)
            edge.set(qn("w:color"), color.lstrip("#"))

    def cell_style(cell: Any, fill: str, *, color: str = "#111827", bold: bool = False, size: float = 8.5, outline: str = "#CBD5E1") -> None:
        shade(cell, fill)
        border(cell, outline)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
                run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
                run.font.size = Pt(size)
                run.bold = bold
                run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))

    def table_widths(header: list[str]) -> list[int]:
        # Email uses a 600px shell with compact utility columns.  Map the same
        # proportions to Word's usable page width (10166 twips at these margins).
        available = 10166
        raw = _email_column_widths(header)
        widths: list[int | None] = []
        for value in raw:
            if value == "auto":
                widths.append(None)
            elif value.endswith("%"):
                widths.append(max(720, int(available * float(value[:-1]) / 100)))
            else:
                widths.append(int(float(value) * 18))
        fixed = sum(value or 0 for value in widths)
        auto_count = sum(value is None for value in widths)
        if fixed >= available:
            return [max(500, int((value or 500) * available / max(fixed, 1))) for value in widths]
        auto_width = max(960, (available - fixed) // max(auto_count, 1))
        return [value if value is not None else auto_width for value in widths]

    def set_width(cell: Any, width: int) -> None:
        properties = cell._tc.get_or_add_tcPr()
        tc_width = properties.first_child_found_in("w:tcW")
        if tc_width is None:
            tc_width = OxmlElement("w:tcW")
            properties.append(tc_width)
        tc_width.set(qn("w:w"), str(width))
        tc_width.set(qn("w:type"), "dxa")

    def set_table_cell_value(
        cell: Any,
        header: str,
        value: str,
        *,
        strong_start_due_divider: bool = False,
    ) -> None:
        if not _is_stacked_start_due_cell(header, value):
            cell.text = value
            return
        start_line, due_line = value.split("\n", 1)
        cell.text = ""
        start_paragraph = cell.paragraphs[0]
        start_paragraph.add_run(start_line)
        paragraph_properties = start_paragraph._p.get_or_add_pPr()
        paragraph_borders = OxmlElement("w:pBdr")
        bottom_border = OxmlElement("w:bottom")
        bottom_border.set(qn("w:val"), "single")
        bottom_border.set(qn("w:sz"), "12" if strong_start_due_divider else "4")
        bottom_border.set(qn("w:space"), "1")
        bottom_border.set(qn("w:color"), "334155" if strong_start_due_divider else "94A3B8")
        paragraph_borders.append(bottom_border)
        paragraph_properties.append(paragraph_borders)
        cell.add_paragraph(due_line)

    def add_text(lines: list[str]) -> None:
        non_empty = [line for line in lines if line.strip()]
        keyed_only = bool(non_empty) and all(_is_keyed_prompt_line(line) for line in non_empty)
        for line in non_empty:
            value = line.strip()
            if keyed_only:
                cell = document.add_table(rows=1, cols=1).cell(0, 0)
                cell.text = value
                cell_style(cell, "#F8FAFC", bold=True, size=9, outline="#E5E7EB")
                document.add_paragraph().paragraph_format.space_after = Pt(1)
                continue
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            if _is_guidance_line(line):
                paragraph.paragraph_format.left_indent = Inches(0.14)
            run = paragraph.add_run(value)
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
            run.font.size = Pt(9)
            run.bold = _is_keyed_prompt_line(line) or (len(value) <= 45 and value.endswith((":", ": 0")))
            if _is_guidance_line(line):
                run.italic = True
                run.font.color.rgb = RGBColor(100, 116, 139)

    def add_table(block: dict[str, Any]) -> None:
        caption = str(block["caption"] or "")
        if caption:
            cell = document.add_table(rows=1, cols=1).cell(0, 0)
            cell.text = caption
            cell_style(cell, "#F8FAFC", bold=True, size=9)
        header = block["header"]
        table = document.add_table(rows=1, cols=len(header))
        table.autofit = False
        column_widths = table_widths(header)
        for index, grid_column in enumerate(table._tbl.tblGrid.gridCol_lst):
            grid_column.w = Twips(column_widths[index])
        for index, value in enumerate(header):
            cell = table.rows[0].cells[index]
            cell.text = value
            cell_style(cell, "#E5E7EB", bold=True, size=8)
            set_width(cell, column_widths[index])
        for row_index, values in enumerate(block["rows"]):
            row = table.add_row()
            tone = block["row_tones"][row_index] if row_index < len(block["row_tones"]) else block["tone"]
            fill, color = _table_tone_styles(tone)
            highlighted = block["highlights"][row_index] if row_index < len(block["highlights"]) else False
            for column, value in enumerate(values):
                header_name = _normalized_table_header(header[column])
                cell_fill, cell_color, cell_bold = fill, color, False
                override = _table_cell_style_override(header_name, str(value))
                if override is not None:
                    cell_fill, cell_color, cell_bold = override
                elif header_name == "DISK" and str(value).strip().upper() == "YES":
                    cell_fill, cell_color, cell_bold = "#DCFCE7", "#166534", True
                elif header_name == "DISK" and str(value).strip().upper() == "NO":
                    cell_fill, cell_color, cell_bold = "#FEE2E2", "#991B1B", True
                elif header_name == "MBAJTUR?" and str(value).strip() == "✓":
                    cell_fill, cell_color, cell_bold = "#DCFCE7", "#166534", True
                elif header_name == "MBAJTUR?" and str(value).strip() == "✕":
                    cell_fill, cell_color, cell_bold = "#FEE2E2", "#991B1B", True
                if highlighted and header_name == "TITLE":
                    cell_color, cell_bold = "#2563EB", True
                cell = row.cells[column]
                set_table_cell_value(
                    cell,
                    header_name,
                    str(value),
                    strong_start_due_divider=_is_m3_start_due_table(caption),
                )
                outline = "#DC2626" if tone == "eight-am" else ("#2563EB" if highlighted else "#CBD5E1")
                cell_style(cell, cell_fill, color=cell_color, bold=cell_bold, size=8, outline=outline)
                set_width(cell, column_widths[column])
        document.add_paragraph().paragraph_format.space_after = Pt(1)

    document = Document()
    page = document.sections[0]
    page.top_margin = page.bottom_margin = Inches(0.6)
    page.left_margin = page.right_margin = Inches(0.72)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(subject)
    run.font.name, run.font.size, run.bold = "Arial", Pt(16), True
    run.font.color.rgb = RGBColor(17, 24, 39)
    label = f"{report_code} · {report_day:%d.%m.%Y}"
    if tomorrow is not None:
        label += f" · Neser: {tomorrow:%d.%m.%Y}"
    subtitle = document.add_paragraph(label)
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.runs[0].font.name = "Arial"
    subtitle.runs[0].font.size = Pt(9)
    subtitle.runs[0].font.color.rgb = RGBColor(71, 85, 105)

    current_group = ""
    for index, report_section in enumerate(sections, 1):
        group = _section_report_group_for_code(
            report_code, report_section.get("title"), report_section.get("section_key")
        )
        if group != current_group:
            group_cell = document.add_table(rows=1, cols=1).cell(0, 0)
            group_cell.text = group
            cell_style(group_cell, "#F1F5F9", color="#334155", bold=True, outline="#D7DEE8")
            current_group = group
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(4)
        heading_run = heading.add_run(f"{index}. {report_section.get('title') or 'Untitled'}")
        heading_run.font.name, heading_run.font.size, heading_run.bold = "Arial", Pt(10.5), True
        heading_run.font.color.rgb = RGBColor(15, 23, 42)
        for block in _section_report_blocks(str(report_section.get("body") or "")):
            add_table(block) if block["kind"] == "table" else add_text(block["lines"])

    output = __import__("io").BytesIO()
    document.save(output)
    return output.getvalue()


def render_section_report_png(
    subject: str,
    report_code: str,
    report_day: date,
    sections: list[dict[str, str]],
    *,
    tomorrow: date | None = None,
    preserve_empty_tables: bool = False,
) -> bytes:
    """Bitmap equivalent of the email report; tables remain real cell grids."""
    from PIL import Image, ImageDraw, ImageFont

    width, margin = 1200, 46
    try:
        font = ImageFont.truetype(os.getenv("PRIMEFLOW_REPORT_FONT_PATH", r"C:\Windows\Fonts\segoeui.ttf"), 18)
        bold = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 18)
        heading = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 29)
    except OSError:
        font = bold = heading = ImageFont.load_default()
    measure = ImageDraw.Draw(Image.new("RGB", (width, 1), "white"))

    def wrap(value: str, text_font: Any, max_width: int) -> list[str]:
        lines: list[str] = []
        for source in (value or "").splitlines() or [""]:
            words, current = source.split() or [""], ""
            for word in words:
                candidate = word if not current else f"{current} {word}"
                if current and measure.textlength(candidate, font=text_font) > max_width:
                    lines.append(current)
                    current = word
                else:
                    current = candidate
            lines.append(current)
        return lines or [""]

    def widths(header: list[str]) -> list[int]:
        available = width - 2 * margin
        result: list[int | None] = []
        for item in _email_column_widths(header):
            result.append(None if item == "auto" else max(64, int(float(item.rstrip("%")) * (available / 100 if item.endswith("%") else 2))))
        fixed, autos = sum(item or 0 for item in result), sum(item is None for item in result)
        if fixed >= available:
            return [max(44, int((item or 44) * available / fixed)) for item in result]
        return [item if item is not None else max(80, (available - fixed) // max(autos, 1)) for item in result]

    layout: list[dict[str, Any]] = []
    current_group = ""
    for index, report_section in enumerate(sections, 1):
        group = _section_report_group_for_code(
            report_code, report_section.get("title"), report_section.get("section_key")
        )
        if group != current_group:
            layout.append({"kind": "group", "value": group})
            current_group = group
        layout.append({"kind": "section", "value": f"{index}. {report_section.get('title') or 'Untitled'}"})
        layout.extend(
            _section_report_blocks(
                str(report_section.get("body") or ""),
                preserve_empty_tables=preserve_empty_tables,
            )
        )

    def block_height(block: dict[str, Any]) -> int:
        if block["kind"] == "group":
            return 39
        if block["kind"] == "section":
            return 37
        if block["kind"] == "text":
            return 8 + sum(23 * len(wrap(line.strip(), font, width - 2 * margin)) for line in block["lines"] if line.strip())
        column_widths = widths(block["header"])
        rows = 34 + (32 if block["caption"] else 0)
        for row in block["rows"]:
            row_height = max(31, 8 + max(len(wrap(str(value), font, max(20, column_widths[index] - 12))) for index, value in enumerate(row)) * 22)
            if any(
                _is_stacked_start_due_cell(block["header"][index], str(value))
                for index, value in enumerate(row)
            ):
                row_height = max(row_height, 60)
            rows += row_height
        return rows + 14

    height = 105 + sum(block_height(block) for block in layout) + margin
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((margin, 28), subject, fill="#111827", font=heading)
    label = f"{report_code} · {report_day:%d.%m.%Y}"
    if tomorrow is not None:
        label += f" · Neser: {tomorrow:%d.%m.%Y}"
    draw.text((margin, 68), label, fill="#475569", font=font)
    y = 105
    for block in layout:
        if block["kind"] == "group":
            draw.rectangle((margin, y, width - margin, y + 30), fill="#F1F5F9", outline="#D7DEE8")
            draw.text((margin + 10, y + 6), block["value"], fill="#334155", font=bold)
            y += 39
        elif block["kind"] == "section":
            draw.text((margin, y + 5), block["value"], fill="#0F172A", font=bold)
            y += 37
        elif block["kind"] == "text":
            non_empty = [line for line in block["lines"] if line.strip()]
            keyed_only = bool(non_empty) and all(_is_keyed_prompt_line(line) for line in non_empty)
            for line in non_empty:
                value = line.strip()
                if keyed_only:
                    draw.rectangle((margin, y, width - margin, y + 28), fill="#F8FAFC", outline="#E5E7EB")
                    draw.text((margin + 10, y + 5), value, fill="#0F172A", font=bold)
                    y += 35
                    continue
                text_font = bold if (_is_keyed_prompt_line(line) or (len(value) <= 45 and value.endswith((":", ": 0")))) else font
                color, x = ("#64748B", margin + 18) if _is_guidance_line(line) else ("#111827", margin)
                for wrapped in wrap(value, text_font, width - margin - x):
                    draw.text((x, y + 3), wrapped, fill=color, font=text_font)
                    y += 23
            y += 7
        else:
            if block["caption"]:
                draw.rectangle((margin, y, width - margin, y + 32), fill="#F8FAFC", outline="#CBD5E1")
                draw.text((margin + 10, y + 6), str(block["caption"]), fill="#111827", font=bold)
                y += 32
            header, column_widths = block["header"], widths(block["header"])
            x = margin
            for index, value in enumerate(header):
                right = x + column_widths[index]
                draw.rectangle((x, y, right, y + 34), fill="#E5E7EB", outline="#CBD5E1")
                draw.text((x + 6, y + 7), value, fill="#111827", font=bold)
                x = right
            y += 34
            for row_index, row in enumerate(block["rows"]):
                cells = [wrap(str(value), font, max(20, column_widths[index] - 12)) for index, value in enumerate(row)]
                row_height = max(31, 8 + max(len(cell) for cell in cells) * 22)
                if any(
                    _is_stacked_start_due_cell(block["header"][index], str(value))
                    for index, value in enumerate(row)
                ):
                    row_height = max(row_height, 60)
                tone = block["row_tones"][row_index] if row_index < len(block["row_tones"]) else block["tone"]
                fill, color = _table_tone_styles(tone)
                highlighted = block["highlights"][row_index] if row_index < len(block["highlights"]) else False
                x = margin
                for index, value in enumerate(row):
                    header_name = _normalized_table_header(header[index])
                    cell_fill, cell_color, cell_font = fill, color, font
                    override = _table_cell_style_override(header_name, str(value))
                    if override is not None:
                        cell_fill, cell_color, override_bold = override
                        cell_font = bold if override_bold else font
                    elif header_name == "DISK" and str(value).strip().upper() == "YES":
                        cell_fill, cell_color, cell_font = "#DCFCE7", "#166534", bold
                    elif header_name == "DISK" and str(value).strip().upper() == "NO":
                        cell_fill, cell_color, cell_font = "#FEE2E2", "#991B1B", bold
                    elif header_name == "MBAJTUR?" and str(value).strip() == "✓":
                        cell_fill, cell_color, cell_font = "#DCFCE7", "#166534", bold
                    elif header_name == "MBAJTUR?" and str(value).strip() == "✕":
                        cell_fill, cell_color, cell_font = "#FEE2E2", "#991B1B", bold
                    if highlighted and header_name == "TITLE":
                        cell_color, cell_font = "#2563EB", bold
                    right = x + column_widths[index]
                    outlined = highlighted or tone == "eight-am"
                    outline = "#DC2626" if tone == "eight-am" else ("#2563EB" if highlighted else "#CBD5E1")
                    draw.rectangle((x, y, right, y + row_height), fill=cell_fill, outline=outline, width=3 if outlined else 1)
                    if _is_stacked_start_due_cell(header[index], str(value)):
                        start_line, due_line = str(value).split("\n", 1)
                        draw.text((x + 6, y + 5), start_line, fill=cell_color, font=cell_font)
                        strong_start_due_divider = _is_m3_start_due_table(str(block["caption"] or ""))
                        draw.line(
                            (x, y + 30, right, y + 30),
                            fill="#334155" if strong_start_due_divider else "#94A3B8",
                            width=3 if strong_start_due_divider else 1,
                        )
                        draw.text((x + 6, y + 34), due_line, fill=cell_color, font=cell_font)
                    else:
                        for line_index, text in enumerate(cells[index]):
                            draw.text((x + 6, y + 5 + line_index * 22), text, fill=cell_color, font=cell_font)
                    x = right
                y += row_height
            y += 14
    output = __import__("io").BytesIO()
    image.save(output, format="PNG", optimize=True)
    return output.getvalue()


def section_report_attachments(
    subject: str,
    report_code: str,
    report_day: date,
    sections: list[dict[str, str]],
    *,
    tomorrow: date | None = None,
) -> list[tuple[str, bytes, str]]:
    filename = f"PrimeFlow-{report_code}-{report_day:%Y-%m-%d}"
    return [
        (
            f"{filename}.docx",
            render_section_report_docx(subject, report_code, report_day, sections, tomorrow=tomorrow),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        (f"{filename}.png", render_section_report_png(subject, report_code, report_day, sections, tomorrow=tomorrow), "image/png"),
    ]


async def send_section_report(
    subject: str,
    recipients: dict[str, list[str]],
    plain_text: str,
    html_body: str,
    *,
    report_code: str,
    report_day: date,
    sections: list[dict[str, str]],
    tomorrow: date | None = None,
) -> dict[str, Any]:
    gmail = GmailService()
    attachments = section_report_attachments(subject, report_code, report_day, sections, tomorrow=tomorrow)
    return await gmail.send_verified(subject, recipients, plain_text, html_body, attachments=attachments)


async def send_meetings_report(
    subject: str,
    recipients: dict[str, list[str]],
    plain_text: str,
    html_body: str,
    *,
    report_day: date,
    tomorrow: date,
    sections: list[dict[str, str]],
) -> dict[str, Any]:
    return await send_section_report(
        subject,
        recipients,
        plain_text,
        html_body,
        report_code="M3",
        report_day=report_day,
        tomorrow=tomorrow,
        sections=sections,
    )
