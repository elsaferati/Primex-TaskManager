from __future__ import annotations

import hashlib
import html
import io
import json
import logging
import os
import re
import uuid
from collections import defaultdict
from datetime import date, datetime, time, timedelta
from pathlib import Path
from typing import Any, Literal
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from pydantic import BaseModel, Field
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import landscape, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import LongTable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy import select, text
from sqlalchemy.ext.asyncio import AsyncSession

from app.db import SessionLocal
from app.models.department import Department
from app.models.enums import GaNoteStatus
from app.models.plan_note import PlanNote
from app.models.primeflow_report_delivery_run import PrimeFlowReportDeliveryRun
from app.models.primeflow_report_snapshot import PrimeFlowReportSnapshot
from app.models.project import Project
from app.models.task import Task
from app.models.task_assignee import TaskAssignee
from app.models.user import User
from app.services.primeflow_report import GmailService, GmailVerificationError, clean_description


logger = logging.getLogger(__name__)
REPORT_TYPE = "px_jav_weekly"
REPORT_SLOT = "15:50"
DEFAULT_RECIPIENT = "334primex.eu@gmail.com"
TERMINAL_STATUSES = {"SENT", "ALREADY_SENT"}
EXCEL_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
WORD_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME = "application/pdf"


class PxJavNoteRow(BaseModel):
    note_id: str
    number: int
    content: str
    comment: str = ""
    note_status: str
    priority: str = "-"
    discussed: bool
    next_week: bool
    created_at: datetime
    created_by: str
    department: str
    project: str
    result: Literal["DETYRË", "VETËM SHËNIM"]
    task_count: int = 0
    active_task_count: int = 0
    assignees: list[str] = Field(default_factory=list)
    task_statuses: list[str] = Field(default_factory=list)
    task_created_dates: list[date] = Field(default_factory=list)
    task_start_dates: list[date] = Field(default_factory=list)
    task_due_dates: list[date] = Field(default_factory=list)
    year_end_task: bool = False
    year_end_task_count: int = 0
    year_end_comment: bool = False


class PxJavWeeklyReport(BaseModel):
    report_date: date
    generated_at: datetime
    period_start: datetime
    period_end: datetime
    timezone: str
    recipient: str
    source_note_count: int = 0
    period_note_count: int = 0
    commented_note_count: int = 0
    year_end_comment_count: int = 0
    excluded_task_count: int = 0
    excluded_next_week_count: int = 0
    rows: list[PxJavNoteRow]

    @property
    def total_notes(self) -> int:
        return len(self.rows)

    @property
    def notes_without_task(self) -> int:
        return sum(row.task_count == 0 for row in self.rows)

    @property
    def next_week_tasks(self) -> int:
        return sum(row.next_week and row.task_count > 0 for row in self.rows)

    @property
    def next_week_without_task(self) -> int:
        return sum(row.next_week and row.task_count == 0 for row in self.rows)

    @property
    def note_only(self) -> int:
        return sum(row.result == "VETËM SHËNIM" for row in self.rows)

    def summary(self) -> dict[str, int]:
        return {
            "period_notes": self.period_note_count,
            "commented_notes": self.commented_note_count,
            "year_end_comments": self.year_end_comment_count,
            "report_notes": self.total_notes,
            "notes_without_task": self.notes_without_task,
            "next_week_tasks": self.next_week_tasks,
            "note_only": self.note_only,
            "excluded_with_task": self.excluded_task_count,
            "next_week_without_task": self.next_week_without_task,
        }


def report_timezone() -> ZoneInfo:
    return ZoneInfo(os.getenv("PX_JAV_WEEKLY_REPORT_TIMEZONE", "Europe/Tirane"))


def configured_recipient() -> str:
    return os.getenv("PX_JAV_WEEKLY_REPORT_RECIPIENT", DEFAULT_RECIPIENT).strip() or DEFAULT_RECIPIENT


def report_subject(day: date) -> str:
    return f"Raporti PX JAV - kontrolli i taskave | {day:%d.%m.%Y} | {REPORT_SLOT}"


def report_filename_stem(report: PxJavWeeklyReport) -> str:
    return f"Raporti_PX_JAV_{report.report_date:%d-%m-%Y}_{REPORT_SLOT.replace(':', '-')}"


def _enum_text(value: Any, default: str = "-") -> str:
    if value is None:
        return default
    raw = getattr(value, "value", value)
    return str(raw or default)


def _display_user(user: User | None) -> str:
    if user is None:
        return "-"
    return (user.full_name or user.username or user.email or "-").strip()


def _initials(value: str) -> str:
    parts = [part for part in value.replace("-", " ").split() if part]
    return "".join(part[0] for part in parts[:3]).upper() or "-"


def _unique_dates(values: list[datetime | None]) -> list[date]:
    return sorted({value.date() for value in values if value is not None})


def _date_text(values: list[date], *, mark_year_end: bool = False) -> str:
    output: list[str] = []
    for value in values:
        rendered = value.strftime("%d.%m.%Y")
        if mark_year_end and value.day == 31 and value.month == 12:
            rendered += " (FUNDVIT)"
        output.append(rendered)
    return "; ".join(output) or "-"


def _clean_note_text(value: str | None) -> str:
    return clean_description(value).strip() or "-"


def _is_year_end_comment(value: str | None) -> bool:
    return bool(re.search(r"(?<!\d)31[.\/-]12(?:[.\/-]\d{2,4})?(?!\d)", value or ""))


def _in_timezone(value: datetime, timezone: ZoneInfo) -> datetime:
    if value.tzinfo is None:
        value = value.replace(tzinfo=ZoneInfo("UTC"))
    return value.astimezone(timezone)


def classify_note_result(
    task_count: int,
    converted_flag: bool,
) -> Literal["DETYRË", "VETËM SHËNIM"]:
    """Use only actual linked tasks as truth; the stored converted flag is informational."""
    if task_count:
        return "DETYRË"
    return "VETËM SHËNIM"


def _review_label(row: PxJavNoteRow) -> str:
    if row.year_end_comment:
        return "31.12 / TASK" if row.task_count > 0 else "31.12 / PA TASK"
    if row.next_week and row.task_count > 0:
        return "TASK PËR J.T"
    if row.task_count > 0:
        return "TASK I KRIJUAR"
    return "PA TASK"


def previous_thursday_cutoff(value: datetime) -> datetime:
    """Return the prior calendar Thursday at the weekly report time."""
    days_back = (value.weekday() - 3) % 7 or 7
    previous_thursday = value.date() - timedelta(days=days_back)
    return datetime.combine(previous_thursday, time(15, 50), tzinfo=value.tzinfo)


async def build_px_jav_weekly_report(
    db: AsyncSession,
    *,
    report_date: date | None = None,
    as_of: datetime | None = None,
    timezone_name: str | None = None,
    recipient: str | None = None,
) -> PxJavWeeklyReport:
    timezone_name = timezone_name or report_timezone().key
    timezone = ZoneInfo(timezone_name)
    generated_at = datetime.now(timezone)
    if as_of is not None:
        period_end = as_of.astimezone(timezone)
    elif report_date is not None:
        period_end = datetime.combine(report_date, time(15, 50), tzinfo=timezone)
    else:
        period_end = generated_at
    report_date = period_end.date()
    period_start = previous_thursday_cutoff(period_end)

    notes = list((
        await db.execute(
            select(PlanNote)
            .where(PlanNote.status != GaNoteStatus.CLOSED)
            .order_by(PlanNote.created_at.desc(), PlanNote.id.desc())
        )
    ).scalars().all())
    notes.sort(key=lambda note: _is_year_end_comment(note.comment))
    note_ids = [note.id for note in notes]

    tasks: list[Task] = []
    if note_ids:
        tasks = (
            await db.execute(
                select(Task)
                .where(Task.plan_note_origin_id.in_(note_ids))
                .order_by(Task.created_at, Task.id)
            )
        ).scalars().all()

    task_ids = [task.id for task in tasks]
    explicit_assignees: dict[uuid.UUID, set[uuid.UUID]] = defaultdict(set)
    if task_ids:
        for task_id, user_id in (
            await db.execute(
                select(TaskAssignee.task_id, TaskAssignee.user_id).where(TaskAssignee.task_id.in_(task_ids))
            )
        ).all():
            explicit_assignees[task_id].add(user_id)

    user_ids = {note.created_by for note in notes if note.created_by}
    user_ids.update(task.assigned_to for task in tasks if task.assigned_to)
    user_ids.update(user_id for values in explicit_assignees.values() for user_id in values)
    users: dict[uuid.UUID, User] = {}
    if user_ids:
        users = {
            user.id: user
            for user in (await db.execute(select(User).where(User.id.in_(user_ids)))).scalars().all()
        }

    department_ids = {note.department_id for note in notes if note.department_id}
    department_ids.update(task.department_id for task in tasks if task.department_id)
    departments: dict[uuid.UUID, Department] = {}
    if department_ids:
        departments = {
            item.id: item
            for item in (
                await db.execute(select(Department).where(Department.id.in_(department_ids)))
            ).scalars().all()
        }

    project_ids = {note.project_id for note in notes if note.project_id}
    project_ids.update(task.project_id for task in tasks if task.project_id)
    projects: dict[uuid.UUID, Project] = {}
    if project_ids:
        projects = {
            item.id: item
            for item in (await db.execute(select(Project).where(Project.id.in_(project_ids)))).scalars().all()
        }

    tasks_by_note: dict[uuid.UUID, list[Task]] = defaultdict(list)
    for task in tasks:
        if task.plan_note_origin_id:
            tasks_by_note[task.plan_note_origin_id].append(task)

    rows: list[PxJavNoteRow] = []
    excluded_task_count = 0
    for note in notes:
        linked_tasks = tasks_by_note.get(note.id, [])
        year_end_comment = _is_year_end_comment(note.comment)
        if linked_tasks and not note.next_week:
            excluded_task_count += 1
            continue

        assignee_ids: set[uuid.UUID] = set()
        for task in linked_tasks:
            if task.assigned_to:
                assignee_ids.add(task.assigned_to)
            assignee_ids.update(explicit_assignees.get(task.id, set()))
        assignees = sorted({_display_user(users.get(user_id)) for user_id in assignee_ids})
        due_dates = _unique_dates([task.due_date for task in linked_tasks])

        result = classify_note_result(len(linked_tasks), bool(note.is_converted_to_task))

        note_department_id = note.department_id or next(
            (task.department_id for task in linked_tasks if task.department_id), None
        )
        note_project_id = note.project_id or next(
            (task.project_id for task in linked_tasks if task.project_id), None
        )
        department = departments.get(note_department_id) if note_department_id else None
        project = projects.get(note_project_id) if note_project_id else None
        rows.append(PxJavNoteRow(
            note_id=str(note.id),
            number=len(rows) + 1,
            content=_clean_note_text(note.content),
            comment=_clean_note_text(note.comment) if note.comment else "",
            note_status=_enum_text(note.status),
            priority=_enum_text(note.priority),
            discussed=bool(note.is_discussed),
            next_week=bool(note.next_week),
            created_at=_in_timezone(note.created_at, timezone),
            created_by=_display_user(users.get(note.created_by)) if note.created_by else "-",
            department=(department.code or department.name) if department else "-",
            project=project.title if project else "-",
            result=result,
            task_count=len(linked_tasks),
            active_task_count=sum(bool(task.is_active) for task in linked_tasks),
            assignees=assignees,
            task_statuses=sorted({_enum_text(task.status) for task in linked_tasks}),
            task_created_dates=_unique_dates([task.created_at for task in linked_tasks]),
            task_start_dates=_unique_dates([task.start_date for task in linked_tasks]),
            task_due_dates=due_dates,
            year_end_task=any(value.day == 31 and value.month == 12 for value in due_dates),
            year_end_task_count=sum(
                bool(task.due_date and task.due_date.day == 31 and task.due_date.month == 12)
                for task in linked_tasks
            ),
            year_end_comment=year_end_comment,
        ))

    return PxJavWeeklyReport(
        report_date=report_date,
        generated_at=generated_at,
        period_start=period_start,
        period_end=period_end,
        timezone=timezone_name,
        recipient=recipient or configured_recipient(),
        source_note_count=len(notes),
        period_note_count=sum(
            period_start <= _in_timezone(note.created_at, timezone) < period_end for note in notes
        ),
        commented_note_count=sum(bool(row.comment) for row in rows),
        year_end_comment_count=sum(row.year_end_comment for row in rows),
        excluded_task_count=excluded_task_count,
        excluded_next_week_count=0,
        rows=rows,
    )


def render_plain_text(report: PxJavWeeklyReport) -> str:
    summary = report.summary()
    return (
        "Përshëndetje,\n\n"
        "Bashkëngjitur është raporti PX JAV për të gjitha shënimet pa task dhe taskat e krijuara për J.T.\n"
        f"Periudha: {report.period_start:%d.%m.%Y %H:%M} - "
        f"{report.period_end:%d.%m.%Y %H:%M}.\n\n"
        "Përmbledhje:\n"
        f"- Shënime të krijuara në periudhë: {summary['period_notes']}\n"
        f"- Shënime me koment: {summary['commented_notes']}\n"
        f"- Shënime me koment 31.12 / fundvit: {summary['year_end_comments']}\n"
        f"- Pa task (në raport): {summary['notes_without_task']}\n"
        f"- Task i krijuar për J.T (në raport): {summary['next_week_tasks']}\n"
        f"- Task i zakonshëm (përjashtuar): {summary['excluded_with_task']}\n"
        f"- J.T pa task real (në raport): {summary['next_week_without_task']}\n"
        "\n"
        "J.T me task real paraqitet si task i krijuar për javën tjetër. Komentet shfaqen kur ekzistojnë; "
        "shënimet me koment 31.12 shënohen si FUNDVIT dhe dalin në fund.\n\n"
        "Me respekt,\nPrimeFlow"
    )


def render_html(report: PxJavWeeklyReport) -> str:
    summary = report.summary()
    cards = [
        ("Në periudhë", summary["period_notes"], "#e2e8f0"),
        ("Me koment", summary["commented_notes"], "#f8fafc"),
        ("Koment 31.12", summary["year_end_comments"], "#fef3c7"),
        ("Pa task", summary["notes_without_task"], "#dbeafe"),
        ("Task për J.T", summary["next_week_tasks"], "#dcfce7"),
        ("Task normal - jashtë", summary["excluded_with_task"], "#f1f5f9"),
        ("J.T pa task", summary["next_week_without_task"], "#ede9fe"),
    ]
    card_html = "".join(
        f'<td style="background:{color};padding:12px;border:1px solid #cbd5e1;text-align:center">'
        f'<div style="font-size:22px;font-weight:700">{value}</div>'
        f'<div style="font-size:12px">{html.escape(label)}</div></td>'
        for label, value, color in cards
    )
    return (
        '<div style="font-family:Arial,sans-serif;color:#0f172a">'
        f'<h2 style="margin-bottom:4px">Raporti PX JAV - kontrolli i taskave</h2>'
        f'<p style="margin-top:0;color:#475569">Periudha: {report.period_start:%d.%m.%Y %H:%M} - '
        f'{report.period_end:%d.%m.%Y %H:%M} | {html.escape(report.timezone)}</p>'
        f'<table role="presentation" cellspacing="0" cellpadding="0"><tr>{card_html}</tr></table>'
        '<p>Raporti liston të gjitha shënimet pa task dhe shënimet J.T që kanë task real të lidhur.</p>'
        '<p><strong>J.T:</strong> shfaqet si “Task për J.T” vetëm kur ekziston tasku real i lidhur.</p>'
        '<p><strong>31.12:</strong> shfaqet pavarësisht datës së krijimit dhe shënohet si fundvit.</p>'
        '<p>Bashkëngjitur: Excel, Word dhe PDF.</p>'
        '</div>'
    )


def render_xlsx(report: PxJavWeeklyReport) -> bytes:
    workbook = Workbook()
    summary_sheet = workbook.active
    summary_sheet.title = "PËRMBLEDHJE"
    detail_sheet = workbook.create_sheet("KONTROLLI PX JAV")

    summary_sheet.append(["RAPORTI PX JAV", report.report_date])
    summary_sheet.append(["Gjeneruar", report.generated_at.replace(tzinfo=None)])
    summary_sheet.append(["Nga", report.period_start.replace(tzinfo=None)])
    summary_sheet.append(["Deri", report.period_end.replace(tzinfo=None)])
    summary_sheet.append(["Timezone", report.timezone])
    summary_sheet.append(["Marrësi", report.recipient])
    summary_sheet.append([])
    summary_sheet.append(["Treguesi", "Vlera"])
    labels = [
        ("Shënime të periudhës", report.period_note_count),
        ("Shënime me koment", report.commented_note_count),
        ("Koment 31.12 / fundvit", report.year_end_comment_count),
        ("Rreshta në raport", report.total_notes),
        ("Pa task - në raport", report.notes_without_task),
        ("Task i krijuar për J.T - në raport", report.next_week_tasks),
        ("Task i zakonshëm - përjashtuar", report.excluded_task_count),
        ("J.T pa task real - në raport", report.next_week_without_task),
    ]
    for label, value in labels:
        summary_sheet.append([label, value])

    headers = [
        "NR", "KONTROLLI", "SHËNIMI", "KOMENT", "STATUS SHËNIMI", "PRIORITET",
        "DISK", "J.T", "31.12", "DATA/ORA", "NGA", "DEPARTAMENTI", "PROJEKTI",
        "TASK / PËR", "STATUS TASK", "DEADLINE", "NOTE ID",
    ]
    detail_sheet.append(headers)
    for row in report.rows:
        detail_sheet.append([
            row.number, _review_label(row), row.content, row.comment, row.note_status, row.priority,
            "YES" if row.discussed else "NO", "YES" if row.next_week else "NO",
            "FUNDVIT" if row.year_end_comment else "NO",
            row.created_at.replace(tzinfo=None), row.created_by, row.department, row.project,
            "; ".join(row.assignees) or "-", "; ".join(row.task_statuses) or "-",
            _date_text(row.task_due_dates, mark_year_end=True), row.note_id,
        ])

    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    for sheet, header_row in ((summary_sheet, 8), (detail_sheet, 1)):
        for cell in sheet[header_row]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        sheet.freeze_panes = f"A{header_row + 1}"
        sheet.auto_filter.ref = sheet.dimensions

    summary_sheet["B1"].number_format = "dd.mm.yyyy"
    summary_sheet["B2"].number_format = "dd.mm.yyyy hh:mm"
    summary_sheet["B3"].number_format = "dd.mm.yyyy hh:mm"
    summary_sheet["B4"].number_format = "dd.mm.yyyy hh:mm"
    summary_sheet.column_dimensions["A"].width = 28
    summary_sheet.column_dimensions["B"].width = 30

    detail_sheet.column_dimensions["A"].width = 6
    detail_sheet.column_dimensions["B"].width = 18
    detail_sheet.column_dimensions["C"].width = 55
    detail_sheet.column_dimensions["D"].width = 35
    for column in range(5, len(headers) + 1):
        detail_sheet.column_dimensions[get_column_letter(column)].width = 16
    detail_sheet.column_dimensions["N"].width = 30
    detail_sheet.column_dimensions["Q"].width = 38
    result_fills = {
        "PA TASK": PatternFill("solid", fgColor="DBEAFE"),
        "TASK PËR J.T": PatternFill("solid", fgColor="DCFCE7"),
        "TASK I KRIJUAR": PatternFill("solid", fgColor="EDE9FE"),
        "31.12 / PA TASK": PatternFill("solid", fgColor="FEF3C7"),
        "31.12 / TASK": PatternFill("solid", fgColor="FDE68A"),
    }
    for current_row in detail_sheet.iter_rows(min_row=2):
        for cell in current_row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
        current_row[1].fill = result_fills.get(str(current_row[1].value), PatternFill())
        current_row[9].number_format = "dd.mm.yyyy hh:mm"

    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    verified = load_workbook(output, read_only=True, data_only=False)
    if verified.sheetnames != ["PËRMBLEDHJE", "KONTROLLI PX JAV"]:
        raise ValueError("Invalid PX JAV workbook layout")
    if verified["KONTROLLI PX JAV"].max_row != report.total_notes + 1:
        raise ValueError("Invalid PX JAV workbook row count")
    verified.close()
    return output.getvalue()


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int = 60, start: int = 60, bottom: int = 60, end: int = 60) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _set_table_geometry(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    total_width = int(sum(widths) * 1440)
    table_width = tbl_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        tbl_pr.append(table_width)
    table_width.set(qn("w:w"), str(total_width))
    table_width.set(qn("w:type"), "dxa")
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "60")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _set_run_font(run, *, size: float, bold: bool = False, color: str = "111827") -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _write_docx_cell(cell, value: str, *, size: float = 6.5, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    _set_run_font(paragraph.add_run(value), size=size, bold=bold)


def render_docx(report: PxJavWeeklyReport) -> bytes:
    document = Document()
    # Named layout override for this dense operational report: the
    # compact-reference visual system is kept, with Letter landscape geometry.
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(8)
    normal.paragraph_format.space_after = Pt(3)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(header.add_run("PrimeFlow | Raporti PX JAV"), size=7, color="64748B")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(footer.add_run(f"Gjeneruar {report.generated_at:%d.%m.%Y %H:%M} | {report.timezone}"), size=7, color="64748B")

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_font(title.add_run("RAPORTI PX JAV"), size=18, bold=True, color="1F4E78")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(7)
    _set_run_font(
        subtitle.add_run(
            f"Të gjitha shënimet pa task + taskat për J.T | {report.period_start:%d.%m.%Y %H:%M} - "
            f"{report.period_end:%d.%m.%Y %H:%M}"
        ),
        size=10,
        color="475569",
    )

    summary_values = [
        ("NË PERIUDHË", report.period_note_count, "E2E8F0"),
        ("ME KOMENT", report.commented_note_count, "F8FAFC"),
        ("KOMENT 31.12", report.year_end_comment_count, "FEF3C7"),
        ("PA TASK", report.notes_without_task, "DBEAFE"),
        ("TASK PËR J.T", report.next_week_tasks, "DCFCE7"),
        ("TASK NORMAL - JASHTË", report.excluded_task_count, "F1F5F9"),
        ("J.T PA TASK", report.next_week_without_task, "EDE9FE"),
    ]
    summary_table = document.add_table(rows=1, cols=len(summary_values))
    summary_widths = [10.3 / len(summary_values)] * len(summary_values)
    _set_table_geometry(summary_table, summary_widths)
    for cell, (label, value, fill) in zip(summary_table.rows[0].cells, summary_values):
        _set_cell_shading(cell, fill)
        _write_docx_cell(cell, f"{value}\n{label}", size=8, bold=True, center=True)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(5)
    note.paragraph_format.space_after = Pt(5)
    _set_run_font(
        note.add_run(
            "Listohen të gjitha shënimet pa task dhe taskat J.T; komentet shfaqen kur ekzistojnë dhe 31.12 del në fund."
        ),
        size=7.5,
        bold=True,
        color="7A5A00",
    )

    headers = ["NR", "KONTROLLI", "SHËNIMI / KOMENT", "ST / PRIO", "DISK / J.T / 31.12", "DATA / NGA", "DEP / PRJK", "TASK / PËR", "STATUS TASK", "DEADLINE"]
    widths = [0.28, 0.9, 3.4, 0.55, 0.65, 1.0, 1.0, 0.9, 0.7, 0.92]
    detail = document.add_table(rows=1, cols=len(headers))
    detail.style = "Table Grid"
    _set_table_geometry(detail, widths)
    _set_repeat_table_header(detail.rows[0])
    for cell, value in zip(detail.rows[0].cells, headers):
        _set_cell_shading(cell, "1F4E78")
        _write_docx_cell(cell, value, size=7, bold=True, center=True)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    result_fills = {
        "PA TASK": "DBEAFE",
        "TASK PËR J.T": "DCFCE7",
        "TASK I KRIJUAR": "EDE9FE",
        "31.12 / PA TASK": "FEF3C7",
        "31.12 / TASK": "FDE68A",
    }
    for item in report.rows:
        row = detail.add_row()
        content = item.content + (f"\n\nKoment: {item.comment}" if item.comment else "")
        values = [
            str(item.number), _review_label(item), content, f"{item.note_status}\n{item.priority}",
            f"D:{'YES' if item.discussed else 'NO'}\nJ.T:{'YES' if item.next_week else 'NO'}\n"
            f"31.12:{'YES' if item.year_end_comment else 'NO'}",
            f"{item.created_at:%d.%m.%Y %H:%M}\n{_initials(item.created_by)}",
            f"{item.department}\n{item.project}",
            "; ".join(item.assignees) or "-",
            "; ".join(item.task_statuses) or "-",
            _date_text(item.task_due_dates, mark_year_end=True),
        ]
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            _write_docx_cell(cell, value, size=6.3, bold=index == 1, center=index in {0, 1, 3, 4, 8, 9})
        _set_cell_shading(row.cells[1], result_fills[_review_label(item)])

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _register_pdf_fonts() -> tuple[str, str]:
    candidates = [
        (Path("C:/Windows/Fonts/arial.ttf"), Path("C:/Windows/Fonts/arialbd.ttf")),
        (Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"), Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")),
    ]
    for regular_path, bold_path in candidates:
        if regular_path.is_file() and bold_path.is_file():
            try:
                pdfmetrics.registerFont(TTFont("PXJavRegular", str(regular_path)))
                pdfmetrics.registerFont(TTFont("PXJavBold", str(bold_path)))
                return "PXJavRegular", "PXJavBold"
            except Exception:
                logger.exception("Could not register PDF font %s", regular_path)
    return "Helvetica", "Helvetica-Bold"


def render_pdf(report: PxJavWeeklyReport) -> bytes:
    regular_font, bold_font = _register_pdf_fonts()
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=landscape(letter),
        leftMargin=0.28 * inch,
        rightMargin=0.28 * inch,
        topMargin=0.35 * inch,
        bottomMargin=0.38 * inch,
        title=f"Raporti PX JAV {report.report_date:%d.%m.%Y}",
        author="PrimeFlow",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "PXTitle", parent=styles["Title"], fontName=bold_font, fontSize=18,
        leading=20, textColor=colors.HexColor("#1F4E78"), alignment=TA_LEFT, spaceAfter=4,
    )
    meta_style = ParagraphStyle(
        "PXMeta", parent=styles["Normal"], fontName=regular_font, fontSize=8.5,
        leading=10, textColor=colors.HexColor("#475569"), spaceAfter=7,
    )
    cell_style = ParagraphStyle(
        "PXCell", parent=styles["Normal"], fontName=regular_font, fontSize=5.8,
        leading=7, textColor=colors.HexColor("#111827"), wordWrap="CJK",
    )
    cell_bold = ParagraphStyle(
        "PXCellBold", parent=cell_style, fontName=bold_font, fontSize=5.8, leading=7,
    )
    header_style = ParagraphStyle(
        "PXHeader", parent=cell_style, fontName=bold_font, fontSize=6, leading=7,
        textColor=colors.white, alignment=1,
    )
    story: list[Any] = [
        Paragraph("RAPORTI PX JAV - KONTROLLI I TASKAVE", title_style),
        Paragraph(
            f"Periudha: {report.period_start:%d.%m.%Y %H:%M} - "
            f"{report.period_end:%d.%m.%Y %H:%M} | {html.escape(report.timezone)}",
            meta_style,
        ),
    ]
    summary_data = [[
        Paragraph(f"<b>{value}</b><br/>{html.escape(label)}", cell_style)
        for label, value in [
            ("NË PERIUDHË", report.period_note_count),
            ("ME KOMENT", report.commented_note_count),
            ("KOMENT 31.12", report.year_end_comment_count),
            ("PA TASK", report.notes_without_task),
            ("TASK PËR J.T", report.next_week_tasks),
            ("TASK NORMAL - JASHTË", report.excluded_task_count),
            ("J.T PA TASK", report.next_week_without_task),
        ]
    ]]
    summary = Table(summary_data, colWidths=[1.44 * inch] * 7)
    summary.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#E2E8F0")),
        ("BACKGROUND", (1, 0), (1, 0), colors.HexColor("#F8FAFC")),
        ("BACKGROUND", (2, 0), (2, 0), colors.HexColor("#FEF3C7")),
        ("BACKGROUND", (3, 0), (3, 0), colors.HexColor("#DBEAFE")),
        ("BACKGROUND", (4, 0), (4, 0), colors.HexColor("#DCFCE7")),
        ("BACKGROUND", (5, 0), (5, 0), colors.HexColor("#F1F5F9")),
        ("BACKGROUND", (6, 0), (6, 0), colors.HexColor("#EDE9FE")),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.extend([
        summary,
        Spacer(1, 5),
        Paragraph(
            "Listohen të gjitha shënimet pa task dhe taskat J.T; komentet shfaqen kur ekzistojnë dhe 31.12 del në fund.",
            meta_style,
        ),
    ])

    headers = ["NR", "KONTROLLI", "SHËNIMI / KOMENT", "ST / PRIO", "DISK / J.T / 31.12", "DATA / NGA", "DEP / PRJK", "TASK / PËR", "STATUS TASK", "DEADLINE"]
    data: list[list[Any]] = [[Paragraph(html.escape(value), header_style) for value in headers]]
    result_rows: dict[str, list[int]] = defaultdict(list)
    for item in report.rows:
        content = html.escape(item.content).replace("\n", "<br/>")
        if item.comment:
            content += f"<br/><br/><b>Koment:</b> {html.escape(item.comment).replace(chr(10), '<br/>')}"
        row_index = len(data)
        result_rows[_review_label(item)].append(row_index)
        data.append([
            Paragraph(str(item.number), cell_style),
            Paragraph(html.escape(_review_label(item)), cell_bold),
            Paragraph(content, cell_style),
            Paragraph(f"{html.escape(item.note_status)}<br/>{html.escape(item.priority)}", cell_style),
            Paragraph(
                f"D:{'YES' if item.discussed else 'NO'}<br/>"
                f"J.T:{'YES' if item.next_week else 'NO'}<br/>"
                f"31.12:{'YES' if item.year_end_comment else 'NO'}",
                cell_style,
            ),
            Paragraph(f"{item.created_at:%d.%m.%Y %H:%M}<br/>{html.escape(_initials(item.created_by))}", cell_style),
            Paragraph(f"{html.escape(item.department)}<br/>{html.escape(item.project)}", cell_style),
            Paragraph(html.escape("; ".join(item.assignees) or "-"), cell_style),
            Paragraph(html.escape("; ".join(item.task_statuses) or "-"), cell_style),
            Paragraph(html.escape(_date_text(item.task_due_dates, mark_year_end=True)), cell_style),
        ])
    widths = [0.28, 0.9, 3.4, 0.55, 0.65, 1.0, 1.0, 0.9, 0.7, 0.92]
    detail = LongTable(data, colWidths=[width * inch for width in widths], repeatRows=1, splitByRow=1)
    commands: list[tuple] = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#94A3B8")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 2), ("RIGHTPADDING", (0, 0), (-1, -1), 2),
        ("TOPPADDING", (0, 0), (-1, -1), 2.5), ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
        ("ALIGN", (0, 1), (1, -1), "CENTER"), ("ALIGN", (3, 1), (5, -1), "CENTER"),
    ]
    fills = {
        "PA TASK": "#DBEAFE",
        "TASK PËR J.T": "#DCFCE7",
        "TASK I KRIJUAR": "#EDE9FE",
        "31.12 / PA TASK": "#FEF3C7",
        "31.12 / TASK": "#FDE68A",
    }
    for result, indices in result_rows.items():
        for index in indices:
            commands.append(("BACKGROUND", (1, index), (1, index), colors.HexColor(fills[result])))
    detail.setStyle(TableStyle(commands))
    story.append(detail)

    def draw_footer(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont(regular_font, 7)
        canvas.setFillColor(colors.HexColor("#64748B"))
        canvas.drawRightString(10.7 * inch, 0.18 * inch, f"PrimeFlow | PX JAV | Faqe {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)
    return output.getvalue()


def build_attachments(report: PxJavWeeklyReport) -> list[tuple[str, bytes, str]]:
    stem = report_filename_stem(report)
    return [
        (f"{stem}.xlsx", render_xlsx(report), EXCEL_MIME),
        (f"{stem}.docx", render_docx(report), WORD_MIME),
        (f"{stem}.pdf", render_pdf(report), PDF_MIME),
    ]


async def deliver_px_jav_weekly_report(
    report_date: date | None = None,
    *,
    send: bool = True,
    trigger_type: str = "SCHEDULED",
    triggered_by_user_id: uuid.UUID | None = None,
    recipient_group: str = "default",
) -> PrimeFlowReportDeliveryRun:
    timezone = report_timezone()
    now = datetime.now(timezone)
    report_date = report_date or now.date()
    recipient = configured_recipient()
    recipient_map = {"to": [recipient], "cc": [], "bcc": []}
    recipients = [recipient]
    subject = report_subject(report_date)
    if trigger_type == "MANUAL" and recipient_group == "default":
        recipient_group = f"manual-{uuid.uuid4().hex}"

    async with SessionLocal() as db:
        async with db.begin():
            lock_key = f"{REPORT_TYPE}|{report_date.isoformat()}|{REPORT_SLOT}|{recipient_group}"
            await db.execute(text("SELECT pg_advisory_xact_lock(hashtext(:key))"), {"key": lock_key})
            run = (
                await db.execute(
                    select(PrimeFlowReportDeliveryRun)
                    .where(
                        PrimeFlowReportDeliveryRun.report_type == REPORT_TYPE,
                        PrimeFlowReportDeliveryRun.report_date == report_date,
                        PrimeFlowReportDeliveryRun.report_slot == REPORT_SLOT,
                        PrimeFlowReportDeliveryRun.recipient_group == recipient_group,
                    )
                    .with_for_update()
                )
            ).scalar_one_or_none()
            if run is None:
                run = PrimeFlowReportDeliveryRun(
                    report_type=REPORT_TYPE,
                    report_date=report_date,
                    report_slot=REPORT_SLOT,
                    recipient_group=recipient_group,
                    scheduled_for=now,
                    scheduled_execution_time=now,
                    subject=subject,
                    recipients=json.dumps(recipients),
                    status="PENDING",
                    trigger_type=trigger_type,
                    triggered_by_user_id=triggered_by_user_id,
                )
                db.add(run)
                await db.flush()
            if run.status in TERMINAL_STATUSES:
                return run
            if run.status == "RUNNING" and run.started_at and run.started_at > now - timedelta(minutes=30):
                return run
            run.status = "RUNNING"
            run.started_at = now
            run.attempt_count += 1
            run.error_code = None
            run.error_message = None

        try:
            report = await build_px_jav_weekly_report(
                db, report_date=report_date, timezone_name=timezone.key, recipient=recipient
            )
            body = render_plain_text(report)
            html_body = render_html(report)
            run.body_hash = hashlib.sha256(body.encode("utf-8")).hexdigest()
            run.data_generated_at = report.generated_at
            snapshot = (
                await db.execute(
                    select(PrimeFlowReportSnapshot).where(PrimeFlowReportSnapshot.delivery_run_id == run.id)
                )
            ).scalar_one_or_none()
            payload = {**report.model_dump(mode="json"), "summary": report.summary()}
            if snapshot is None:
                snapshot = PrimeFlowReportSnapshot(
                    delivery_run_id=run.id,
                    normalized_report_json=payload,
                    plain_text_body=body,
                    html_body=html_body,
                    content_version=1,
                )
                db.add(snapshot)
            else:
                snapshot.normalized_report_json = payload
                snapshot.plain_text_body = body
                snapshot.html_body = html_body
                snapshot.content_version += 1

            if not send:
                run.status = "PENDING"
                run.finished_at = datetime.now(timezone)
                await db.commit()
                setattr(run, "dry_run_report", report)
                return run

            gmail = GmailService()
            message = await gmail.send_verified(
                subject,
                recipient_map,
                body,
                html_body,
                attachments=build_attachments(report),
            )
            run.status = "SENT"
            run.gmail_message_id = message.get("id")
            run.gmail_thread_id = message.get("threadId")
        except ValueError as exc:
            run.status, run.error_code, run.error_message = "FAILED_DATA", type(exc).__name__, str(exc)[:2000]
        except GmailVerificationError as exc:
            run.status, run.error_code, run.error_message = "FAILED_VERIFICATION", type(exc).__name__, str(exc)[:2000]
            run.gmail_message_id = exc.response.get("id")
            run.gmail_thread_id = exc.response.get("threadId")
        except Exception as exc:
            logger.exception("px_jav_weekly_report_failed report_date=%s", report_date)
            run.status, run.error_code, run.error_message = "FAILED_EMAIL", type(exc).__name__, str(exc)[:2000]
        run.finished_at = datetime.now(timezone)
        await db.commit()
        return run
