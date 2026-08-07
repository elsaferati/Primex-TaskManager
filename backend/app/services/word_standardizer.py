from __future__ import annotations

import io
import hashlib
import re
import unicodedata
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Any
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Emu, Inches, Pt
from PIL import Image


KOSOVO_TIMEZONE = ZoneInfo("Europe/Belgrade")
LOGO_PATH = Path(__file__).resolve().parents[1] / "assets" / "primex_logo.png"
FILENAME_DATE = re.compile(r"(?<!\d)\d{1,2}[._-]\d{1,2}[._-]\d{2,4}(?!\d)")
COMPANY_LINES = (
    "PrimEx SH.P.K.",
    "Phone: +383 49 937 863",
    "Email: info@primex.com",
    "Website: www.primexeu.com",
)
NARROW_MARGIN = Inches(0.5)
HEADER_FOOTER_DISTANCE = Inches(0.08)


class WordStandardizationError(ValueError):
    pass


@dataclass(frozen=True)
class ComplianceCheck:
    id: str
    label: str
    compliant: bool


@dataclass(frozen=True)
class WordAnalysis:
    filename: str
    suggested_description: str
    paragraphs: int
    tables: int
    sections: int
    checks: list[ComplianceCheck]

    def to_dict(self) -> dict[str, Any]:
        payload = asdict(self)
        payload["is_compliant"] = all(check.compliant for check in self.checks)
        return payload


def _text(value: Any) -> str:
    return "" if value is None else str(value).strip()


def _normalize(value: Any) -> str:
    raw = unicodedata.normalize("NFKD", _text(value))
    ascii_value = "".join(char for char in raw if not unicodedata.combining(char))
    return re.sub(r"\s+", " ", ascii_value.casefold()).strip()


def _sanitize_description(value: str) -> str:
    clean = FILENAME_DATE.sub("", value)
    clean = re.sub(r"[^0-9A-Za-zÀ-ž]+", "_", clean.upper())
    return re.sub(r"_+", "_", clean).strip("_")[:120]


def _description_from_filename(filename: str) -> str:
    return _sanitize_description(Path(filename).stem) or "WORD_STANDARD"


def _load_document(content: bytes, filename: str):
    if Path(filename).suffix.lower() != ".docx":
        raise WordStandardizationError("Lejohet vetëm formati modern Word .docx.")
    if not content:
        raise WordStandardizationError("Dokumenti Word është bosh.")
    try:
        return Document(io.BytesIO(content))
    except Exception as exc:
        raise WordStandardizationError("Dokumenti .docx është i dëmtuar, i enkriptuar ose nuk mund të lexohet.") from exc


def _container_text(container) -> str:
    values = [paragraph.text for paragraph in container.paragraphs]
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                values.append(cell.text)
    return "\n".join(value for value in values if value)


def _field_codes(container) -> list[str]:
    return [
        _text(node.text)
        for node in container._element.iter(qn("w:instrText"))
        if _text(node.text)
    ]


def _has_field(container, field_name: str) -> bool:
    field_pattern = re.compile(rf"^\s*{re.escape(field_name)}(?:\s|\\|$)", re.IGNORECASE)
    return any(field_pattern.search(code) for code in _field_codes(container))


def _official_logo_relationship_ids(container) -> set[str]:
    official_digest = hashlib.sha256(LOGO_PATH.read_bytes()).digest() if LOGO_PATH.is_file() else b""
    return {
        relationship_id
        for relationship_id, relationship in container.part.rels.items()
        if relationship.reltype == RT.IMAGE
        and official_digest
        and hashlib.sha256(relationship.target_part.blob).digest() == official_digest
    }


def _has_official_logo(container) -> bool:
    return bool(_official_logo_relationship_ids(container))


def _official_logo_keeps_proportions(container) -> bool:
    relationship_ids = _official_logo_relationship_ids(container)
    if not relationship_ids:
        return False
    with Image.open(LOGO_PATH) as image:
        expected_ratio = image.width / image.height
    for blip in container._element.iter(qn("a:blip")):
        if blip.get(qn("r:embed")) not in relationship_ids:
            continue
        parent = blip
        while parent is not None and parent.tag != qn("wp:inline"):
            parent = parent.getparent()
        if parent is None:
            continue
        extent = parent.find(qn("wp:extent"))
        if extent is None:
            continue
        width = int(extent.get("cx", "0"))
        height = int(extent.get("cy", "0"))
        if height and abs((width / height) - expected_ratio) / expected_ratio <= 0.01:
            return True
    return False


def _company_information_present(container) -> bool:
    footer_text = _normalize(_container_text(container))
    return all(_normalize(line) in footer_text for line in COMPANY_LINES)


def _updates_fields_on_open(document) -> bool:
    update_fields = document.settings._element.find(qn("w:updateFields"))
    if update_fields is None:
        return False
    return update_fields.get(qn("w:val"), "true").casefold() not in {"0", "false", "off", "no"}


def _uses_narrow_margins(document) -> bool:
    tolerance = Inches(0.01)
    for section in document.sections:
        margins = (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin)
        if any(margin is None for margin in margins):
            return False
        if any(abs(int(margin) - int(NARROW_MARGIN)) > int(tolerance) for margin in margins):
            return False
    return True


def analyze_word_document(content: bytes, filename: str) -> WordAnalysis:
    document = _load_document(content, filename)
    sections = list(document.sections)
    header_logo = all(_has_official_logo(section.header) for section in sections)
    logo_proportions = all(_official_logo_keeps_proportions(section.header) for section in sections)
    automatic_date = all(_has_field(section.header, "DATE") for section in sections)
    footer_company = all(_company_information_present(section.footer) for section in sections)
    automatic_pages = all(
        _has_field(section.footer, "PAGE") and _has_field(section.footer, "NUMPAGES")
        for section in sections
    )
    same_first_page = all(not section.different_first_page_header_footer for section in sections)
    checks = [
        ComplianceCheck("header_logo", "Logoja zyrtare PrimEx ndodhet në header", header_logo),
        ComplianceCheck("logo_proportions", "Logoja ruan proporcionet origjinale pa shtrembërim", logo_proportions),
        ComplianceCheck("automatic_date", "Data është fushë automatike DATE (DD/MM/YYYY)", automatic_date),
        ComplianceCheck("footer_company", "Footer-i përmban informacionin zyrtar të kompanisë", footer_company),
        ComplianceCheck("automatic_pages", "Faqet përdorin fushat automatike PAGE dhe NUMPAGES", automatic_pages),
        ComplianceCheck("same_first_page", "I njëjti header/footer përdoret edhe në faqen e parë", same_first_page),
        ComplianceCheck("update_fields", "Fushat automatike përditësohen kur dokumenti hapet", _updates_fields_on_open(document)),
        ComplianceCheck("narrow_margins", "Dokumenti përdor margjina Narrow 0.5 inç", _uses_narrow_margins(document)),
    ]
    return WordAnalysis(
        filename=filename,
        suggested_description=_description_from_filename(filename),
        paragraphs=len(document.paragraphs),
        tables=len(document.tables),
        sections=len(sections),
        checks=checks,
    )


def _clear_container(container) -> None:
    for child in list(container._element):
        container._element.remove(child)


def _set_run_font(run, *, size: float, bold: bool = False, color: str = "1F2937") -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = None
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    color_element = run._element.get_or_add_rPr().find(qn("w:color"))
    if color_element is None:
        color_element = OxmlElement("w:color")
        run._element.get_or_add_rPr().append(color_element)
    color_element.set(qn("w:val"), color)


def _add_field(paragraph, instruction: str, display_text: str, *, size: float = 9, bold: bool = False) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction_element = OxmlElement("w:instrText")
    instruction_element.set(qn("xml:space"), "preserve")
    instruction_element.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    begin_run._r.extend((begin, instruction_element, separate))
    _set_run_font(begin_run, size=size, bold=bold)

    display_run = paragraph.add_run(display_text)
    _set_run_font(display_run, size=size, bold=bold)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)
    _set_run_font(end_run, size=size, bold=bold)


def _remove_table_borders(table) -> None:
    table_properties = table._tbl.tblPr
    borders = table_properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_element = borders.find(qn(f"w:{edge}"))
        if edge_element is None:
            edge_element = OxmlElement(f"w:{edge}")
            borders.append(edge_element)
        edge_element.set(qn("w:val"), "nil")


def _set_cell_margins(cell, *, top: int = 0, start: int = 0, bottom: int = 0, end: int = 0) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_table_geometry(table, total_width: int, column_widths: tuple[int, int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_properties = table._tbl.tblPr
    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(total_width))
    table_width.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in column_widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)
    for index, width in enumerate(column_widths):
        for cell in table.columns[index].cells:
            cell.width = Emu(width * 635)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell, start=45, end=45)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _remove_table_borders(table)


def _set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 0, line: float = 1) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def _create_header(section, now: datetime) -> None:
    header = section.header
    _clear_container(header)
    usable_width_emu = int(section.page_width - section.left_margin - section.right_margin)
    usable_width_twips = int(round(usable_width_emu / 635))
    left_width = int(usable_width_twips * 0.62)
    right_width = usable_width_twips - left_width
    table = header.add_table(rows=1, cols=2, width=usable_width_emu)
    _set_table_geometry(table, usable_width_twips, (left_width, right_width))

    logo_paragraph = table.cell(0, 0).paragraphs[0]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(logo_paragraph)
    logo_run = logo_paragraph.add_run()
    logo = logo_run.add_picture(str(LOGO_PATH), width=Inches(1.0))
    logo._inline.docPr.set("descr", "PrimEx company logo")

    date_paragraph = table.cell(0, 1).paragraphs[0]
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing(date_paragraph)
    _add_field(date_paragraph, 'DATE \\@ "dd/MM/yyyy"', now.strftime("%d/%m/%Y"), size=8, bold=True)


def _create_footer(section) -> None:
    footer = section.footer
    _clear_container(footer)
    usable_width_emu = int(section.page_width - section.left_margin - section.right_margin)
    usable_width_twips = int(round(usable_width_emu / 635))
    left_width = int(usable_width_twips * 0.70)
    right_width = usable_width_twips - left_width
    table = footer.add_table(rows=1, cols=2, width=usable_width_emu)
    _set_table_geometry(table, usable_width_twips, (left_width, right_width))

    company_paragraph = table.cell(0, 0).paragraphs[0]
    company_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(company_paragraph, line=0.95)
    company_rows = (
        f"{COMPANY_LINES[0]} | {COMPANY_LINES[1]}",
        f"{COMPANY_LINES[2]} | {COMPANY_LINES[3]}",
    )
    for index, line in enumerate(company_rows):
        run = company_paragraph.add_run(line)
        _set_run_font(run, size=6.5, bold=index == 0, color="334155")
        if index < len(company_rows) - 1:
            run.add_break()

    page_paragraph = table.cell(0, 1).paragraphs[0]
    page_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing(page_paragraph)
    prefix = page_paragraph.add_run("Page ")
    _set_run_font(prefix, size=7.5, bold=True)
    _add_field(page_paragraph, "PAGE", "1", size=7.5, bold=True)
    separator = page_paragraph.add_run(" of ")
    _set_run_font(separator, size=7.5, bold=True)
    _add_field(page_paragraph, "NUMPAGES", "1", size=7.5, bold=True)


def _configure_document_fields(document) -> None:
    settings = document.settings._element
    even_and_odd = settings.find(qn("w:evenAndOddHeaders"))
    if even_and_odd is not None:
        settings.remove(even_and_odd)
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def standardize_word_document(
    content: bytes,
    filename: str,
    initials: str,
    description: str | None = None,
    generated_at: datetime | None = None,
) -> tuple[bytes, str, dict[str, Any]]:
    clean_initials = re.sub(r"[^0-9A-Za-zÀ-ž]", "", initials).upper()[:10]
    if not clean_initials:
        raise WordStandardizationError("Inicialet janë të detyrueshme për dokumentin aktual.")
    if not LOGO_PATH.is_file():
        raise WordStandardizationError("Logoja zyrtare PrimEx mungon në server.")

    analysis = analyze_word_document(content, filename)
    document = _load_document(content, filename)
    now = generated_at.astimezone(KOSOVO_TIMEZONE) if generated_at else datetime.now(KOSOVO_TIMEZONE)
    description_value = _sanitize_description(description or "") or analysis.suggested_description
    output_filename = f"{description_value}_{now.strftime('%d.%m.%Y')}_{clean_initials}.docx"

    _configure_document_fields(document)
    for section in document.sections:
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.top_margin = NARROW_MARGIN
        section.right_margin = NARROW_MARGIN
        section.bottom_margin = NARROW_MARGIN
        section.left_margin = NARROW_MARGIN
        section.header_distance = HEADER_FOOTER_DISTANCE
        section.footer_distance = HEADER_FOOTER_DISTANCE
        _create_header(section, now)
        _create_footer(section)

    output = io.BytesIO()
    try:
        document.save(output)
    except Exception as exc:
        raise WordStandardizationError("Dokumenti Word nuk mund të ruhej me standardet PrimEx.") from exc
    document_bytes = output.getvalue()

    verification = analyze_word_document(document_bytes, output_filename)
    failed_checks = [check.label for check in verification.checks if not check.compliant]
    if failed_checks:
        raise WordStandardizationError(
            "Dokumenti nuk u gjenerua sepse kontrolli final dështoi: " + "; ".join(failed_checks)
        )

    original_by_id = {check.id: check for check in analysis.checks}
    corrections = [
        {
            "category": check.id,
            "detail": (
                f"U korrigjua: {check.label}."
                if not original_by_id[check.id].compliant
                else f"U verifikua dhe u ruajt: {check.label}."
            ),
        }
        for check in verification.checks
    ]
    corrections.append(
        {
            "category": "header_footer_layout",
            "detail": "Header/footer u aplikuan në çdo seksion me hapësirë të sigurt dhe pa mbivendosje me përmbajtjen.",
        }
    )
    report = {
        "filename": output_filename,
        "generated_at": now.isoformat(),
        "summary": f"Dokumenti u standardizua në {len(document.sections)} seksion(e) dhe kaloi kontrollin final PrimEx.",
        "corrections": corrections,
        "checks": [asdict(check) for check in verification.checks],
    }
    return document_bytes, output_filename, report
