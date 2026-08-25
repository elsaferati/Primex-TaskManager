from __future__ import annotations

import asyncio
import base64
import hashlib
import json
import logging
import os
import re
import smtplib
import ssl
import html
import io
from dataclasses import dataclass
from datetime import date, datetime, time, timedelta
from email.message import EmailMessage
from email.utils import make_msgid
from typing import Any, Awaitable, Callable
from zoneinfo import ZoneInfo

import httpx
from pydantic import BaseModel, Field

from app.config import settings

logger = logging.getLogger(__name__)
REPORT_TYPE = "primeflow_1h"
SLOTS = ("10:00", "11:00", "11:50", "14:10", "14:20", "15:50")
STATUS_ORDER = {"IN_PROGRESS": 0, "TODO": 1, "DONE": 2}
STATUS_MARKERS = {"IN_PROGRESS": "🟡 IN PROGRESS", "TODO": "⚪ TODO", "DONE": "✅ DONE"}
REMINDER_CATEGORY_NORMALIZED = "pyetjet per 1h"
BOARD_REMINDER_SECTION_TITLE = "PYETJET PER 1H - BORD"
REMINDER_SECTION_TITLE = "STAFF - HAPAT PER 1H"
UNDISCUSSED_NOTES_SECTION_TITLE = "NOTA PA DISKUTUARA"
TECHNICAL_TAGS = re.compile(
    r"\[\[\s*/?\s*(?:added|done(?:\s*:\s*(?:grey|gray|blue|green))?)\s*\]\]",
    re.IGNORECASE,
)
ADDED_TAGS = re.compile(r"\[\[\s*/?\s*added\s*\]\]", re.IGNORECASE)
DONE_BLOCK = re.compile(
    r"\[\[\s*done(?:\s*:\s*(grey|gray|blue|green))?\s*\]\](.*?)\[\[\s*/\s*done\s*\]\]",
    re.IGNORECASE | re.DOTALL,
)
STRIKE_COLORS = {"grey": "#6b7280", "gray": "#6b7280", "blue": "#2563eb", "green": "#16a34a"}
STRIKE_TIMESTAMP_DISPLAY = re.compile(r"\s+\d{2}:\d{2}\s+\d{2}\.\d{2}\s*$", re.MULTILINE)
NUMBERED_ITEM = re.compile(r"(?<!\S)(\d+)\.\s*")
TRANSIENT_CODES = {429, 500, 502, 503, 504}
STATUS_COLORS = {
    "TODO": ("#fbcfe8", "#111827", "#ec4899"),
    "IN_PROGRESS": ("#fef3c7", "#111827", "#d97706"),
    "DONE": ("#d4ffe1", "#14532d", "#22c55e"),
}
REPORT_STRIKE_LEGEND = (
    ("Blue strike", "Kryer ne intervalin e caktuar", STRIKE_COLORS["blue"]),
    ("Green strike", "Kryer me heret", STRIKE_COLORS["green"]),
    ("Grey strike", "Kryer dje", STRIKE_COLORS["grey"]),
)
BLOCKED_SECTION_TITLE_PREFIX = "BLLOK 14:30-15:30"


class GmailVerificationError(RuntimeError):
    def __init__(self, message: str, response: dict[str, Any] | None = None) -> None:
        super().__init__(message)
        self.response = response or {}


class ReportTask(BaseModel):
    title: str
    description: str
    marked_title: str = ""
    marked_description: str = ""
    department: str = "-"
    status: str
    marker: str


class ReportEmployee(BaseModel):
    name: str
    tasks: list[ReportTask] = Field(default_factory=list)


class ReportSection(BaseModel):
    title: str
    employees: list[ReportEmployee] = Field(default_factory=list)


class ReportReminderQuestion(BaseModel):
    text: str
    guidance: str = ""


class ReportUndiscussedNote(BaseModel):
    content: str
    author: str = "-"
    created_at: datetime | None = None


def _board_reminder_questions() -> list[ReportReminderQuestion]:
    return [
        ReportReminderQuestion(text="Slotin paraprak/aktual"),
        ReportReminderQuestion(text="A ke filluar me slotin aktual?"),
        ReportReminderQuestion(text="Nese jo, kur?"),
        ReportReminderQuestion(text="A kryhet sot?"),
        ReportReminderQuestion(text="A kryhet kete jave?"),
        ReportReminderQuestion(text="A arrihet RLZ javor?"),
        ReportReminderQuestion(text="Done? / Strikes? / Notes te reja? Data? AM/PM? Kujt?"),
        ReportReminderQuestion(text="BZ Notes", guidance="Secili i lexon vet para BZ me GA"),
    ]


class ReportDocument(BaseModel):
    subject: str
    report_date: date
    report_slot: str
    generated_at: datetime
    source_generated_at: datetime
    recipients: dict[str, list[str]]
    sections: list[ReportSection]
    board_reminders: list[ReportReminderQuestion] = Field(default_factory=_board_reminder_questions)
    reminders: list[ReportReminderQuestion] = Field(default_factory=list)
    undiscussed_notes: list[ReportUndiscussedNote] = Field(default_factory=list)
    truncated: bool = False

    @property
    def task_count(self) -> int:
        return sum(len(employee.tasks) for section in self.sections for employee in section.employees)


def report_timezone() -> ZoneInfo:
    return ZoneInfo(os.getenv("PRIMEFLOW_REPORT_TIMEZONE", "Europe/Tirane"))


def previous_working_day(day: date, holidays: set[date] | None = None) -> date:
    holidays = holidays or set()
    candidate = day - timedelta(days=1)
    while candidate.weekday() >= 5 or candidate in holidays:
        candidate -= timedelta(days=1)
    return candidate


def report_subject(day: date, slot: str) -> str:
    if slot not in SLOTS:
        raise ValueError(f"Unsupported report slot: {slot}")
    # Subjects are plain text in Gmail. Brackets give the leading time a
    # strong, consistent visual treatment without relying on unsupported HTML.
    return f"【{slot}】 PrimeFlow 1H - {day:%d.%m.%Y}"


def exact_subject(headers: list[dict[str, str]], expected: str) -> bool:
    return any(h.get("name", "").lower() == "subject" and h.get("value") == expected for h in headers)


def clean_description(value: str | None) -> str:
    cleaned = TECHNICAL_TAGS.sub("", value or "")
    return STRIKE_TIMESTAMP_DISPLAY.sub("", re.sub(r"[ \t]+", " ", cleaned)).strip()


def clean_title(value: str | None) -> str:
    return clean_description(value)


def preserve_done_marks(value: str | None) -> str:
    return ADDED_TAGS.sub("", value or "").strip()


def split_numbered_text(value: str) -> tuple[str, list[str]]:
    matches = list(NUMBERED_ITEM.finditer(value))
    if not matches:
        return value.strip(), []
    heading = value[:matches[0].start()].strip()
    items = [
        value[match.start():matches[index + 1].start() if index + 1 < len(matches) else len(value)].strip()
        for index, match in enumerate(matches)
    ]
    return heading, items


def split_task_display(title: str) -> tuple[str, list[str]]:
    """First line stays the black title; every following line renders grey."""
    heading, numbered_items = split_numbered_text(title)
    heading_lines = [line.strip() for line in heading.splitlines() if line.strip()]
    if not heading_lines:
        return "", numbered_items
    return heading_lines[0], [*heading_lines[1:], *numbered_items]


def done_ranges(value: str, marked_source: str) -> list[tuple[int, int, str]]:
    ranges = []
    cursor = 0
    for match in DONE_BLOCK.finditer(marked_source):
        selected = TECHNICAL_TAGS.sub("", match.group(2))
        start = value.find(selected, cursor)
        if start >= 0:
            colour = STRIKE_COLORS.get((match.group(1) or "grey").lower(), STRIKE_COLORS["grey"])
            ranges.append((start, start + len(selected), colour))
            cursor = start + len(selected)
    return ranges


def _task_date(item: dict[str, Any]) -> date | None:
    raw = item.get("report_date") or item.get("date") or item.get("day")
    if not raw:
        raw = item.get("planned_for") or item.get("due_date") or item.get("start_date")
    try:
        return date.fromisoformat(str(raw)[:10]) if raw else None
    except ValueError:
        return None


def _slot(item: dict[str, Any]) -> str | None:
    return item.get("one_h_report_slot") or item.get("slot") or item.get("time_slot")


def _source_slot_for_report_slot(slot: str) -> str:
    """The internal 14:10 report remains distinct from the 14:20 Today digest."""
    return "14:20" if slot == "14:10" else slot


def _employee(item: dict[str, Any]) -> str:
    return str(
        item.get("employee") or item.get("person") or item.get("owner")
        or item.get("user_name") or item.get("assignee_name") or item.get("user") or ""
    ).strip()


def employee_initials(name: str) -> str:
    parts = re.findall(r"[^\W\d_]+", name, flags=re.UNICODE)
    return "".join(part[0] for part in parts).upper()


def _weekly_planner_user_sort_key(item: dict[str, Any]) -> tuple[int, str, int, int, str] | None:
    """Read the user order supplied by the Common View payload.

    Common View assigns this key to 1H rows using the same department and
    Weekly Planner user order used by M1, M2, and M3.  Keep it as payload
    metadata so the report can use one normalized order for every output.
    """
    raw = item.get("weekly_planner_sort") or item.get("weeklyPlannerSort")
    if not isinstance(raw, (list, tuple)) or len(raw) < 5:
        return None
    try:
        return (
            int(raw[0]),
            str(raw[1]).casefold(),
            int(raw[2]),
            int(raw[3]),
            str(raw[4]).casefold(),
        )
    except (TypeError, ValueError):
        return None


def _employee_sort_key(employee: str, tasks: list[dict[str, Any]]) -> tuple[int, str, int, int, str]:
    keys = [key for task in tasks if (key := _weekly_planner_user_sort_key(task)) is not None]
    return min(keys) if keys else (10**6, "~", 1, 10**6, employee.casefold())


def _group_tasks_by_employee(tasks: list[dict[str, Any]]) -> list[tuple[str, list[dict[str, Any]]]]:
    grouped: dict[str, list[dict[str, Any]]] = {}
    for task in tasks:
        grouped.setdefault(_employee(task), []).append(task)
    return sorted(grouped.items(), key=lambda group: _employee_sort_key(*group))


def filter_tasks(
    items: list[dict[str, Any]], day: date, slot: str | None = None,
) -> list[dict[str, Any]]:
    seen: set[str] = set()
    result = []
    for item in items:
        employee = _employee(item)
        title = item.get("task_title") or item.get("title") or item.get("task")
        if not employee or not str(title or "").strip() or _task_date(item) != day:
            continue
        if slot is not None and _slot(item) != slot:
            continue
        key = str(item.get("id") or json.dumps(item, sort_keys=True, default=str))
        if key in seen:
            continue
        seen.add(key)
        status = str(item.get("status") or "").upper()
        if status not in STATUS_ORDER:
            logger.error("Unexpected task status task_id=%s status=%s", key, status)
            continue
        result.append(item)
    return result


def _render_section(title: str, tasks: list[dict[str, Any]]) -> str:
    lines = [title]
    employee_groups = _group_tasks_by_employee(tasks)
    for employee_index, (employee, employee_tasks) in enumerate(employee_groups, 1):
        lines.append(f"{employee_index}. {employee}")
        ordered = sorted(employee_tasks, key=lambda x: (STATUS_ORDER[str(x.get("status")).upper()], str(x.get("task_title") or x.get("title") or x.get("task"))))
        for task_index, task in enumerate(ordered, 1):
            status = str(task.get("status")).upper()
            title_value = clean_title(str(task.get("task_title") or task.get("title") or task.get("task")))
            description = clean_description(task.get("description") if "description" in task else task.get("note"))
            lines.extend([f"{employee_index}.{task_index} {STATUS_MARKERS[status]} {title_value}", "Përshkrimi:", description])
    if not employee_groups:
        lines.append("(Asnjë detyrë)")
    return "\n".join(lines)


def _document_section(
    title: str,
    tasks: list[dict[str, Any]],
    title_overrides: dict[str, tuple[str, str]] | None = None,
    description_overrides: dict[str, tuple[str, str]] | None = None,
    department_codes: dict[str, str] | None = None,
) -> ReportSection:
    employees = []
    for employee, employee_tasks in _group_tasks_by_employee(tasks):
        ordered = sorted(employee_tasks, key=lambda x: (
            STATUS_ORDER[str(x.get("status")).upper()],
            str(x.get("task_title") or x.get("title") or x.get("task")),
        ))
        report_tasks = []
        for task in ordered:
            # Common View's `title` is the full source-note content when a
            # task comes from GA/KA or PX notes. `task_title` is only its
            # one-line task label, so it must not hide the numbered points.
            raw_title = str(task.get("title") or task.get("task_title") or task.get("task"))
            raw_description = task.get("description") if "description" in task else task.get("note")
            task_id = str(task.get("task_id") or task.get("id"))
            title_override = (title_overrides or {}).get(task_id)
            description_override = (description_overrides or {}).get(task_id)
            description_duplicates_title = bool(raw_description) and (
                clean_description(raw_description) == clean_title(raw_title)
            )
            department_id = str(task.get("department_id") or task.get("departmentId") or "")
            department = (
                task.get("department_code")
                or task.get("departmentCode")
                or (department_codes or {}).get(department_id)
                or "-"
            )
            report_tasks.append(ReportTask(
                title=clean_title(title_override[0]) if title_override else clean_title(raw_title),
                description="" if description_duplicates_title else (
                    clean_description(description_override[0]) if description_override else clean_description(raw_description)
                ),
                marked_title=title_override[1] if title_override else preserve_done_marks(raw_title),
                marked_description="" if description_duplicates_title else (
                    description_override[1] if description_override else preserve_done_marks(raw_description)
                ),
                department=str(department).strip() or "-",
                status=str(task.get("status")).upper(),
                marker=STATUS_MARKERS[str(task.get("status")).upper()],
            ))
        employees.append(ReportEmployee(
            name=employee_initials(employee),
            tasks=report_tasks,
        ))
    return ReportSection(title=title, employees=employees)


def build_report_document(
    data: dict[str, Any],
    report_day: date,
    slot: str,
    recipients: dict[str, list[str]] | None = None,
    reminders: list[ReportReminderQuestion] | None = None,
    title_overrides: dict[str, tuple[str, str]] | None = None,
    description_overrides: dict[str, tuple[str, str]] | None = None,
    undiscussed_notes: list[ReportUndiscussedNote] | None = None,
) -> ReportDocument:
    guardrails = data.get("guardrails") or {}
    truncated = any((guardrails.get("truncated") or {}).values())
    if truncated:
        raise ValueError("Common View contains truncated buckets")
    items = data.get("items") or {}
    department_codes = {
        str(department.get("id")): str(department.get("code") or "").strip()
        for department in (data.get("departments") or [])
        if isinstance(department, dict) and department.get("id")
    }
    one_h = items.get("oneH") or data.get("tasks") or []
    definitions: list[tuple[str, list[dict[str, Any]]]] = []
    if slot in {"10:00", "14:20"}:
        # The morning report is the full-day baseline. The new 14:20 report
        # repeats today's work completed or planned through 14:20 only.
        candidate_slots = ("10:00", "11:00", "11:50", "14:20", "15:50") if slot == "10:00" else (
            "10:00", "11:00", "11:50", "14:20"
        )
        for candidate in candidate_slots:
            definitions.append(
                (
                    f"{candidate} SLOTI {report_day:%d.%m.%Y}",
                    filter_tasks(one_h, report_day, candidate),
                )
            )
        definitions.extend([
            ("DETYRA PA SLOT – E GJITHË DITA", [
                task for task in filter_tasks(one_h, report_day, None) if _slot(task) is None
            ]),
            ("P: PERSONALE", filter_tasks(items.get("personal") or [], report_day)),
            ("R1 = 1H", filter_tasks(items.get("r1") or [], report_day)),
        ])
    else:
        previous_slot = SLOTS[SLOTS.index(slot) - 1]
        definitions.extend([
            (
                f"{slot} SLOTI {report_day:%d.%m.%Y}",
                filter_tasks(one_h, report_day, _source_slot_for_report_slot(slot)),
            ),
            (
                f"{previous_slot} SLOTI PARAPRAK {report_day:%d.%m.%Y}",
                filter_tasks(one_h, report_day, _source_slot_for_report_slot(previous_slot)),
            ),
        ])
        # The normal afternoon report uses the internal 14:10 identity so it
        # stays distinct from the Today digest when both deliver at 14:20.
        # It covers the 14:20 task bucket and keeps BLL work visible.
        if slot == "14:10":
            definitions.append((
                f"{BLOCKED_SECTION_TITLE_PREFIX} {report_day:%d.%m.%Y}",
                filter_tasks(items.get("blocked") or [], report_day),
            ))
    generated_value = data.get("generated_at") or datetime.now(report_timezone()).isoformat()
    source_generated = datetime.fromisoformat(str(generated_value).replace("Z", "+00:00"))
    return ReportDocument(
        subject=report_subject(report_day, slot),
        report_date=report_day,
        report_slot=slot,
        generated_at=datetime.now(report_timezone()),
        source_generated_at=source_generated,
        recipients=recipients or {"to": [], "cc": [], "bcc": []},
        sections=[
            _document_section(
                title,
                tasks,
                title_overrides,
                description_overrides,
                department_codes,
            )
            for title, tasks in definitions
        ],
        board_reminders=_board_reminder_questions(),
        reminders=list(reminders or []),
        undiscussed_notes=list(undiscussed_notes or []),
    )


def render_plain_text(document: ReportDocument) -> str:
    blocks = [document.subject, f"Generated: {document.generated_at.isoformat()}", ""]
    for reminder_title, questions in (
        (BOARD_REMINDER_SECTION_TITLE, document.board_reminders),
        (REMINDER_SECTION_TITLE, document.reminders),
    ):
        if not questions:
            continue
        reminder_lines = [reminder_title]
        for index, question in enumerate(questions, 1):
            reminder_lines.append(f"{index}. {question.text}")
            if question.guidance:
                reminder_lines.append(f"   {question.guidance}")
        blocks.append("\n".join(reminder_lines))
    for section in document.sections:
        lines = [section.title]
        if not section.employees:
            lines.append("(Asnjë detyrë)")
        for employee in section.employees:
            lines.append(employee.name)
            for task in employee.tasks:
                heading, detail_lines = split_task_display(task.title)
                if heading:
                    lines.append(heading)
                lines.extend(detail_lines)
                if task.description:
                    lines.append(task.description)
        blocks.append("\n".join(lines))
    if document.undiscussed_notes:
        note_lines = [UNDISCUSSED_NOTES_SECTION_TITLE, "NR | NOTE | KUSH | DATA"]
        for index, note in enumerate(document.undiscussed_notes, 1):
            created = note.created_at.astimezone(report_timezone()).strftime("%d.%m %H:%M") if note.created_at else "-"
            note_lines.append(f"{index} | {note.content} | {note.author} | {created}")
        blocks.append("\n".join(note_lines))
    return "\n\n".join(blocks)


def render_html(
    document: ReportDocument,
    *,
    pre_sections_html: str = "",
    content_width: int = 600,
) -> str:
    """Outlook-safe 1H HTML: nested tables, inline styles, bgcolor (Word engine)."""

    def marked_html(value: str, marked_source: str) -> str:
        ranges = done_ranges(value, marked_source)
        boundaries = sorted({0, len(value), *(point for start, end, _colour in ranges for point in (start, end))})
        parts = []
        for start, end in zip(boundaries, boundaries[1:]):
            content = html.escape(value[start:end]).replace(chr(10), "<br>")
            mark = next((item for item in ranges if item[0] <= start and item[1] >= end), None)
            if mark is not None:
                content = (
                    f"<span style=\"color:{mark[2]};text-decoration:line-through;text-decoration-thickness:2px;\">"
                    f"{content}</span>"
                )
            parts.append(content)
        return "".join(parts)

    def detail_row(content: str) -> str:
        return (
            f"<div style=\"font-family:Arial,sans-serif;font-size:13px;line-height:1.4;"
            f"color:#64748b;font-weight:400;margin:5px 0 0;\">{content}</div>"
        )

    def task_card(background: str, accent: str, title_html: str, body_html: str) -> str:
        # Left accent as its own cell — Outlook ignores CSS border-left reliably.
        return (
            "<table role=\"presentation\" width=\"100%\" cellspacing=\"0\" cellpadding=\"0\" border=\"0\" "
            "style=\"width:100%;border-collapse:collapse;margin:8px 0;\">"
            "<tr>"
            f"<td width=\"6\" bgcolor=\"{accent}\" style=\"width:6px;background-color:{accent};"
            f"font-size:0;line-height:0;\">&nbsp;</td>"
            f"<td bgcolor=\"{background}\" style=\"background-color:{background};border:1px solid {accent};"
            f"border-left:0;padding:10px 12px;font-family:Arial,sans-serif;\">"
            f"<div style=\"font-family:Arial,sans-serif;font-size:14px;line-height:1.35;"
            f"font-weight:700;color:#050505;\">{title_html}</div>"
            f"{body_html}"
            "</td></tr></table>"
        )

    def section_title_block(title: str) -> str:
        return (
            "<table role=\"presentation\" width=\"100%\" cellspacing=\"0\" cellpadding=\"0\" border=\"0\" "
            "style=\"width:100%;border-collapse:collapse;margin:18px 0 10px;\">"
            "<tr>"
            "<td width=\"6\" bgcolor=\"#2563eb\" style=\"width:6px;background-color:#2563eb;"
            "font-size:0;line-height:0;\">&nbsp;</td>"
            "<td bgcolor=\"#eef2ff\" style=\"background-color:#eef2ff;padding:11px 14px;"
            "font-family:Arial,sans-serif;font-size:16px;font-weight:800;color:#0f172a;\">"
            f"{html.escape(title)}"
            "</td></tr></table>"
        )

    def section_separator() -> str:
        """A reliable visual break between report slots in Gmail and Outlook."""
        return (
            '<table role="presentation" width="100%" cellspacing="0" cellpadding="0" border="0" '
            'data-report-section-separator="true" '
            'style="width:100%;border-collapse:collapse;margin:30px 0 4px;">'
            '<tr><td height="4" bgcolor="#334155" '
            'style="height:4px;line-height:4px;font-size:0;background-color:#334155;">&nbsp;</td></tr>'
            '</table>'
        )

    def bll_task_table(section: ReportSection) -> str:
        """Compact BLL table used at the end of the normal afternoon report."""
        rows: list[str] = []
        number = 0
        for employee in section.employees:
            for task in employee.tasks:
                number += 1
                background, _, accent = STATUS_COLORS[task.status]
                heading, detail_lines = split_task_display(task.title)
                details = "".join(
                    detail_row(marked_html(item, task.marked_title)) for item in detail_lines
                )
                if task.description:
                    details += detail_row(marked_html(task.description, task.marked_description))
                rows.append(
                    "<tr>"
                    f'<td style="padding:9px;border:1px solid #cbd5e1;text-align:center;font-family:Arial,sans-serif;">{number}</td>'
                    f'<td style="padding:9px;border:1px solid #cbd5e1;font-family:Arial,sans-serif;font-weight:700;">{html.escape(employee.name)}</td>'
                    f'<td style="padding:9px;border:1px solid #cbd5e1;font-family:Arial,sans-serif;">{html.escape(task.department)}</td>'
                    f'<td bgcolor="{background}" style="padding:9px;border:1px solid {accent};font-family:Arial,sans-serif;">'
                    f'<div style="font-size:14px;font-weight:700;color:#050505;">{marked_html(heading or task.title, task.marked_title)}</div>{details}'
                    "</td></tr>"
                )
        if not rows:
            rows.append(
                '<tr><td colspan="4" style="padding:10px;border:1px solid #cbd5e1;font-family:Arial,sans-serif;color:#64748b;">(Asnjë detyrë)</td></tr>'
            )
        return (
            '<table role="presentation" width="100%" cellspacing="0" cellpadding="0" border="0" '
            'data-bll-task-table="true" style="width:100%;border-collapse:collapse;margin:8px 0 4px;">'
            '<tr bgcolor="#e2e8f0">'
            '<th style="padding:8px;border:1px solid #cbd5e1;font-family:Arial,sans-serif;text-align:center;font-size:12px;">NR</th>'
            '<th style="padding:8px;border:1px solid #cbd5e1;font-family:Arial,sans-serif;text-align:left;font-size:12px;">KUSH</th>'
            '<th style="padding:8px;border:1px solid #cbd5e1;font-family:Arial,sans-serif;text-align:left;font-size:12px;">DEP</th>'
            '<th style="padding:8px;border:1px solid #cbd5e1;font-family:Arial,sans-serif;text-align:left;font-size:12px;">TITULLI</th>'
            "</tr>"
            f"{''.join(rows)}"
            "</table>"
        )

    def undiscussed_notes_table(notes: list[ReportUndiscussedNote]) -> str:
        rows = []
        for index, note in enumerate(notes, 1):
            created = note.created_at.astimezone(report_timezone()).strftime("%d.%m %H:%M") if note.created_at else "-"
            rows.append(
                '<tr bgcolor="#dbeafe" style="background-color:#dbeafe;">'
                f'<td bgcolor="#dbeafe" style="background-color:#dbeafe;padding:8px;border:1px solid #93c5fd;text-align:center;font-family:Arial,sans-serif;vertical-align:top;">{index}</td>'
                f'<td bgcolor="#dbeafe" style="background-color:#dbeafe;padding:8px;border:1px solid #93c5fd;font-family:Arial,sans-serif;vertical-align:top;white-space:pre-wrap;">{html.escape(note.content).replace(chr(10), "<br>")}</td>'
                f'<td bgcolor="#dbeafe" style="background-color:#dbeafe;padding:8px;border:1px solid #93c5fd;font-family:Arial,sans-serif;vertical-align:top;">{html.escape(note.author)}</td>'
                f'<td bgcolor="#dbeafe" style="background-color:#dbeafe;padding:8px;border:1px solid #93c5fd;font-family:Arial,sans-serif;vertical-align:top;white-space:nowrap;">{created}</td>'
                "</tr>"
            )
        return (
            '<table role="presentation" width="100%" cellspacing="0" cellpadding="0" border="0" '
            'data-undiscussed-notes-table="true" style="width:100%;border-collapse:collapse;margin:8px 0 4px;">'
            '<tr bgcolor="#e2e8f0">'
            '<th style="padding:8px;border:1px solid #cbd5e1;font-family:Arial,sans-serif;text-align:center;font-size:12px;">NR</th>'
            '<th style="padding:8px;border:1px solid #cbd5e1;font-family:Arial,sans-serif;text-align:left;font-size:12px;">NOTE</th>'
            '<th style="padding:8px;border:1px solid #cbd5e1;font-family:Arial,sans-serif;text-align:left;font-size:12px;">KUSH</th>'
            '<th style="padding:8px;border:1px solid #cbd5e1;font-family:Arial,sans-serif;text-align:left;font-size:12px;">DATA</th>'
            '</tr>'
            f"{''.join(rows)}"
            "</table>"
        )

    def reminder_column(title: str, questions: list[ReportReminderQuestion]) -> str:
        question_parts = []
        for index, question in enumerate(questions, 1):
            guidance = (
                f' <span style="color:#64748b;font-weight:400;">'
                f'({html.escape(question.guidance).replace(chr(10), " / ")})</span>'
                if question.guidance else ""
            )
            question_parts.append(
                f'<span style="white-space:normal;"><strong>{index}.</strong> '
                f'{html.escape(question.text)}{guidance}</span>'
            )
        return (
            section_title_block(title)
            + '<table role="presentation" width="100%" cellspacing="0" cellpadding="0" border="0" '
            'data-compact-reminder-row="true" style="width:100%;border-collapse:collapse;margin:0;">'
            '<tr><td bgcolor="#f8fafc" style="background-color:#f8fafc;border:1px solid #94a3b8;'
            'border-left:6px solid #64748b;padding:9px 10px;font-family:Arial,sans-serif;'
            'font-size:13px;line-height:1.45;color:#0f172a;">'
            + ' <strong style="display:inline-block;color:#1e3a5f;font-size:20px;line-height:1;'
            'font-weight:900;vertical-align:-2px;padding:0 5px;">/</strong> '.join(question_parts)
            + '</td></tr></table>'
        )

    def board_reminder_column(questions: list[ReportReminderQuestion]) -> str:
        return (
            '<div data-board-reminder-columns="true">'
            + reminder_column(BOARD_REMINDER_SECTION_TITLE, questions)
            + '</div>'
        )

    def report_legend_html() -> str:
        def strike_item(label: str, description: str, color: str) -> str:
            return (
                '<td width="33.33%" style="width:33.33%;padding:6px 8px;border:1px solid #cbd5e1;'
                'font-family:Arial,sans-serif;font-size:11px;line-height:1.3;vertical-align:middle;">'
                f'<strong style="color:{color};text-decoration:line-through;'
                f'text-decoration-thickness:2px;">{html.escape(label)}</strong><br>'
                f'<span style="color:#64748b;">{html.escape(description)}</span></td>'
            )

        return (
            '<table role="presentation" width="100%" cellspacing="0" cellpadding="0" border="0" '
            'data-report-color-legend="true" style="width:100%;border-collapse:collapse;margin:0 0 12px;">'
            '<tr><th colspan="3" bgcolor="#f1f5f9" style="background-color:#f1f5f9;'
            'border:1px solid #cbd5e1;padding:7px 9px;font-family:Arial,sans-serif;'
            'font-size:12px;text-align:left;color:#0f172a;">COLOR LEGEND</th></tr><tr>'
            + "".join(strike_item(*item) for item in REPORT_STRIKE_LEGEND)
            + '</tr></table>'
        )

    body_chunks: list[str] = [report_legend_html()]
    if pre_sections_html:
        body_chunks.append(pre_sections_html)

    if document.board_reminders and document.reminders:
        body_chunks.append(
            '<table role="presentation" width="100%" cellspacing="0" cellpadding="0" border="0" '
            'data-reminder-columns="true" style="width:100%;border-collapse:collapse;">'
            '<tr>'
            '<td width="50%" valign="top" style="width:50%;padding:0 6px 0 0;vertical-align:top;">'
            f"{reminder_column(REMINDER_SECTION_TITLE, document.reminders)}"
            '</td>'
            '<td width="50%" valign="top" style="width:50%;padding:0 0 0 6px;vertical-align:top;">'
            f"{board_reminder_column(document.board_reminders)}"
            '</td>'
            '</tr></table>'
        )
    elif document.board_reminders:
        body_chunks.append(board_reminder_column(document.board_reminders))
    elif document.reminders:
        body_chunks.append(reminder_column(REMINDER_SECTION_TITLE, document.reminders))

    for section_index, section in enumerate(document.sections):
        if section_index:
            body_chunks.append(section_separator())
        body_chunks.append(section_title_block(section.title))
        if section.title.startswith(BLOCKED_SECTION_TITLE_PREFIX):
            body_chunks.append(bll_task_table(section))
            continue
        if not section.employees:
            body_chunks.append(
                "<div style=\"font-family:Arial,sans-serif;color:#64748b;padding:8px 0;\">(Asnjë detyrë)</div>"
            )
        for employee in section.employees:
            body_chunks.append(
                f"<div style=\"font-family:Arial,sans-serif;font-size:15px;font-weight:800;"
                f"color:#0f172a;margin:12px 0 6px;\">{html.escape(employee.name)}</div>"
            )
            for task in employee.tasks:
                background, _, accent = STATUS_COLORS[task.status]
                heading, detail_lines = split_task_display(task.title)
                detail_html = "".join(
                    detail_row(marked_html(item, task.marked_title)) for item in detail_lines
                )
                if task.description:
                    detail_html += detail_row(
                        marked_html(task.description, task.marked_description)
                    )
                body_chunks.append(
                    task_card(
                        background,
                        accent,
                        marked_html(heading or task.title, task.marked_title),
                        detail_html,
                    )
                )

    if document.undiscussed_notes:
        body_chunks.append(section_separator())
        body_chunks.append(section_title_block(UNDISCUSSED_NOTES_SECTION_TITLE))
        body_chunks.append(undiscussed_notes_table(document.undiscussed_notes))

    meta = (
        f"Generated {html.escape(document.generated_at.isoformat())} · {document.task_count} tasks"
    )
    return f"""<!doctype html>
<html><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<!--[if mso]>
<style type="text/css">
table, td {{ font-family: Arial, sans-serif !important; }}
</style>
<![endif]-->
</head>
<body style="margin:0;padding:0;background:#f8fafc;font-family:Arial,sans-serif;color:#0f172a;">
<table role="presentation" width="100%" cellspacing="0" cellpadding="0" border="0" style="width:100%;background:#f8fafc;border-collapse:collapse;">
<tr><td align="center" style="padding:16px 8px;">
<!--[if mso]>
<table role="presentation" width="{content_width}" cellspacing="0" cellpadding="0" border="0"><tr><td>
<![endif]-->
<table role="presentation" width="{content_width}" cellspacing="0" cellpadding="0" border="0" style="width:{content_width}px;max-width:100%;background:#ffffff;border:1px solid #e5e7eb;border-collapse:collapse;">
<tr><td bgcolor="#8799b2" style="background-color:#8799b2;padding:18px 20px;font-family:Arial,sans-serif;">
<div style="font-family:Arial,sans-serif;font-size:22px;line-height:1.25;font-weight:700;color:#ffffff;margin:0 0 4px;">{html.escape(document.subject)}</div>
<div style="font-family:Arial,sans-serif;font-size:13px;line-height:1.35;color:#ffffff;">{meta}</div>
</td></tr>
<tr><td style="padding:14px 16px;font-family:Arial,sans-serif;background:#ffffff;">
{''.join(body_chunks)}
</td></tr>
</table>
<!--[if mso]>
</td></tr></table>
<![endif]-->
</td></tr></table>
</body></html>"""


def render_docx(document: ReportDocument) -> bytes:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    def shade(cell: Any, color: str) -> None:
        fill = OxmlElement("w:shd")
        fill.set(qn("w:fill"), color.lstrip("#"))
        cell._tc.get_or_add_tcPr().append(fill)

    def add_marked_runs(paragraph: Any, value: str, marked_source: str, *, bold: bool, color: str) -> None:
        ranges = done_ranges(value, marked_source)
        boundaries = sorted({0, len(value), *(point for start, end, _colour in ranges for point in (start, end))})
        for start, end in zip(boundaries, boundaries[1:]):
            run = paragraph.add_run(value[start:end])
            run.bold = bold
            mark = next((item for item in ranges if item[0] <= start and item[1] >= end), None)
            run.font.strike = mark is not None
            run.font.size = Pt(9.5 if bold else 9)
            run.font.color.rgb = RGBColor.from_string((mark[2] if mark else color).lstrip("#"))

    output = io.BytesIO()
    doc = Document()
    doc.sections[0].top_margin = doc.sections[0].bottom_margin = Inches(0.65)
    doc.sections[0].left_margin = doc.sections[0].right_margin = Inches(0.7)
    title_cell = doc.add_table(rows=1, cols=1).cell(0, 0)
    shade(title_cell, "#8799b2")
    title_run = title_cell.paragraphs[0].add_run(document.subject)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(255, 255, 255)
    meta = title_cell.add_paragraph(f"Generated {document.generated_at.isoformat()} · {document.task_count} tasks")
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    legend_header = doc.add_table(rows=1, cols=1).cell(0, 0)
    shade(legend_header, "#f1f5f9")
    legend_title = legend_header.paragraphs[0].add_run("COLOR LEGEND")
    legend_title.bold = True
    legend_title.font.size = Pt(10)
    legend_table = doc.add_table(rows=1, cols=3)
    legend_table.style = "Table Grid"
    for cell, (label, description, color) in zip(legend_table.rows[0].cells, REPORT_STRIKE_LEGEND):
        label_run = cell.paragraphs[0].add_run(label)
        label_run.bold = True
        label_run.font.strike = True
        label_run.font.size = Pt(8.5)
        label_run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
        description_run = cell.add_paragraph().add_run(description)
        description_run.font.size = Pt(7.5)
        description_run.font.color.rgb = RGBColor.from_string("475569")
    for reminder_title, questions in (
        (BOARD_REMINDER_SECTION_TITLE, document.board_reminders),
        (REMINDER_SECTION_TITLE, document.reminders),
    ):
        if not questions:
            continue
        doc.add_paragraph()
        reminder_header = doc.add_table(rows=1, cols=1).cell(0, 0)
        shade(reminder_header, "#eef2ff")
        reminder_run = reminder_header.paragraphs[0].add_run(reminder_title)
        reminder_run.bold = True
        reminder_run.font.size = Pt(13)
        for index, question in enumerate(questions, 1):
            card_cell = doc.add_table(rows=1, cols=1).cell(0, 0)
            shade(card_cell, "#f8fafc")
            add_marked_runs(
                card_cell.paragraphs[0],
                f"{index}. {question.text}",
                f"{index}. {question.text}",
                bold=True,
                color="#050505",
            )
            if question.guidance:
                guidance = card_cell.add_paragraph()
                add_marked_runs(
                    guidance, question.guidance, question.guidance, bold=False, color="#64748b"
                )
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
    for section in document.sections:
        doc.add_paragraph()
        section_cell = doc.add_table(rows=1, cols=1).cell(0, 0)
        shade(section_cell, "#eef2ff")
        section_run = section_cell.paragraphs[0].add_run(section.title)
        section_run.bold = True
        section_run.font.size = Pt(13)
        if not section.employees:
            doc.add_paragraph("(Asnjë detyrë)")
        for employee in section.employees:
            employee_paragraph = doc.add_paragraph()
            employee_paragraph.paragraph_format.space_before = Pt(8)
            employee_paragraph.paragraph_format.space_after = Pt(3)
            employee_run = employee_paragraph.add_run(employee.name)
            employee_run.bold = True
            employee_run.font.size = Pt(11)
            for task in employee.tasks:
                background, _, _ = STATUS_COLORS[task.status]
                card_cell = doc.add_table(rows=1, cols=1).cell(0, 0)
                shade(card_cell, background)
                heading, detail_lines = split_task_display(task.title)
                add_marked_runs(card_cell.paragraphs[0], heading or task.title, task.marked_title, bold=True, color="#050505")
                for item in detail_lines:
                    detail = card_cell.add_paragraph()
                    detail.paragraph_format.space_after = Pt(2)
                    add_marked_runs(detail, item, task.marked_title, bold=False, color="#64748b")
                if task.description:
                    description = card_cell.add_paragraph()
                    add_marked_runs(
                        description, task.description, task.marked_description, bold=False, color="#64748b"
                    )
                doc.add_paragraph().paragraph_format.space_after = Pt(0)
    if document.undiscussed_notes:
        doc.add_paragraph()
        notes_header = doc.add_table(rows=1, cols=1).cell(0, 0)
        shade(notes_header, "#eef2ff")
        notes_title = notes_header.paragraphs[0].add_run(UNDISCUSSED_NOTES_SECTION_TITLE)
        notes_title.bold = True
        notes_title.font.size = Pt(13)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for cell, label in zip(table.rows[0].cells, ("NR", "NOTE", "KUSH", "DATA")):
            shade(cell, "#e2e8f0")
            run = cell.paragraphs[0].add_run(label)
            run.bold = True
        for index, note in enumerate(document.undiscussed_notes, 1):
            cells = table.add_row().cells
            created = note.created_at.astimezone(report_timezone()).strftime("%d.%m %H:%M") if note.created_at else "-"
            for cell, value in zip(cells, (str(index), note.content, note.author, created)):
                cell.text = value
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)
    doc.save(output)
    return output.getvalue()


def render_png(document: ReportDocument) -> bytes:
    from PIL import Image, ImageDraw, ImageFont
    width, margin = 1400, 55
    font_path = os.getenv("PRIMEFLOW_REPORT_FONT_PATH", r"C:\Windows\Fonts\seguiemj.ttf")
    fallback = r"C:\Windows\Fonts\arial.ttf"
    try:
        font = ImageFont.truetype(font_path if os.path.exists(font_path) else fallback, 20)
        bold = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 21)
        heading = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 30)
    except OSError:
        font = bold = heading = ImageFont.load_default()
    import textwrap

    def draw_line_with_marks(x: int, line_y: int, line: str, marked_source: str, line_font: Any, color: str) -> None:
        marked_values = [
            (TECHNICAL_TAGS.sub("", match.group(2)).strip(), STRIKE_COLORS.get((match.group(1) or "grey").lower(), STRIKE_COLORS["grey"]))
            for match in DONE_BLOCK.finditer(marked_source)
        ]
        mark_colour = next(
            (strike_colour for selected, strike_colour in marked_values if selected and (line.strip() in selected or selected in line)),
            None,
        )
        draw.text((x, line_y), line, fill=mark_colour or color, font=line_font)
        if mark_colour:
            bounds = draw.textbbox((x, line_y), line, font=line_font)
            strike_y = (bounds[1] + bounds[3]) // 2
            draw.line((bounds[0], strike_y, bounds[2], strike_y), fill=mark_colour, width=2)

    estimated_lines = (
        13
        + len(document.sections) * 3
        + sum(
            2 + len(textwrap.wrap(question.text, 90)) + len(textwrap.wrap(question.guidance or "", 95))
            for questions in (document.board_reminders, document.reminders)
            for question in questions
        )
        + sum(
            3 + sum(
                4 + len(textwrap.wrap(task.title, 85)) + len(textwrap.wrap(task.description, 90))
                for task in employee.tasks
            )
            for section in document.sections
            for employee in section.employees
        )
        + sum(2 + len(textwrap.wrap(note.content, 82)) for note in document.undiscussed_notes)
    )
    height = max(650, 140 + estimated_lines * 30)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((margin, 35, width - margin, 135), radius=12, fill="#8799b2")
    draw.text((margin + 24, 55), document.subject, fill="white", font=heading)
    draw.text((margin + 24, 99), f"Generated {document.generated_at.isoformat()} · {document.task_count} tasks", fill="white", font=font)
    y = 160
    legend_height = 76
    draw.rounded_rectangle(
        (margin, y, width - margin, y + legend_height),
        radius=8,
        fill="#ffffff",
        outline="#cbd5e1",
        width=1,
    )
    draw.rectangle((margin, y, width - margin, y + 30), fill="#f1f5f9", outline="#cbd5e1")
    draw.text((margin + 10, y + 5), "COLOR LEGEND", fill="#0f172a", font=bold)
    legend_column_width = (width - (2 * margin)) // 3
    for column, (label, description, color) in enumerate(REPORT_STRIKE_LEGEND):
        left = margin + (column * legend_column_width)
        top = y + 30
        if column:
            draw.line((left, top, left, top + 46), fill="#cbd5e1", width=1)
        text_left = left + 10
        draw.text((text_left, top + 5), label, fill=color, font=bold)
        bounds = draw.textbbox((text_left, top + 5), label, font=bold)
        strike_y = (bounds[1] + bounds[3]) // 2
        draw.line((bounds[0], strike_y, bounds[2], strike_y), fill=color, width=2)
        draw.text((text_left, top + 26), description, fill="#64748b", font=font)
    y += legend_height + 18
    for reminder_title, questions in (
        (BOARD_REMINDER_SECTION_TITLE, document.board_reminders),
        (REMINDER_SECTION_TITLE, document.reminders),
    ):
        if not questions:
            continue
        draw.rectangle((margin, y, width - margin, y + 48), fill="#eef2ff")
        draw.rectangle((margin, y, margin + 7, y + 48), fill="#2563eb")
        draw.text((margin + 18, y + 11), reminder_title, fill="#0f172a", font=bold)
        y += 62
        for index, question in enumerate(questions, 1):
            title_lines = textwrap.wrap(f"{index}. {question.text}", 95) or [""]
            guidance_lines = textwrap.wrap(question.guidance, 105) if question.guidance else []
            card_height = 25 + 28 * (len(title_lines) + len(guidance_lines))
            draw.rounded_rectangle(
                (margin + 5, y, width - margin, y + card_height),
                radius=8,
                fill="#f8fafc",
                outline="#64748b",
                width=2,
            )
            draw.rectangle((margin + 5, y + 4, margin + 11, y + card_height - 4), fill="#64748b")
            line_y = y + 12
            for line in title_lines:
                draw.text((margin + 25, line_y), line, fill="#050505", font=bold)
                line_y += 28
            for line in guidance_lines:
                draw.text((margin + 25, line_y), line, fill="#64748b", font=font)
                line_y += 28
            y += card_height + 12
        y += 14
    for section in document.sections:
        draw.rectangle((margin, y, width - margin, y + 48), fill="#eef2ff")
        draw.rectangle((margin, y, margin + 7, y + 48), fill="#2563eb")
        draw.text((margin + 18, y + 11), section.title, fill="#0f172a", font=bold)
        y += 62
        if not section.employees:
            draw.text((margin + 18, y), "(Asnjë detyrë)", fill="#64748b", font=font)
            y += 40
        for employee in section.employees:
            draw.text((margin + 5, y), employee.name, fill="#0f172a", font=bold)
            y += 38
            for task in employee.tasks:
                background, _, accent = STATUS_COLORS[task.status]
                task_heading, detail_items = split_task_display(task.title)
                title_lines = textwrap.wrap(task_heading or task.title, 95) or [""]
                detail_lines = [
                    line for item in detail_items
                    for line in (textwrap.wrap(item, 105, subsequent_indent="   ") or [""])
                ]
                description_lines = textwrap.wrap(task.description, 105) if task.description else []
                card_height = 25 + 28 * (len(title_lines) + len(detail_lines) + len(description_lines))
                draw.rounded_rectangle((margin + 5, y, width - margin, y + card_height), radius=8, fill=background, outline=accent, width=2)
                draw.rectangle((margin + 5, y + 4, margin + 11, y + card_height - 4), fill=accent)
                line_y = y + 12
                for line in title_lines:
                    draw_line_with_marks(margin + 25, line_y, line, task.marked_title, bold, "#050505")
                    line_y += 28
                for line in detail_lines:
                    draw_line_with_marks(margin + 25, line_y, line, task.marked_title, font, "#64748b")
                    line_y += 28
                for line in description_lines:
                    draw_line_with_marks(margin + 25, line_y, line, task.marked_description, font, "#64748b")
                    line_y += 28
                y += card_height + 12
        y += 14
    if document.undiscussed_notes:
        draw.rectangle((margin, y, width - margin, y + 48), fill="#eef2ff")
        draw.rectangle((margin, y, margin + 7, y + 48), fill="#2563eb")
        draw.text((margin + 18, y + 11), UNDISCUSSED_NOTES_SECTION_TITLE, fill="#0f172a", font=bold)
        y += 62
        table_left, table_right = margin + 5, width - margin
        columns = (table_left, table_left + 60, table_left + 900, table_left + 1080, table_right)
        header_height = 34
        draw.rectangle((table_left, y, table_right, y + header_height), fill="#e2e8f0", outline="#94a3b8", width=1)
        for x in columns[1:-1]:
            draw.line((x, y, x, y + header_height), fill="#94a3b8", width=1)
        for label, x in zip(("NR", "NOTE", "KUSH", "DATA"), columns[:-1]):
            draw.text((x + 7, y + 7), label, fill="#0f172a", font=bold)
        y += header_height
        for index, note in enumerate(document.undiscussed_notes, 1):
            note_lines = textwrap.wrap(note.content, 68) or [""]
            row_height = max(32, 26 * len(note_lines))
            draw.rectangle((table_left, y, table_right, y + row_height), fill="#ffffff", outline="#cbd5e1", width=1)
            for x in columns[1:-1]:
                draw.line((x, y, x, y + row_height), fill="#cbd5e1", width=1)
            draw.text((columns[0] + 7, y + 6), str(index), fill="#0f172a", font=font)
            for line_index, line in enumerate(note_lines):
                draw.text((columns[1] + 7, y + 6 + 25 * line_index), line, fill="#0f172a", font=font)
            draw.text((columns[2] + 7, y + 6), note.author, fill="#0f172a", font=font)
            created = note.created_at.astimezone(report_timezone()).strftime("%d.%m %H:%M") if note.created_at else "-"
            draw.text((columns[3] + 7, y + 6), created, fill="#0f172a", font=font)
            y += row_height
        y += 14
    image = image.crop((0, 0, width, y + margin))
    output = io.BytesIO()
    image.save(output, format="PNG", optimize=True)
    return output.getvalue()


def build_report(data: dict[str, Any], report_day: date, slot: str) -> str:
    return render_plain_text(build_report_document(data, report_day, slot))


async def retry(operation: Callable[[], Awaitable[Any]], *, delays: tuple[float, ...] = (0, 2, 5)) -> Any:
    last: Exception | None = None
    for attempt, delay in enumerate(delays, 1):
        if delay:
            await asyncio.sleep(delay)
        try:
            return await operation()
        except (httpx.ConnectError, httpx.ConnectTimeout, httpx.ReadTimeout, httpx.RemoteProtocolError) as exc:
            last = exc
        except httpx.HTTPStatusError as exc:
            if exc.response.status_code not in TRANSIENT_CODES:
                raise
            last = exc
        logger.warning("transient_operation_failure attempt=%s error=%s", attempt, type(last).__name__)
    assert last is not None
    raise last


@dataclass
class PrimeFlowClient:
    base_url: str
    email: str | None
    password: str | None
    access_token: str | None = None

    async def _token(self, client: httpx.AsyncClient) -> str:
        if self.access_token:
            return self.access_token
        response = await client.post("/api/auth/login", json={"email": self.email, "password": self.password})
        response.raise_for_status()
        self.access_token = response.json().get("access_token") or response.json().get("accessToken")
        if not self.access_token:
            raise ValueError("PrimeFlow login response contained no access token")
        return self.access_token

    async def common_view(self, day: date) -> dict[str, Any]:
        async with httpx.AsyncClient(base_url=self.base_url, timeout=30) as client:
            async def retrieve(week_day: date) -> dict[str, Any]:
                params = {
                    "week_start": week_day.isoformat(),
                    "include_all_departments": "true",
                    "freeze_one_h_slots": "false",
                    "max_items_per_bucket": 5000,
                }
                token = await self._token(client)
                response = await client.get(
                    "/api/common-view",
                    params=params,
                    headers={"Authorization": f"Bearer {token}", "Cache-Control": "no-cache"},
                )
                if response.status_code == 401:
                    self.access_token = None
                    token = await self._token(client)
                    response = await client.get(
                        "/api/common-view",
                        params=params,
                        headers={"Authorization": f"Bearer {token}", "Cache-Control": "no-cache"},
                    )
                response.raise_for_status()
                payload = response.json()
                if any((payload.get("guardrails", {}).get("truncated") or {}).values()):
                    raise ValueError("Common View contains truncated buckets")
                return payload
            current = await retry(lambda: retrieve(day))
            if day.weekday() != 0:
                return current
            previous = await retry(lambda: retrieve(previous_working_day(day)))
            for bucket, values in (previous.get("items") or {}).items():
                current.setdefault("items", {}).setdefault(bucket, []).extend(values)
            current["generated_at"] = max(current["generated_at"], previous["generated_at"])
            return current


class GmailService:
    def __init__(self) -> None:
        sender = os.getenv("EMAIL_USER") or settings.EMAIL_USER
        password = os.getenv("EMAIL_PASSWORD") or settings.EMAIL_PASSWORD
        if not sender or not password:
            missing = [
                name for name, value in (("EMAIL_USER", sender), ("EMAIL_PASSWORD", password))
                if not value
            ]
            raise ValueError(f"Missing email configuration: {', '.join(missing)}")
        self.sender = sender.strip()
        self.password = password.replace(" ", "")
        self.host = (os.getenv("EMAIL_HOST") or settings.EMAIL_HOST).strip()
        self.port = int(os.getenv("EMAIL_PORT") or settings.EMAIL_PORT)

    async def find_exact(self, subject: str, recipients: list[str] | None = None) -> dict[str, Any] | None:
        # SMTP has no mailbox-search operation. Database uniqueness and row locks
        # in the delivery service provide the idempotency guard for SMTP sends.
        return None

    async def send_verified(
        self, subject: str, recipients: list[str] | dict[str, list[str]], body: str,
        html_body: str | None = None,
        attachments: list[tuple[str, bytes, str]] | None = None,
        message_id: str | None = None,
    ) -> dict[str, Any]:
        recipient_map = recipients if isinstance(recipients, dict) else {"to": recipients, "cc": [], "bcc": []}
        all_recipients = sum(recipient_map.values(), [])
        if not recipient_map["to"]:
            raise ValueError("At least one To recipient is required")
        message = EmailMessage()
        message_id = message_id or make_msgid(domain=self.sender.rsplit("@", 1)[-1])
        if not message_id.startswith("<"):
            message_id = f"<{message_id}>"
        message["From"] = self.sender
        message["To"] = ", ".join(recipient_map["to"])
        message["Subject"] = subject
        message["Message-ID"] = message_id
        message["X-Entity-Ref-ID"] = message_id.strip("<>")
        if recipient_map["cc"]:
            message["Cc"] = ", ".join(recipient_map["cc"])
        if recipient_map["bcc"]:
            message["Bcc"] = ", ".join(recipient_map["bcc"])
        message.set_content(body, charset="utf-8")
        if html_body:
            # multipart/alternative with HTML last — Outlook/Gmail prefer the last part.
            message.add_alternative(html_body, subtype="html", charset="utf-8")
        for filename, content, mime_type in attachments or []:
            maintype, subtype = mime_type.split("/", 1)
            message.add_attachment(content, maintype=maintype, subtype=subtype, filename=filename)

        def send_smtp() -> None:
            context = ssl.create_default_context()
            with smtplib.SMTP(self.host, self.port, timeout=30) as smtp:
                smtp.ehlo()
                smtp.starttls(context=context)
                smtp.ehlo()
                smtp.login(self.sender, self.password)
                smtp.send_message(message, from_addr=self.sender, to_addrs=all_recipients)

        await asyncio.to_thread(send_smtp)
        return {
            "id": message_id.strip("<>"),
            "threadId": None,
            "provider_message_id": None,
            "transport": "smtp",
        }


def predecessor(day: date, slot: str) -> tuple[date, str]:
    index = SLOTS.index(slot)
    return (previous_working_day(day), "15:50") if index == 0 else (day, SLOTS[index - 1])
