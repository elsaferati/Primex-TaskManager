from __future__ import annotations

import csv
import html
import mimetypes
import uuid
from datetime import datetime, timedelta, timezone
from pathlib import Path
from zoneinfo import ZoneInfo

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from fastapi.responses import FileResponse, HTMLResponse
from sqlalchemy import or_, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload
from pydantic import BaseModel
from starlette.concurrency import run_in_threadpool

from app.api.access import ensure_department_access
from app.api.deps import get_current_user
from app.db import get_db
from app.models.ga_note import GaNote
from app.models.ga_note_attachment import GaNoteAttachment
from app.models.enums import GaNotePriority, GaNoteStatus, GaNoteType, NotificationType, TaskStatus, UserRole
from app.models.project import Project
from app.models.task import Task
from app.models.user import User
from app.models.task_user_comment import TaskUserComment
from app.schemas.ga_note import (
    GaNoteAttachmentOut,
    GaNoteCreate,
    GaNoteOut,
    GaNoteTaskBundleResponse,
    GaNoteTaskBundleUpdate,
    GaNoteTaskDeadlineUpdate,
    GaNoteUpdate,
)
from app.services.audit import add_audit_log
from app.services.daily_realization_baseline import ensure_daily_baselines_for_departments
from app.services.daily_realization_events import record_task_semantic_events, task_semantic_state
from app.services.ga_note_task import ga_note_default_task_description, ga_note_task_title
from app.services.task_strike_events import record_description_strike_events, record_title_strike_events
from app.services.task_daily_progress import upsert_explicit_task_daily_status
from app.services.ga_note_task_instances import (
    GaNoteAssigneeExecutionState,
    apply_ga_note_assignee_execution_states,
    apply_ga_note_shared_task_fields,
    reconcile_ga_note_task_assignees,
)
from app.services.notifications import (
    add_notification,
    notification_task_preview,
    publish_notification,
)
from app.config import settings


router = APIRouter()


class MarkWaitingDoneResponse(BaseModel):
    updated_count: int
    skipped_count: int


class GaNoteTaskDeadlineResponse(BaseModel):
    updated_count: int
    start_date: datetime | None = None
    due_date: datetime | None = None
    is_deadline_important: bool | None = None


# Backward-compatible aliases retained for router tests and callers.
_ga_note_task_title = ga_note_task_title
_ga_note_default_task_description = ga_note_default_task_description


def _attachment_out(attachment: GaNoteAttachment) -> GaNoteAttachmentOut:
    return GaNoteAttachmentOut(
        id=attachment.id,
        note_id=attachment.note_id,
        original_filename=attachment.original_filename,
        stored_filename=attachment.stored_filename,
        content_type=attachment.content_type,
        size_bytes=attachment.size_bytes,
        created_by=attachment.created_by,
        created_at=attachment.created_at,
    )


def _note_out(note: GaNote) -> GaNoteOut:
    return GaNoteOut(
        id=note.id,
        content=note.content,
        created_by=note.created_by,
        note_type=note.note_type,
        status=note.status,
        priority=note.priority,
        start_date=note.start_date,
        due_date=note.due_date,
        completed_at=note.completed_at,
        is_converted_to_task=note.is_converted_to_task,
        is_discussed=note.is_discussed,
        project_id=note.project_id,
        department_id=note.department_id,
        created_at=note.created_at,
        updated_at=note.updated_at,
        attachments=[_attachment_out(a) for a in (note.attachments or [])],
    )


def _ga_note_upload_base_dir() -> Path:
    upload_base = Path(settings.GA_NOTES_UPLOAD_DIR)
    if not upload_base.is_absolute():
        upload_base = Path(__file__).resolve().parents[3] / upload_base
    return upload_base


def _preview_page(title: str, body: str) -> str:
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(title)}</title>
  <style>
    :root {{ color-scheme: light; font-family: Arial, sans-serif; }}
    body {{ margin: 0; padding: 16px; color: #0f172a; background: #fff; }}
    h2 {{ margin: 0 0 14px; font-size: 16px; }}
    h3 {{ margin: 18px 0 8px; font-size: 14px; }}
    .sheet {{ overflow: auto; margin-bottom: 20px; border: 1px solid #cbd5e1; border-radius: 6px; }}
    table {{ width: 100%; border-collapse: collapse; font-size: 12px; }}
    td {{ min-width: 90px; max-width: 320px; padding: 6px 8px; border: 1px solid #cbd5e1; vertical-align: top; white-space: pre-wrap; overflow-wrap: anywhere; }}
    tr:first-child td {{ background: #f1f5f9; font-weight: 600; }}
    pre {{ margin: 0; white-space: pre-wrap; overflow-wrap: anywhere; font: 12px/1.5 Consolas, monospace; }}
    p {{ margin: 0 0 10px; white-space: pre-wrap; overflow-wrap: anywhere; }}
    .muted {{ color: #64748b; font-size: 12px; }}
  </style>
</head>
<body>
  <h2>{html.escape(title)}</h2>
  {body}
</body>
</html>"""


def _preview_table(rows: list[list[object]]) -> str:
    rendered_rows = []
    for row in rows:
        cells = "".join(f"<td>{html.escape('' if value is None else str(value))}</td>" for value in row)
        rendered_rows.append(f"<tr>{cells}</tr>")
    return f'<div class="sheet"><table>{"".join(rendered_rows)}</table></div>'


def _build_attachment_preview(path: Path, original_filename: str) -> str:
    suffix = path.suffix.lower()

    if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        from openpyxl import load_workbook

        workbook = load_workbook(path, read_only=True, data_only=True)
        sections: list[str] = []
        try:
            for worksheet in workbook.worksheets[:8]:
                rows: list[list[object]] = []
                for row_index, row in enumerate(worksheet.iter_rows(values_only=True)):
                    if row_index >= 250:
                        break
                    rows.append(list(row[:40]))
                sections.append(f"<h3>{html.escape(worksheet.title)}</h3>")
                sections.append(_preview_table(rows) if rows else '<p class="muted">Empty sheet</p>')
            if len(workbook.worksheets) > 8:
                sections.append('<p class="muted">Only the first 8 sheets are shown.</p>')
        finally:
            workbook.close()
        return _preview_page(original_filename, "".join(sections))

    if suffix in {".csv", ".tsv"}:
        delimiter = "\t" if suffix == ".tsv" else ","
        rows: list[list[object]] = []
        with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as stream:
            for row_index, row in enumerate(csv.reader(stream, delimiter=delimiter)):
                if row_index >= 500:
                    break
                rows.append(list(row[:40]))
        return _preview_page(original_filename, _preview_table(rows))

    if suffix == ".docx":
        from docx import Document

        document = Document(path)
        sections = [f"<p>{html.escape(paragraph.text)}</p>" for paragraph in document.paragraphs[:500] if paragraph.text]
        for table in document.tables[:20]:
            rows = [[cell.text for cell in row.cells[:20]] for row in table.rows[:200]]
            sections.append(_preview_table(rows))
        return _preview_page(original_filename, "".join(sections) or '<p class="muted">Empty document</p>')

    text_suffixes = {
        ".txt", ".log", ".md", ".json", ".xml", ".html", ".htm", ".css", ".js", ".ts", ".tsx",
        ".jsx", ".py", ".sql", ".yaml", ".yml", ".ini", ".cfg",
    }
    if suffix in text_suffixes:
        content = path.read_text(encoding="utf-8", errors="replace")[:300_000]
        return _preview_page(original_filename, f"<pre>{html.escape(content)}</pre>")

    body = (
        '<p>This file format cannot be displayed directly in the browser.</p>'
        '<p class="muted">Use the Download button in the preview window to open it with the appropriate application.</p>'
    )
    return _preview_page(original_filename, body)


async def _ensure_note_access(note: GaNote, user, db: AsyncSession) -> None:
    # Every authenticated PrimeFlow user may view and edit GA/KA notes.
    return


async def _get_note_or_404(note_id: uuid.UUID, db: AsyncSession) -> GaNote:
    note = (await db.execute(select(GaNote).where(GaNote.id == note_id))).scalar_one_or_none()
    if note is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="GA note not found")
    return note


async def _save_ga_note_attachments(
    note: GaNote,
    files: list[UploadFile],
    db: AsyncSession,
    user,
) -> list[GaNoteAttachmentOut]:
    await _ensure_note_access(note, user, db)

    if not files:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No files uploaded")
    if len(files) > settings.GA_NOTES_MAX_FILES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Too many files. Max {settings.GA_NOTES_MAX_FILES}.",
        )

    max_bytes = settings.GA_NOTES_MAX_FILE_MB * 1024 * 1024
    upload_base = _ga_note_upload_base_dir()
    note_dir = upload_base / str(note.id)
    note_dir.mkdir(parents=True, exist_ok=True)

    created: list[GaNoteAttachment] = []
    stored_paths: list[Path] = []
    try:
        for upload in files:
            attachment_id = uuid.uuid4()
            original_name = (upload.filename or "file").strip()
            extension = Path(original_name).suffix
            stored_name = f"{attachment_id}{extension}"
            stored_path = note_dir / stored_name

            size_bytes = 0
            with stored_path.open("wb") as buffer:
                while True:
                    chunk = await upload.read(1024 * 1024)
                    if not chunk:
                        break
                    size_bytes += len(chunk)
                    if size_bytes > max_bytes:
                        buffer.close()
                        if stored_path.exists():
                            stored_path.unlink()
                        await upload.close()
                        raise HTTPException(
                            status_code=status.HTTP_400_BAD_REQUEST,
                            detail=f"File too large. Max {settings.GA_NOTES_MAX_FILE_MB}MB.",
                        )
                    buffer.write(chunk)
            stored_paths.append(stored_path)

            attachment = GaNoteAttachment(
                id=attachment_id,
                note_id=note.id,
                original_filename=original_name,
                stored_filename=stored_name,
                content_type=upload.content_type,
                size_bytes=size_bytes,
                created_by=user.id,
            )
            db.add(attachment)
            created.append(attachment)
            await upload.close()

        await db.commit()
    except Exception:
        for path in stored_paths:
            if path.exists():
                path.unlink()
        await db.rollback()
        raise

    return [_attachment_out(a) for a in created]


@router.get("", response_model=list[GaNoteOut])
async def list_ga_notes(
    project_id: uuid.UUID | None = None,
    department_id: uuid.UUID | None = None,
    db: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
) -> list[GaNoteOut]:
    # Filter out closed notes that were closed more than 30 days ago
    closed_cutoff = datetime.utcnow() - timedelta(days=30)

    stmt = select(GaNote).options(selectinload(GaNote.attachments)).order_by(
        GaNote.updated_at.desc(),
        GaNote.created_at.desc(),
    )

    # Include all open notes; include closed notes only if recently closed
    stmt = stmt.where(
        or_(
            GaNote.status != GaNoteStatus.CLOSED,
            GaNote.completed_at.is_(None),
            GaNote.completed_at >= closed_cutoff,
        )
    )
    
    if project_id is not None:
        project = (await db.execute(select(Project).where(Project.id == project_id))).scalar_one_or_none()
        if project is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
        stmt = stmt.where(GaNote.project_id == project_id)
    elif department_id is not None:
        stmt = stmt.where(GaNote.department_id == department_id)

    notes = (await db.execute(stmt)).scalars().all()
    return [_note_out(n) for n in notes]


@router.post("", response_model=GaNoteOut, status_code=status.HTTP_201_CREATED)
async def create_ga_note(
    payload: GaNoteCreate,
    db: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
) -> GaNoteOut:
    if payload.project_id is None and payload.department_id is None:
        if user.role not in (UserRole.ADMIN, UserRole.MANAGER):
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="project_id or department_id required")

    project = None
    department_id = payload.department_id
    if payload.project_id is not None:
        project = (await db.execute(select(Project).where(Project.id == payload.project_id))).scalar_one_or_none()
        if project is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
        if project.department_id is not None:
            ensure_department_access(user, project.department_id)
        if department_id is not None and project.department_id != department_id:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Project department mismatch")
        department_id = project.department_id
    elif department_id is not None:
        ensure_department_access(user, department_id)


    note = GaNote(
        content=payload.content,
        created_by=payload.created_by or user.id,
        note_type=payload.note_type or GaNoteType.GA,
        status=payload.status or GaNoteStatus.OPEN,
        priority=payload.priority,
        start_date=payload.start_date,
        due_date=payload.due_date,
        completed_at=payload.completed_at,
        is_converted_to_task=payload.is_converted_to_task or False,
        is_discussed=payload.is_discussed or False,
        project_id=payload.project_id,
        department_id=department_id,
    )
    db.add(note)
    await db.commit()
    await db.refresh(note)
    return _note_out(note)


@router.patch("/{note_id}", response_model=GaNoteOut)
async def update_ga_note(
    note_id: uuid.UUID,
    payload: GaNoteUpdate,
    db: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
) -> GaNoteOut:
    note = (await db.execute(select(GaNote).where(GaNote.id == note_id))).scalar_one_or_none()
    if note is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="GA note not found")
    await _ensure_note_access(note, user, db)

    old_content = note.content
    if payload.content is not None:
        note.content = payload.content
    if payload.status is not None:
        note.status = payload.status
        if payload.status == GaNoteStatus.CLOSED:
            note.completed_at = note.completed_at or datetime.utcnow()
    if payload.priority is not None:
        note.priority = payload.priority
    if payload.is_converted_to_task is not None:
        note.is_converted_to_task = payload.is_converted_to_task
    if payload.is_discussed is not None:
        note.is_discussed = payload.is_discussed

    if payload.content is not None and payload.content != old_content:
        new_task_title = _ga_note_task_title(note.content)
        old_default_description = _ga_note_default_task_description(old_content)
        new_default_description = _ga_note_default_task_description(note.content)

        linked_tasks = (
            await db.execute(select(Task).where(Task.ga_note_origin_id == note.id))
        ).scalars().all()

        for task in linked_tasks:
            # Tasks created from GA/KA notes should always track the note title.
            # Keep the stored title aligned with the full note text, only truncating at the task schema limit.
            task.title = new_task_title
            record_title_strike_events(
                db,
                task_id=task.id,
                actor_user_id=user.id,
                before_title=old_content,
                after_title=note.content,
            )
            if task.description == old_default_description:
                task.description = new_default_description
                record_description_strike_events(
                    db,
                    task_id=task.id,
                    actor_user_id=user.id,
                    before_description=old_default_description,
                    after_description=new_default_description,
                )

    await db.commit()
    await db.refresh(note)
    return _note_out(note)


def _same_timestamp(left: datetime, right: datetime) -> bool:
    def as_utc(value: datetime) -> datetime:
        if value.tzinfo is None:
            return value.replace(tzinfo=timezone.utc)
        return value.astimezone(timezone.utc)

    return abs((as_utc(left) - as_utc(right)).total_seconds()) < 0.001


@router.patch("/{note_id}/task-bundle", response_model=GaNoteTaskBundleResponse)
async def update_ga_note_task_bundle(
    note_id: uuid.UUID,
    payload: GaNoteTaskBundleUpdate,
    db: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
) -> GaNoteTaskBundleResponse:
    """Atomically update a GA note and its independent per-assignee tasks.

    Shared task fields are applied to every active copy. Membership changes
    create/deactivate copies without changing the status or progress of people
    who remain assigned.
    """

    note = (
        await db.execute(
            select(GaNote)
            .options(selectinload(GaNote.attachments))
            .where(GaNote.id == note_id)
            .with_for_update()
        )
    ).scalar_one_or_none()
    if note is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="GA note not found")
    await _ensure_note_access(note, user, db)

    existing_tasks = (
        await db.execute(select(Task).where(Task.ga_note_origin_id == note.id).with_for_update())
    ).scalars().all()
    semantic_before = {task.id: task_semantic_state(task) for task in existing_tasks}
    affected_department_ids = {task.department_id for task in existing_tasks if task.department_id}
    if note.department_id:
        affected_department_ids.add(note.department_id)
    if payload.assignee_ids:
        affected_department_ids.update(
            department_id for department_id in (
                await db.execute(select(User.department_id).where(User.id.in_(payload.assignee_ids)))
            ).scalars().all() if department_id
        )
    await ensure_daily_baselines_for_departments(
        db, department_ids=affected_department_ids, actor=user,
        day=datetime.now(ZoneInfo(settings.REALIZATION_TIMEZONE)).date(),
    )

    if (
        payload.expected_updated_at is not None
        and note.updated_at is not None
        and not _same_timestamp(payload.expected_updated_at, note.updated_at)
    ):
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="This GA note was changed by another user. Reload it before saving.",
        )

    fields_set = getattr(payload, "model_fields_set", getattr(payload, "__fields_set__", set()))
    old_content = note.content
    old_project_id = note.project_id
    if "content" in fields_set:
        cleaned_content = (payload.content or "").strip()
        if not cleaned_content:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Note text cannot be empty")
        note.content = cleaned_content

    if "project_id" in fields_set:
        project = None
        if payload.project_id is not None:
            project = (await db.execute(select(Project).where(Project.id == payload.project_id))).scalar_one_or_none()
            if project is None:
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
        note.project_id = payload.project_id
        if project is not None:
            note.department_id = project.department_id

    if payload.assignee_ids is not None:
        if not payload.assignee_ids:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="A converted GA note must keep at least one task assignee",
            )
        try:
            reconcile_result = await reconcile_ga_note_task_assignees(
                db,
                note=note,
                desired_assignee_ids=payload.assignee_ids,
                actor_user_id=user.id,
            )
        except ValueError as exc:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc)) from exc
        active_tasks = reconcile_result.active_tasks
    else:
        active_tasks = (
            await db.execute(
                select(Task)
                .where(Task.ga_note_origin_id == note.id, Task.is_active.is_(True))
                .order_by(Task.created_at.asc(), Task.id.asc())
                .with_for_update()
            )
        ).scalars().all()
        for task in active_tasks:
            task.fast_task_group_id = None
        reconcile_result = None

    title_before = {
        task.id: old_content if "content" in fields_set else task.title
        for task in active_tasks
    }
    description_before = {task.id: task.description for task in active_tasks}

    title = _ga_note_task_title(note.content) if "content" in fields_set else None
    description_is_set = "description" in fields_set
    updated_count = apply_ga_note_shared_task_fields(
        active_tasks,
        title=title,
        description_is_set=description_is_set,
        description=payload.description,
    )
    if "project_id" in fields_set:
        for task in active_tasks:
            changed = False
            if task.project_id != payload.project_id:
                task.project_id = payload.project_id
                changed = True
            if project is not None and task.department_id != project.department_id:
                task.department_id = project.department_id
                changed = True
            if changed:
                updated_count += 1

    if payload.assignee_states is not None:
        try:
            from app.api.routers.tasks import _validate_waiting_confirmation_assignee

            status_before_by_task_id = {
                task.id: task.status.value if isinstance(task.status, TaskStatus) else str(task.status)
                for task in active_tasks
            }
            for item in payload.assignee_states:
                if item.status == TaskStatus.WAITING_CONFIRMATION:
                    await _validate_waiting_confirmation_assignee(db, item.confirmation_assignee_id)
            updated_count += apply_ga_note_assignee_execution_states(
                active_tasks,
                [
                    GaNoteAssigneeExecutionState(
                        assignee_id=item.assignee_id,
                        status=item.status,
                        confirmation_assignee_id=item.confirmation_assignee_id,
                        start_date=item.start_date,
                        due_date=item.due_date,
                        finish_period=item.finish_period,
                        is_deadline_important=item.is_deadline_important,
                        priority=item.priority,
                        is_bllok=item.is_bllok,
                        is_1h_report=item.is_1h_report,
                        is_r1=item.is_r1,
                        is_personal=item.is_personal,
                    )
                    for item in payload.assignee_states
                ],
            )
            tasks_by_assignee = {
                task.assigned_to: task
                for task in active_tasks
                if task.is_active and task.assigned_to is not None
            }
            today = datetime.now(ZoneInfo(settings.REALIZATION_TIMEZONE)).date()
            for item in payload.assignee_states:
                matching_task = tasks_by_assignee.get(item.assignee_id)
                if matching_task is None:
                    continue
                if status_before_by_task_id.get(matching_task.id) == item.status.value:
                    continue
                await upsert_explicit_task_daily_status(
                    db,
                    task_id=matching_task.id,
                    day_date=today,
                    status=item.status,
                    finish_period=item.finish_period.value if item.finish_period else None,
                )
        except ValueError as exc:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc)) from exc

    # Preserve the legacy default-description behavior when only note content
    # changed and the description was never customized.
    if "content" in fields_set and not description_is_set:
        old_default = _ga_note_default_task_description(old_content)
        new_default = _ga_note_default_task_description(note.content)
        for task in active_tasks:
            if task.description == old_default and task.description != new_default:
                task.description = new_default
                updated_count += 1

    for task in active_tasks:
        previous_title = title_before.get(task.id)
        current_title = note.content if "content" in fields_set else task.title
        if previous_title != current_title:
            record_title_strike_events(
                db,
                task_id=task.id,
                actor_user_id=user.id,
                before_title=previous_title,
                after_title=current_title,
            )
        previous_description = description_before.get(task.id)
        if previous_description != task.description:
            record_description_strike_events(
                db,
                task_id=task.id,
                actor_user_id=user.id,
                before_description=previous_description,
                after_description=task.description,
            )

    # Re-scheduling through the note editor should restore the task on each
    # assignee's weekly planner (clear stale task/project day exclusions).
    if payload.assignee_ids is not None or payload.assignee_states is not None:
        from app.api.routers.tasks import _clear_task_planner_exclusions_for_current_plan

        for task in active_tasks:
            if task.assigned_to is None:
                continue
            await _clear_task_planner_exclusions_for_current_plan(
                db,
                task=task,
                user_ids=[task.assigned_to],
            )

    semantic_tasks = {task.id: task for task in existing_tasks}
    semantic_tasks.update({task.id: task for task in active_tasks})
    if reconcile_result is not None:
        semantic_tasks.update({task.id: task for task in reconcile_result.deactivated_tasks})
    for task in semantic_tasks.values():
        after_state = task_semantic_state(task)
        before_state = semantic_before.get(task.id, {**after_state, "is_active": False, "assigned_to": None})
        old_owner = before_state.get("assigned_to") if before_state.get("is_active") else None
        new_owner = after_state.get("assigned_to") if after_state.get("is_active") else None
        record_task_semantic_events(
            db=db, task_id=task.id, actor_user_id=user.id,
            before=before_state, after=after_state,
            old_assignee_ids=[old_owner] if old_owner else [],
            new_assignee_ids=[new_owner] if new_owner else [],
        )

    before_assignees = None
    after_assignees = [str(task.assigned_to) for task in active_tasks if task.assigned_to]
    if payload.assignee_ids is not None:
        before_assignees = "reconciled"
    add_audit_log(
        db=db,
        actor_user_id=user.id,
        entity_type="ga_note",
        entity_id=note.id,
        action="task_bundle_updated",
        before={
            "content": old_content,
            "project_id": str(old_project_id) if old_project_id else None,
            "assignees": before_assignees,
        },
        after={
            "content": note.content,
            "project_id": str(note.project_id) if note.project_id else None,
            "assignees": after_assignees,
            "assignee_states": [
                {
                    "assignee_id": str(item.assignee_id),
                    "status": item.status.value,
                    "start_date": item.start_date.isoformat() if item.start_date else None,
                    "due_date": item.due_date.isoformat() if item.due_date else None,
                    "finish_period": item.finish_period.value if item.finish_period else None,
                    "priority": item.priority.value,
                    "is_bllok": item.is_bllok,
                    "is_1h_report": item.is_1h_report,
                    "is_r1": item.is_r1,
                    "is_personal": item.is_personal,
                }
                for item in (payload.assignee_states or [])
            ],
        },
    )

    created_notifications = []
    if reconcile_result is not None:
        for created_task in reconcile_result.created_tasks:
            if created_task.assigned_to is None:
                continue
            created_notifications.append(
                add_notification(
                    db=db,
                    user_id=created_task.assigned_to,
                    type=NotificationType.assignment,
                    title="Task assigned",
                    body=notification_task_preview(created_task.title),
                    data={"task_id": str(created_task.id)},
                )
            )

    await db.commit()
    for notification in created_notifications:
        try:
            await publish_notification(user_id=notification.user_id, notification=notification)
        except Exception:
            pass
    await db.refresh(note)
    return GaNoteTaskBundleResponse(
        note=_note_out(note),
        active_task_ids=[task.id for task in active_tasks],
        assignee_ids=[task.assigned_to for task in active_tasks if task.assigned_to is not None],
        created_count=reconcile_result.created_count if reconcile_result else 0,
        deactivated_count=reconcile_result.deactivated_count if reconcile_result else 0,
        deduplicated_count=reconcile_result.deduplicated_count if reconcile_result else 0,
        updated_count=updated_count,
    )


@router.patch("/{note_id}/task-deadline", response_model=GaNoteTaskDeadlineResponse)
async def update_ga_note_task_deadline(
    note_id: uuid.UUID,
    payload: GaNoteTaskDeadlineUpdate,
    db: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
) -> GaNoteTaskDeadlineResponse:
    """Reject the legacy global-date operation for independent GA task copies."""
    note = await _get_note_or_404(note_id, db)
    await _ensure_note_access(note, user, db)
    raise HTTPException(
        status_code=status.HTTP_409_CONFLICT,
        detail="GA task scheduling is per assignee; update assignee_states through the task bundle",
    )

    linked_tasks = (
        await db.execute(
            select(Task)
            .where(Task.ga_note_origin_id == note_id)
            .where(Task.is_active.is_(True))
        )
    ).scalars().all()

    if not linked_tasks:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No active tasks found for this GA/KA note",
        )

    new_start_date = None if payload.clear_start else payload.start_date
    update_start = payload.clear_start or payload.start_date is not None
    new_due_date = None if payload.clear else payload.due_date
    update_due = payload.clear or payload.due_date is not None
    update_important = payload.is_deadline_important is not None

    if not update_start and not update_due and not update_important:
        return GaNoteTaskDeadlineResponse(
            updated_count=0,
            start_date=linked_tasks[0].start_date,
            due_date=linked_tasks[0].due_date,
            is_deadline_important=linked_tasks[0].is_deadline_important,
        )

    updated_count = 0
    for task in linked_tasks:
        before = {
            "start_date": task.start_date.isoformat() if task.start_date else None,
            "due_date": task.due_date.isoformat() if task.due_date else None,
            "is_deadline_important": task.is_deadline_important,
        }
        changed = False

        if update_start and task.start_date != new_start_date:
            task.start_date = new_start_date
            changed = True

        if update_due:
            if (
                task.due_date is not None
                and new_due_date is not None
                and new_due_date != task.due_date
                and task.original_due_date is None
            ):
                task.original_due_date = task.due_date
            if task.due_date != new_due_date:
                task.due_date = new_due_date
                changed = True

        if update_important and task.is_deadline_important != payload.is_deadline_important:
            task.is_deadline_important = bool(payload.is_deadline_important)
            changed = True

        if changed:
            updated_count += 1
            after = {
                "start_date": task.start_date.isoformat() if task.start_date else None,
                "due_date": task.due_date.isoformat() if task.due_date else None,
                "is_deadline_important": task.is_deadline_important,
            }
            add_audit_log(
                db=db,
                actor_user_id=user.id,
                entity_type="task",
                entity_id=task.id,
                action="ga_note_deadline_update",
                before=before,
                after=after,
            )

    await db.commit()

    sample = linked_tasks[0]
    return GaNoteTaskDeadlineResponse(
        updated_count=updated_count,
        start_date=sample.start_date,
        due_date=sample.due_date,
        is_deadline_important=sample.is_deadline_important,
    )


@router.post("/{note_id}/mark-waiting-done", response_model=MarkWaitingDoneResponse)
async def mark_note_waiting_tasks_done(
    note_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
) -> MarkWaitingDoneResponse:
    note = await _get_note_or_404(note_id, db)
    await _ensure_note_access(note, user, db)

    tasks = (
        await db.execute(
            select(Task).where(Task.ga_note_origin_id == note_id)
        )
    ).scalars().all()

    waiting = [task for task in tasks if task.status == TaskStatus.WAITING_CONFIRMATION]
    comments = {
        row.task_id: (row.comment or "").strip()
        for row in (
            await db.execute(
                select(TaskUserComment).where(
                    TaskUserComment.task_id.in_([task.id for task in waiting]),
                    TaskUserComment.user_id == user.id,
                )
            )
        ).scalars().all()
    } if waiting else {}
    if any(task.assigned_to != user.id or not comments.get(task.id) for task in waiting):
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail={
                "code": "RLZ_TASK_COMMENT_REQUIRED",
                "message": "Para mbylljes, shto komentin personal për secilën detyrë; përjashtimet mbyllen individualisht me arsye.",
            },
        )

    updated_count = 0
    for task in tasks:
        if task.status == TaskStatus.WAITING_CONFIRMATION:
            task.status = TaskStatus.DONE
            task.completed_at = task.completed_at or datetime.now(timezone.utc)
            await upsert_explicit_task_daily_status(
                db,
                task_id=task.id,
                day_date=datetime.now(timezone.utc).date(),
                status=TaskStatus.DONE,
                finish_period=task.finish_period.value if hasattr(task.finish_period, "value") else task.finish_period,
            )
            updated_count += 1

    await db.commit()

    return MarkWaitingDoneResponse(
        updated_count=updated_count,
        skipped_count=len(tasks) - updated_count,
    )


@router.get("/{note_id}", response_model=GaNoteOut)
async def get_ga_note(
    note_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
) -> GaNoteOut:
    note = (
        await db.execute(
            select(GaNote)
            .options(selectinload(GaNote.attachments))
            .where(GaNote.id == note_id)
        )
    ).scalar_one_or_none()
    if note is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="GA note not found")
    await _ensure_note_access(note, user, db)
    return _note_out(note)


@router.post("/{note_id}/attachments", response_model=list[GaNoteAttachmentOut], status_code=status.HTTP_201_CREATED)
async def upload_ga_note_attachments(
    note_id: uuid.UUID,
    files: list[UploadFile] = File(...),
    db: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
) -> list[GaNoteAttachmentOut]:
    note = await _get_note_or_404(note_id, db)
    return await _save_ga_note_attachments(note, files, db, user)


@router.post("/attachments", response_model=list[GaNoteAttachmentOut], status_code=status.HTTP_201_CREATED)
async def upload_ga_note_attachments_by_form(
    note_id: uuid.UUID = Form(...),
    files: list[UploadFile] = File(...),
    db: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
) -> list[GaNoteAttachmentOut]:
    note = await _get_note_or_404(note_id, db)
    return await _save_ga_note_attachments(note, files, db, user)


@router.get("/attachments/{attachment_id}/preview")
async def preview_ga_note_attachment(
    attachment_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
):
    attachment = (
        await db.execute(
            select(GaNoteAttachment)
            .options(selectinload(GaNoteAttachment.note))
            .where(GaNoteAttachment.id == attachment_id)
        )
    ).scalar_one_or_none()
    if attachment is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Attachment not found")

    await _ensure_note_access(attachment.note, user, db)

    stored_path = _ga_note_upload_base_dir() / str(attachment.note_id) / attachment.stored_filename
    if not stored_path.exists():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found on server")

    content_type = attachment.content_type or mimetypes.guess_type(attachment.original_filename)[0] or "application/octet-stream"
    suffix = stored_path.suffix.lower()
    if content_type.startswith(("image/", "audio/", "video/")) or content_type == "application/pdf" or suffix == ".pdf":
        return FileResponse(stored_path, media_type=content_type)

    try:
        preview_html = await run_in_threadpool(
            _build_attachment_preview,
            stored_path,
            attachment.original_filename,
        )
    except Exception:
        preview_html = _preview_page(
            attachment.original_filename,
            '<p>The preview could not be generated.</p><p class="muted">Use Download to open the original file.</p>',
        )
    return HTMLResponse(preview_html)


@router.get("/attachments/{attachment_id}")
async def download_ga_note_attachment(
    attachment_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
):
    attachment = (
        await db.execute(
            select(GaNoteAttachment)
            .options(selectinload(GaNoteAttachment.note))
            .where(GaNoteAttachment.id == attachment_id)
        )
    ).scalar_one_or_none()
    if attachment is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Attachment not found")

    await _ensure_note_access(attachment.note, user, db)

    upload_base = _ga_note_upload_base_dir()
    stored_path = upload_base / str(attachment.note_id) / attachment.stored_filename
    if not stored_path.exists():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found on server")

    return FileResponse(
        stored_path,
        media_type=attachment.content_type or "application/octet-stream",
        filename=attachment.original_filename,
    )


@router.delete("/attachments/{attachment_id}", status_code=status.HTTP_204_NO_CONTENT, response_model=None)
async def delete_ga_note_attachment(
    attachment_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
):
    attachment = (
        await db.execute(
            select(GaNoteAttachment)
            .options(selectinload(GaNoteAttachment.note))
            .where(GaNoteAttachment.id == attachment_id)
        )
    ).scalar_one_or_none()
    if attachment is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Attachment not found")

    await _ensure_note_access(attachment.note, user, db)

    upload_base = _ga_note_upload_base_dir()
    note_dir = upload_base / str(attachment.note_id)
    stored_path = note_dir / attachment.stored_filename

    await db.delete(attachment)
    await db.commit()

    try:
        if stored_path.exists():
            stored_path.unlink()
        if note_dir.exists() and not any(note_dir.iterdir()):
            note_dir.rmdir()
    except OSError:
        pass
